""".docx bytes → (DocumentModel, LossReport). The read half of the docx writer.

This module is the SECOND place OOXML wordprocessingml vocabulary is allowed, and for the
same reason as the first: inverting a file format is the parser's whole job. The contract
is one-directional — vendor vocabulary comes IN, only `documents/model.py` types go OUT.
Nothing downstream of `parse_docx` ever sees a `w:` tag.

(Note the package docstring's "no vendor file-format vocabulary appears outside
``writers/``" predates a reader existing at all. The honest statement of the rule is
"vendor vocabulary is confined to the format-specific edge modules"; a future `readers/`
subpackage is where this belongs once there is more than one.)

Two things make a parser trustworthy, and both are contracts, not effort:

**Order is preserved by construction.** python-docx exposes `doc.paragraphs` and
`doc.tables` as two independent sequences, so any parser that iterates them in turn emits
every paragraph before every table and silently reorders the document. This one walks the
body element's own children, which IS document order.

**Every loss is named.** The model deliberately describes a small, portable document; a
real .docx holds far more (footnotes, comments, merged cells, tracked changes, text
boxes, arbitrary character formatting). Dropping those quietly is the failure mode that
makes a "fidelity" editor untrustworthy, so each one appends a `LossItem` saying WHAT was
dropped and WHERE. A construct the model cannot express and that produces no item is a
bug in this module, not an acceptable simplification.

## The writer's conventions, inverted

`writers/docx_writer.py` expresses two model properties through a convention rather than
a dedicated OOXML feature, because the default template cannot be relied on to define the
styles that would otherwise carry them. Each inverse below is therefore load-bearing: get
it wrong and a document THIS repo wrote fails to round-trip.

* A `code` run is written as the font `Courier New` (no code character style exists), so
  a run in that font reads back as `Run(code=True)`. A run in any *other* explicit font
  is a `run_property` loss — the model has no font field.
* A hyperlink is written as a real `w:hyperlink` plus an explicit colour `0563C1` and a
  single underline (the `Hyperlink` character style is not guaranteed to exist). Inside a
  hyperlink those two exact values are therefore the link's own decoration and are
  CONSUMED by `Run.link`; any other colour or underline style inside a link, and any
  colour or underline outside one, is reported.
* An `image` block is written as the placeholder paragraph ``[image: <slug>]`` (the block
  references an artifact and the writer has no store), so a lone unformatted paragraph of
  exactly that shape reads back as an `image` block. A user's literal ``[image: x]``
  paragraph is reclassified by this rule; the text survives byte-identically either way,
  which is why the round trip is stable regardless.
* `code` vs `paragraph` is genuinely ambiguous in the output: a `code` block and a
  `paragraph` whose every run is `code=True` render to the same OOXML. This parser reads
  an all-monospace paragraph as a `code` block, which is the inverse of the writer's
  plain-text path; the alternative reading loses nothing either, and both are stable
  under re-render.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any

from personalclaw.documents.model import (
    Block,
    Cell,
    DocumentModel,
    PageSetup,
    ParagraphStyle,
    Run,
)

#: An lxml element or a python-docx object. `Any` rather than a precise type: the ones
#: python-docx exposes (`ProvidesStoryPart` and friends) are its internals, and naming
#: them here would pull vendor types into a module whose whole contract is that only
#: `documents/model.py` types come OUT. These handles never escape this file.
_Xml = Any

#: Must match `writers/docx_writer.py`. Duplicated rather than imported: the writer is a
#: pure model→bytes function and importing it here to read two constants would couple the
#: read path to the write path's import graph (and its optional dependency) for nothing.
_MONOSPACE = "Courier New"
_LINK_COLOR = "0563C1"

#: The writer builds on python-docx's default template, which is US Letter. A parsed page
#: of any other size cannot be expressed — `PageSetup` carries orientation, not size.
_TEMPLATE_PAGE_IN = {8.5, 11.0}

_IMAGE_PLACEHOLDER = re.compile(r"^\[image: (.*)\]$")
_HEADING_STYLE = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)

#: Paragraph styles this parser CONSUMES into the model (as a kind, a level, or the
#: document title). Any other named style is a `paragraph_style` loss.
_CONSUMED_STYLES = {"normal", "title", "list bullet", "list number", "list paragraph"}

#: The table style the writer always emits. Any other is a `table_style` loss.
_WRITER_TABLE_STYLE = "table grid"

#: Character properties that carry visible formatting the model has no field for.
#: A DENYLIST, deliberately: an allowlist over the full `w:rPr` vocabulary would report
#: `w:lang`, `w:noProof` and revision ids on nearly every Word-authored run, and a loss
#: report that cries wolf is one nobody reads. Anything visible belongs in this map.
_RUN_PROPS = {
    "color": "text colour",
    "u": "underline",
    "sz": "font size",
    "szCs": "complex-script font size",
    "strike": "strikethrough",
    "dstrike": "double strikethrough",
    "highlight": "highlight",
    "shd": "run shading",
    "vertAlign": "superscript/subscript",
    "smallCaps": "small caps",
    "caps": "all caps",
    "rStyle": "character style",
    "spacing": "character spacing",
    "position": "raised/lowered text",
    "bdr": "character border",
    "outline": "outline",
    "emboss": "emboss",
    "imprint": "imprint",
    "vanish": "hidden text",
    "w": "character scaling",
}

#: Paragraph properties with no model field. Same denylist reasoning as `_RUN_PROPS`;
#: `w:widowControl`, `w:contextualSpacing` and revision ids are omitted as invisible.
_PARA_PROPS = {
    "ind": "indentation",
    "pBdr": "paragraph border",
    "shd": "paragraph shading",
    "tabs": "tab stops",
    "pageBreakBefore": "page-break-before",
    "keepNext": "keep with next",
    "keepLines": "keep lines together",
    "framePr": "text frame",
    "textDirection": "text direction",
    "outlineLvl": "outline level",
    "bidi": "right-to-left paragraph",
    "numPr": "list numbering",
}

#: OOXML element local names that ARE a construct the model cannot hold, mapped to the
#: loss kind they raise. Matched on local name rather than a fully-qualified tag so a
#: construct wrapped in a different namespace (`mc:AlternateContent`, VML, DrawingML)
#: still registers; the cost is that an unrelated namespace reusing one of these names
#: would report a phantom, which no known wordprocessing namespace does.
_CONSTRUCTS = {
    "footnoteReference": "footnote",
    "endnoteReference": "endnote",
    "commentReference": "comment",
    "commentRangeStart": "comment",
    "ins": "tracked_change",
    "del": "tracked_change",
    "fldSimple": "field",
    "fldChar": "field",
    "instrText": "field",
    "bookmarkStart": "bookmark",
    "oMath": "math",
    "oMathPara": "math",
    "object": "embedded_object",
    "sdt": "content_control",
}

#: Every loss a parse can report. A closed vocabulary so a caller can branch on a kind
#: (and a test can assert one construct family produced its item) instead of matching
#: prose. `detail` carries the specifics.
LOSS_KINDS = (
    "footnote",
    "endnote",
    "comment",
    "nested_table",
    "text_box",
    "merged_cells",
    "multi_section",
    "header_footer",
    "embedded_image",
    "embedded_object",
    "tracked_change",
    "field",
    "bookmark",
    "math",
    "content_control",
    "internal_link",
    "run_property",
    "explicit_off_toggle",
    "paragraph_property",
    "paragraph_style",
    "table_style",
    "line_break",
    "tab",
    "line_spacing_exact",
    "page_property",
    "list_item_formatting",
    "nested_list_level",
    "heading_level_clamped",
    "multi_paragraph_cell",
    "inline_page_break",
)


@dataclass(frozen=True)
class LossItem:
    """One thing the .docx said and the model cannot hold.

    Frozen because a report is evidence: a caller that could edit an item in place could
    make the report disagree with the file it describes.
    """

    kind: str
    detail: str
    #: Index into `DocumentModel.blocks` of the block this loss belongs to; -1 when the
    #: loss is document-level (page setup, sections, headers) and belongs to no block.
    block_index: int = -1
    #: 0-based ordinal of the source paragraph among ALL body paragraphs, including ones
    #: that merged into a list block or were consumed as the title. -1 when not
    #: paragraph-scoped. Kept alongside `block_index` because several paragraphs can
    #: collapse into one block, and "which bullet" is the answer a user needs.
    paragraph_ordinal: int = -1

    def __post_init__(self) -> None:
        if self.kind not in LOSS_KINDS:
            raise ValueError(f"unknown loss kind {self.kind!r}; expected one of {LOSS_KINDS}")

    @property
    def where(self) -> str:
        parts = []
        if self.block_index >= 0:
            parts.append(f"block {self.block_index}")
        if self.paragraph_ordinal >= 0:
            parts.append(f"paragraph {self.paragraph_ordinal}")
        return ", ".join(parts) or "document"

    def __str__(self) -> str:
        return f"{self.kind} at {self.where}: {self.detail}"


@dataclass
class LossReport:
    """Everything a parse could not represent, in document order."""

    items: list[LossItem] = field(default_factory=list)

    def add(
        self,
        kind: str,
        detail: str,
        *,
        block_index: int = -1,
        paragraph_ordinal: int = -1,
    ) -> None:
        self.items.append(
            LossItem(
                kind=kind,
                detail=detail,
                block_index=block_index,
                paragraph_ordinal=paragraph_ordinal,
            )
        )

    @property
    def lossless(self) -> bool:
        """True when the model holds everything the file said.

        Named for the QUESTION a caller asks ("can I edit and rewrite this safely?")
        rather than as `__bool__` on the report, because `if report:` reads as "there is a
        report" and would invert the meaning at exactly the call site that matters.
        """
        return not self.items

    def kinds(self) -> list[str]:
        """The distinct kinds present, in `LOSS_KINDS` order (stable, not insertion)."""
        present = {item.kind for item in self.items}
        return [kind for kind in LOSS_KINDS if kind in present]

    def of_kind(self, kind: str) -> list[LossItem]:
        if kind not in LOSS_KINDS:
            raise ValueError(f"unknown loss kind {kind!r}; expected one of {LOSS_KINDS}")
        return [item for item in self.items if item.kind == kind]

    def summary(self) -> str:
        if not self.items:
            return "no losses"
        counts = {kind: len(self.of_kind(kind)) for kind in self.kinds()}
        return ", ".join(f"{kind}×{count}" for kind, count in counts.items())

    def to_dict(self) -> dict[str, object]:
        """JSON-ready view for the surface that must warn before an edit (§C5).

        Carries the DERIVED answers (``lossless``, ``kinds``, ``summary``) beside the
        items rather than leaving a client to re-derive them. A report is evidence, and
        the verdict a user is shown must be the one this module computes — a frontend
        that re-implemented ``lossless`` as ``items.length === 0`` would be right today
        and wrong the first time a purely informational item is added.
        """
        return {
            "lossless": self.lossless,
            "kinds": self.kinds(),
            "summary": self.summary(),
            "items": [
                {
                    "kind": item.kind,
                    "detail": item.detail,
                    "block_index": item.block_index,
                    "paragraph_ordinal": item.paragraph_ordinal,
                    "where": item.where,
                }
                for item in self.items
            ],
        }


def parse_docx(data: bytes) -> tuple[DocumentModel, LossReport]:
    """Parse .docx bytes into a `DocumentModel` and the report of what did not fit.

    Raises whatever python-docx raises for a file that is not a .docx (a corrupt package
    is not a "loss" — there is no document to report about, and swallowing it would hand
    the caller an empty model that looks like an empty document).
    """
    from docx import Document

    return _Parser(Document(io.BytesIO(data))).run()


def _local(element: _Xml) -> str:
    """The local name of an lxml element's tag, namespace stripped."""
    return str(getattr(element, "tag", "")).rsplit("}", 1)[-1]


def _attr(element: _Xml, name: str) -> str | None:
    """A `w:`-namespaced attribute, by local name."""
    items = getattr(element, "attrib", {})
    for key, value in items.items():
        if str(key).rsplit("}", 1)[-1] == name:
            return str(value)
    return None


def _is_off(element: _Xml) -> bool:
    """True when a toggle element (`w:b`, `w:i`) is explicitly switched OFF."""
    return _attr(element, "val") in ("0", "false", "off")


class _Parser:
    """One parse. Holds the growing block list and the report so every emitter can name
    the block index a loss belongs to without threading it through ten signatures."""

    def __init__(self, doc: _Xml) -> None:
        self._doc = doc
        self.report = LossReport()
        self._blocks: list[Block] = []
        self._title = ""
        self._ordinal = -1  # incremented per body paragraph, INCLUDING consumed ones
        self._pending_list: Block | None = None

    # -- entry point ---------------------------------------------------------------

    def run(self) -> tuple[DocumentModel, LossReport]:
        page = self._page()
        self._walk(self._doc.element.body)
        self._flush_list()
        model = DocumentModel(title=self._title, blocks=self._blocks, page=page)
        return model, self.report

    @property
    def _next_index(self) -> int:
        """The index the block being built will occupy.

        A pending list block has not been appended yet but will land at exactly this
        index, because nothing else can be appended before it is flushed.
        """
        return len(self._blocks)

    # -- document order ------------------------------------------------------------

    def _walk(self, parent: _Xml) -> None:
        """Dispatch the body's own children, in document order.

        This walk IS the order guarantee. `doc.paragraphs` and `doc.tables` are separate
        sequences, so consuming them in turn would emit an interleaved document as "all
        paragraphs, then all tables" — the defect this method exists to prevent.
        """
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        for child in parent.iterchildren():
            name = _local(child)
            if name == "p":
                self._paragraph(Paragraph(child, self._doc))
            elif name == "tbl":
                self._flush_list()
                self._table(Table(child, self._doc))
            elif name == "sdt":
                # A block-level content control. Reported, then DESCENDED INTO: its
                # paragraphs are the user's content and skipping the wrapper would drop
                # them, which is the one thing a loss report must never let happen.
                self.report.add(
                    "content_control",
                    "block-level content control; its wrapper is dropped, content kept",
                    block_index=self._next_index,
                )
                for content in child.iterchildren():
                    if _local(content) == "sdtContent":
                        self._walk(content)

    # -- page ----------------------------------------------------------------------

    def _page(self) -> PageSetup:
        """Page geometry from the FIRST section (the one that governs page 1).

        Always returns a `PageSetup` rather than None: the file genuinely declares a
        geometry, so reporting "the writer's default" would discard a fact we read.
        """
        sections = list(self._doc.sections)
        if len(sections) > 1:
            self.report.add(
                "multi_section",
                f"{len(sections)} sections; the model holds one page setup "
                "(read from the first section)",
            )
        section = sections[0]
        width = float(section.page_width.inches) if section.page_width else 0.0
        height = float(section.page_height.inches) if section.page_height else 0.0
        if width and height and {round(width, 2), round(height, 2)} != _TEMPLATE_PAGE_IN:
            self.report.add(
                "page_property",
                f"page size {width:.2f}x{height:.2f}in; the model carries orientation "
                "only, so a re-render uses the template's Letter page",
            )
        orientation = "landscape" if width > height else "portrait"
        margins = [
            float(getattr(section, name).inches) if getattr(section, name) else 0.0
            for name in ("top_margin", "bottom_margin", "left_margin", "right_margin")
        ]
        margin = margins[2]  # left: the one a reader notices
        if len(set(round(m, 4) for m in margins)) > 1:
            self.report.add(
                "page_property",
                "margins differ (top/bottom/left/right = "
                + "/".join(f"{m:.2f}" for m in margins)
                + f"in); the model holds one value, using left ({margin:.2f}in)",
            )
        elif margin == 0.0:
            self.report.add(
                "page_property",
                "zero page margins; 0.0 means 'writer default' in the model, so a "
                "genuinely zero-margin page cannot be expressed",
            )
        self._headers_footers(sections)
        return PageSetup(orientation=orientation, margin_in=margin)

    def _headers_footers(self, sections: list) -> None:
        for number, section in enumerate(sections):
            for name in ("header", "footer"):
                part = getattr(section, name, None)
                if part is None:
                    continue
                text = " ".join(p.text for p in part.paragraphs).strip()
                if text:
                    self.report.add(
                        "header_footer",
                        f"section {number} {name}: {text!r}; the model has no " f"{name} field",
                    )

    # -- paragraphs ----------------------------------------------------------------

    def _paragraph(self, para: _Xml) -> None:
        self._ordinal += 1
        ordinal = self._ordinal
        index = self._next_index
        self._constructs(para._p, index, ordinal)
        style_name = _style_name(para)
        runs = self._runs(para, index, ordinal)
        breaks = _page_breaks(para._p)
        text = "".join(run.text for run in runs)

        if breaks and not text:
            self._flush_list()
            self._para_props(para, index, ordinal, consumed_numbering=False)
            self._blocks.append(Block(kind="pagebreak"))
            return
        if breaks:
            self.report.add(
                "inline_page_break",
                f"{breaks} page break(s) share a paragraph with text; the model's "
                "pagebreak is a block of its own, so the break is dropped",
                block_index=index,
                paragraph_ordinal=ordinal,
            )

        lowered = style_name.lower()
        heading = _HEADING_STYLE.match(style_name)
        list_kind = self._list_kind(para, lowered)

        if list_kind:
            self._para_props(para, index, ordinal, consumed_numbering=True)
            self._list_item(list_kind, para, runs, text, index, ordinal)
            return

        self._flush_list()
        index = self._next_index
        self._para_props(para, index, ordinal, consumed_numbering=False)
        style = self._style(para, index, ordinal)

        if lowered == "title":
            if not self._blocks and not self._title:
                self._title = text
                return
            # A second Title-styled paragraph has nowhere to go: the model holds one
            # title. Kept as a top-level heading so the text survives, and reported so
            # the demotion is visible.
            self.report.add(
                "paragraph_style",
                "'Title' style on a paragraph that is not the document title; kept as "
                "a level-1 heading",
                block_index=index,
                paragraph_ordinal=ordinal,
            )
            self._blocks.append(self._text_block("heading", runs, text, style=style, level=1))
            return

        if heading:
            level = int(heading.group(1))
            if level > 6:
                self.report.add(
                    "heading_level_clamped",
                    f"heading level {level} clamped to 6 (the model's maximum)",
                    block_index=index,
                    paragraph_ordinal=ordinal,
                )
            self._blocks.append(self._text_block("heading", runs, text, style=style, level=level))
            return

        if lowered not in _CONSUMED_STYLES:
            self.report.add(
                "paragraph_style",
                f"paragraph style {style_name!r} has no model field; kept as a plain " "paragraph",
                block_index=index,
                paragraph_ordinal=ordinal,
            )

        image = _IMAGE_PLACEHOLDER.match(text)
        if image and len(runs) == 1 and _plain(runs):
            self._blocks.append(Block(kind="image", artifact_slug=image.group(1)))
            return
        if runs and all(run.code for run in runs):
            # Every run monospace: the writer's `code` block. See the module docstring
            # for why this reading is chosen and why the alternative is also stable.
            self._blocks.append(self._text_block("code", runs, text, style=style))
            return
        self._blocks.append(self._text_block("paragraph", runs, text, style=style))

    def _text_block(
        self,
        kind: str,
        runs: list[Run],
        text: str,
        *,
        style: ParagraphStyle | None,
        level: int = 1,
    ) -> Block:
        """A text-bearing block, rich only when the formatting needs it.

        An unformatted paragraph becomes plain `text` with no `runs`: carrying a single
        featureless Run would make every parsed document look formatted, and the writer
        renders both identically. A `code` block's monospace is already carried by the
        KIND, so an otherwise-plain code block is plain too.
        """
        rich = not _plain(runs, ignore_code=kind == "code")
        return Block(
            kind=kind,
            text="" if rich else text,
            level=level,
            runs=list(runs) if rich else [],
            style=style,
        )

    # -- lists ---------------------------------------------------------------------

    def _list_kind(self, para: _Xml, lowered: str) -> str:
        """ "bullets" / "numbered" / "" for a paragraph that is not a list item.

        A resolvable `w:numPr` first, then the style name (what this repo's writer emits —
        "List Bullet" / "List Number" with no numbering reference at all, so a
        numbering-only check would read the writer's own lists as plain paragraphs).

        The direct reference outranks the style because that is the OOXML precedence Word
        itself renders: a paragraph's own `w:numPr` overrides whatever numbering its style
        supplies. A real Word document disagrees with itself here routinely — typing at the
        end of a numbered list and clicking the bullet button leaves the new paragraph
        styled "List Number" with a bullet `w:numPr` on it, and Word shows a bullet.
        Reading the style name first turned those bullets into extra items on the numbered
        list above them; `tests/test_docx_word_authored.py` owns that case on a file Word
        actually saved.
        """
        # `w:numPr` is a child of `w:pPr`, never of `w:p` — reading it off the paragraph
        # element finds nothing and reads every Word list as plain paragraphs.
        properties = _find(para._p, "pPr")
        num_pr = _find(properties, "numPr") if properties is not None else None
        if num_pr is not None:
            level = _find(num_pr, "ilvl")
            if level is not None and (_attr(level, "val") or "0") != "0":
                self.report.add(
                    "nested_list_level",
                    f"list item at indent level {_attr(level, 'val')}; the model's items "
                    "are a flat list of strings",
                    block_index=self._next_index,
                    paragraph_ordinal=self._ordinal,
                )
            fmt = self._num_format(num_pr)
            if fmt:
                return "bullets" if fmt == "bullet" else "numbered"
        if lowered == "list bullet":
            return "bullets"
        if lowered == "list number":
            return "numbered"
        # A numbering reference whose definition we could not read is still a list item.
        # `numbered` rather than `bullets` keeps the more visible wrong answer off the
        # table: stripping the numbers off a numbered list.
        return "numbered" if num_pr is not None else ""

    def _num_format(self, num_pr: _Xml) -> str:
        """The `w:numFmt` behind a `w:numPr`, or "" when it cannot be resolved.

        Unresolvable is not an error: a numbering definition can live in a part this
        document does not carry, and guessing "bullet" there would silently renumber a
        numbered list. The caller treats anything but "bullet" as numbered.
        """
        num_id = _find(num_pr, "numId")
        target = _attr(num_id, "val") if num_id is not None else None
        if not target:
            return ""
        try:
            numbering = self._doc.part.numbering_part.element
        except Exception:  # noqa: BLE001 — no numbering part is a normal document shape
            return ""
        abstract = ""
        for num in numbering.iterchildren():
            if _local(num) != "num" or _attr(num, "numId") != target:
                continue
            ref = _find(num, "abstractNumId")
            abstract = (_attr(ref, "val") or "") if ref is not None else ""
        if not abstract:
            return ""
        for definition in numbering.iterchildren():
            if _local(definition) != "abstractNum":
                continue
            if _attr(definition, "abstractNumId") != abstract:
                continue
            for level in definition.iterchildren():
                if _local(level) != "lvl":
                    continue
                if (_attr(level, "ilvl") or "0") != "0":
                    continue
                fmt = _find(level, "numFmt")
                return (_attr(fmt, "val") or "") if fmt is not None else ""
        return ""

    def _list_item(
        self,
        kind: str,
        para: _Xml,
        runs: list[Run],
        text: str,
        index: int,
        ordinal: int,
    ) -> None:
        if not _plain(runs):
            self.report.add(
                "list_item_formatting",
                "formatted runs in a list item; the model's `items` are plain strings, "
                f"so the formatting of {text!r} is dropped",
                block_index=index,
                paragraph_ordinal=ordinal,
            )
        if self._pending_list is not None and self._pending_list.kind != kind:
            self._flush_list()
        if self._pending_list is None:
            # The style of the FIRST item owns the whole block: the model has one style
            # per block and the writer applies it to every item, so a later item's
            # divergence is reported rather than silently winning.
            self._pending_list = Block(kind=kind, items=[], style=self._style(para, index, ordinal))
        else:
            later = self._style(para, index, ordinal)
            if later != self._pending_list.style:
                self.report.add(
                    "paragraph_property",
                    "list item styled differently from the first item in its list; the "
                    "model holds one style per block, so the first item's wins",
                    block_index=index,
                    paragraph_ordinal=ordinal,
                )
        self._pending_list.items.append(text)

    def _flush_list(self) -> None:
        if self._pending_list is not None:
            self._blocks.append(self._pending_list)
            self._pending_list = None

    # -- tables --------------------------------------------------------------------

    def _table(self, table: _Xml) -> None:
        index = self._next_index
        style = getattr(table, "style", None)
        name = getattr(style, "name", "") or ""
        if name.lower() != _WRITER_TABLE_STYLE:
            self.report.add(
                "table_style",
                f"table style {name!r}; the model has no table style and a re-render "
                f"uses '{_WRITER_TABLE_STYLE}'",
                block_index=index,
            )
        rows: list[list[Cell]] = []
        spans: list[str] = []
        for row_number, row in enumerate(table._tbl.iterchildren()):
            if _local(row) != "tr":
                continue
            cells: list[Cell] = []
            for column, tc in enumerate(row.iterchildren()):
                if _local(tc) != "tc":
                    continue
                span = self._span(tc)
                if span:
                    spans.append(f"r{row_number}c{column} {span}")
                cells.append(self._cell(tc, index, row_number, column))
            rows.append(cells)
        if spans:
            self.report.add(
                "merged_cells",
                "merged cells (" + "; ".join(spans) + "); the model's table is a plain "
                "grid, so the spans are dropped and short rows padded",
                block_index=index,
            )
        width = max((len(row) for row in rows), default=0)
        for row in rows:
            # Padded to the widest row, exactly as the writer normalizes on the way out.
            # Without this a merged-cell table would parse ragged and re-parse padded,
            # which would make the round trip disagree with itself.
            row.extend(Cell() for _ in range(width - len(row)))
        self._blocks.append(Block(kind="table", cells=rows))

    def _span(self, tc: _Xml) -> str:
        properties = _find(tc, "tcPr")
        if properties is None:
            return ""
        parts = []
        grid = _find(properties, "gridSpan")
        if grid is not None:
            parts.append(f"gridSpan={_attr(grid, 'val')}")
        vertical = _find(properties, "vMerge")
        if vertical is not None:
            parts.append(f"vMerge={_attr(vertical, 'val') or 'continue'}")
        return " ".join(parts)

    def _cell(self, tc: _Xml, index: int, row: int, column: int) -> Cell:
        from docx.text.paragraph import Paragraph

        for child in tc.iterchildren():
            if _local(child) == "tbl":
                self.report.add(
                    "nested_table",
                    f"nested table in cell r{row}c{column}; the model's cells hold runs, "
                    "not blocks, so the inner table is dropped",
                    block_index=index,
                )
        paragraphs = [
            Paragraph(child, self._doc) for child in tc.iterchildren() if _local(child) == "p"
        ]
        if len(paragraphs) > 1:
            self.report.add(
                "multi_paragraph_cell",
                f"{len(paragraphs)} paragraphs in cell r{row}c{column}; the model's cell "
                "is one paragraph, so they are joined",
                block_index=index,
            )
        runs: list[Run] = []
        align = ""
        for number, para in enumerate(paragraphs):
            self._constructs(para._p, index, -1)
            runs.extend(self._runs(para, index, -1))
            style_name = _style_name(para).lower()
            if style_name and style_name != "normal":
                self.report.add(
                    "paragraph_style",
                    f"cell paragraph style {_style_name(para)!r} in cell r{row}c{column} "
                    "has no model field",
                    block_index=index,
                )
            if number == 0:
                align = self._align(para, index, -1)
        runs = _merge(runs)
        rich = not _plain(runs)
        # `bold` stays False and bold lives on the RUNS: the writer bolds every header
        # cell by contract, so a cell-level flag read back from row 0 would be
        # indistinguishable from the convention. Run-level bold survives either way.
        return Cell(
            runs=runs if rich else [],
            text="" if rich else "".join(run.text for run in runs),
            align=align,
        )

    # -- runs ----------------------------------------------------------------------

    def _runs(self, para: _Xml, index: int, ordinal: int) -> list[Run]:
        """Model runs for one paragraph, in document order, adjacent equals merged.

        `iter_inner_content()` is what makes a hyperlink's position survive: a
        `w:hyperlink` is not a `w:r` child of the paragraph, so `paragraph.runs` skips it
        entirely and a link would move to the end of its paragraph (or vanish).
        """
        from docx.text.hyperlink import Hyperlink

        out: list[Run] = []
        for item in para.iter_inner_content():
            if isinstance(item, Hyperlink):
                link = item.address or ""
                if not link:
                    self.report.add(
                        "internal_link",
                        f"internal link to anchor {item.fragment!r}; the model's `link` "
                        "is an external URL, so the text is kept unlinked",
                        block_index=index,
                        paragraph_ordinal=ordinal,
                    )
                for inner in item.runs:
                    run = self._run(inner, index, ordinal, link=link)
                    if run is not None:
                        out.append(run)
                continue
            run = self._run(item, index, ordinal)
            if run is not None:
                out.append(run)
        return _merge(out)

    def _run(self, source: _Xml, index: int, ordinal: int, *, link: str = "") -> Run | None:
        """One model Run, or None for a run that carries no text.

        A textless run is dropped rather than reported: `w:r` elements with only
        properties are how Word records formatting at a cursor position, and they say
        nothing about the document's content.
        """
        text = str(source.text or "")
        code = (source.font.name or "") == _MONOSPACE
        self._run_props(source, index, ordinal, link=link)
        if not text:
            return None
        run = Run(
            text=text,
            bold=bool(source.bold),
            italic=bool(source.italic),
            code=code,
            link=link,
        )
        if link:
            self._intra_run_breaks(run, index, ordinal)
        return run

    def _run_props(self, source: _Xml, index: int, ordinal: int, *, link: str) -> None:
        properties = _find(source._r, "rPr")
        if properties is None:
            return
        for child in properties.iterchildren():
            name = _local(child)
            value = _attr(child, "val")
            if name in ("b", "i") and _is_off(child):
                self.report.add(
                    "explicit_off_toggle",
                    f"{name!r} explicitly switched OFF; the model's False means "
                    "'inherit' and the writer only ever turns a property on",
                    block_index=index,
                    paragraph_ordinal=ordinal,
                )
                continue
            if name == "rFonts":
                font = _attr(child, "ascii") or _attr(child, "hAnsi") or ""
                if font and font != _MONOSPACE:
                    self.report.add(
                        "run_property",
                        f"font {font!r}; the model has no font field (only "
                        f"code=True, written as {_MONOSPACE})",
                        block_index=index,
                        paragraph_ordinal=ordinal,
                    )
                continue
            # Inside a hyperlink these two exact values ARE the writer's link decoration
            # and are already carried by `Run.link` — reporting them would make every
            # link this repo writes look lossy. Any other value is a real loss.
            if link and name == "color" and (value or "").upper() == _LINK_COLOR:
                continue
            if link and name == "u" and value == "single":
                continue
            if name in _RUN_PROPS:
                described = f"{_RUN_PROPS[name]} ({name}" + (f"={value})" if value else ")")
                self.report.add(
                    "run_property",
                    f"{described}; the model's Run holds bold/italic/code/link only",
                    block_index=index,
                    paragraph_ordinal=ordinal,
                )

    # -- paragraph properties ------------------------------------------------------

    def _para_props(
        self, para: _Xml, index: int, ordinal: int, *, consumed_numbering: bool
    ) -> None:
        properties = _find(para._p, "pPr")
        if properties is None:
            return
        for child in properties.iterchildren():
            name = _local(child)
            if name == "numPr" and consumed_numbering:
                continue  # read as the list kind, not dropped
            if name == "spacing":
                self._spacing_props(child, index, ordinal)
                continue
            if name in _PARA_PROPS:
                self.report.add(
                    "paragraph_property",
                    f"{_PARA_PROPS[name]} ({name}); the model's ParagraphStyle holds "
                    "alignment and spacing only",
                    block_index=index,
                    paragraph_ordinal=ordinal,
                )

    def _spacing_props(self, spacing: _Xml, index: int, ordinal: int) -> None:
        rule = _attr(spacing, "lineRule")
        if rule and rule != "auto":
            self.report.add(
                "line_spacing_exact",
                f"line spacing rule {rule!r}; the model's line_spacing is a multiple of "
                "single spacing, so an absolute height cannot be expressed",
                block_index=index,
                paragraph_ordinal=ordinal,
            )
        for attribute, label in (("before", "space_before_pt"), ("after", "space_after_pt")):
            if _attr(spacing, attribute) == "0":
                self.report.add(
                    "paragraph_property",
                    f"explicit zero {attribute} spacing; 0.0 means 'unset' in the "
                    f"model's {label}, so a deliberate zero cannot be expressed",
                    block_index=index,
                    paragraph_ordinal=ordinal,
                )

    def _style(self, para: _Xml, index: int, ordinal: int) -> ParagraphStyle | None:
        """The block's `ParagraphStyle`, or None when the paragraph declares nothing.

        None rather than an all-zero style so a parsed model keeps the writer's "the
        template decides" reading; an empty style object would say the same thing but
        make every block look styled.
        """
        fmt = para.paragraph_format
        align = self._align(para, index, ordinal)
        before = fmt.space_before.pt if fmt.space_before is not None else 0.0
        after = fmt.space_after.pt if fmt.space_after is not None else 0.0
        spacing = fmt.line_spacing
        # A float is a multiple of single spacing (lineRule="auto"); a Length is an
        # absolute height, which the model cannot hold — `_spacing_props` reports it.
        line = float(spacing) if isinstance(spacing, float) else 0.0
        if not align and not before and not after and not line:
            return None
        return ParagraphStyle(
            align=align,
            space_before_pt=float(before),
            space_after_pt=float(after),
            line_spacing=line,
        )

    def _align(self, para: _Xml, index: int, ordinal: int) -> str:
        alignment = para.alignment
        if alignment is None:
            return ""
        name = str(getattr(alignment, "name", "")).lower()
        if name in ("left", "center", "right"):
            return name
        if name == "justify":
            return "justify"
        self.report.add(
            "paragraph_property",
            f"alignment {name!r} is outside the model's set (left/center/right/justify)",
            block_index=index,
            paragraph_ordinal=ordinal,
        )
        return ""

    # -- constructs the model cannot hold ------------------------------------------

    def _constructs(self, element: _Xml, index: int, ordinal: int) -> None:
        """Report every unrepresentable construct inside one paragraph element.

        Counted per kind and reported once, so a paragraph with eleven footnote
        references produces one actionable item rather than eleven identical ones.
        """
        counts: dict[str, int] = {}
        names: set[str] = set()
        for node in element.iter():
            name = _local(node)
            names.add(name)
            kind = _CONSTRUCTS.get(name)
            if kind is not None:
                counts[kind] = counts.get(kind, 0) + 1
        # A text box lives inside a `w:pict` or a `w:drawing`, so the enclosing wrapper
        # must NOT also be reported as an image — one construct, one item.
        if "txbxContent" in names:
            counts["text_box"] = counts.get("text_box", 0) + 1
        elif "drawing" in names or "pict" in names:
            counts["embedded_image"] = counts.get("embedded_image", 0) + 1
        for kind, count in counts.items():
            self.report.add(
                kind,
                f"{count} {kind.replace('_', ' ')}"
                + ("s" if count > 1 else "")
                + f"; {_CONSTRUCT_REASON[kind]}",
                block_index=index,
                paragraph_ordinal=ordinal,
            )

    def _intra_run_breaks(self, run: Run, index: int, ordinal: int) -> None:
        """Report a `w:br` / `w:tab` that the writer's hyperlink path cannot re-emit.

        MEASURED, not assumed: python-docx's `Run.text` translates `w:br` → "\\n" and
        `w:tab` → "\\t" in BOTH directions, so on an ordinary run these two survive the
        model's plain `text` field and are NOT a loss. The writer's hyperlink path is the
        exception — it builds its `w:t` element directly, so a newline inside a link ends
        up as literal whitespace in the XML and Word renders no break at all.
        """
        for character, kind, label in (("\n", "line_break", "line break"), ("\t", "tab", "tab")):
            count = run.text.count(character)
            if not count:
                continue
            self.report.add(
                kind,
                f"{count} {label}(s) inside a hyperlink; the writer emits a link's text "
                "as a single XML text node, so the break is not re-emitted (on an "
                "ordinary run it round-trips through the model's text)",
                block_index=index,
                paragraph_ordinal=ordinal,
            )


#: Why each construct cannot be represented. In the item's `detail` so a caller reading
#: one line of the report learns what to do about it.
_CONSTRUCT_REASON = {
    "footnote": "the model has no notes, so the marker and its text are dropped",
    "endnote": "the model has no notes, so the marker and its text are dropped",
    "comment": "the model has no annotations, so the comment is dropped",
    "tracked_change": "the model holds one document state, not a revision history",
    "field": "the model holds literal text, so the field's code is dropped",
    "bookmark": "the model has no anchors, so the bookmark is dropped",
    "math": "the model has no equation, so the formula is dropped",
    "embedded_object": "the model has no embedded object",
    "content_control": "the model has no content control; its content is kept",
    "embedded_image": "an image block REFERENCES an artifact and cannot carry bytes",
    "text_box": "the model is one linear flow, so a floating text box has no place",
}


def _style_name(para: _Xml) -> str:
    style = getattr(para, "style", None)
    return str(getattr(style, "name", "") or "")


def _find(parent: _Xml, name: str) -> _Xml:
    for child in parent.iterchildren():
        if _local(child) == name:
            return child
    return None


def _page_breaks(element: _Xml) -> int:
    """Count explicit `w:br w:type="page"` breaks in a paragraph.

    Matched on the element, NOT on python-docx's `Run.contains_page_break`: that property
    reports a RENDERED break (`w:lastRenderedPageBreak`, which Word writes and this repo's
    writer does not), so it reads False for every break `add_page_break()` produced.
    """
    return sum(
        1 for node in element.iter() if _local(node) == "br" and _attr(node, "type") == "page"
    )


def _plain(runs: list[Run], *, ignore_code: bool = False) -> bool:
    """True when no run carries formatting the plain `text` field would lose."""
    return all(
        not run.bold and not run.italic and not run.link and (ignore_code or not run.code)
        for run in runs
    )


def _merge(runs: list[Run]) -> list[Run]:
    """Collapse adjacent runs with identical formatting.

    Word splits a single formatted phrase across many `w:r` elements (spell-check state,
    revision ids, editing history), so without this a parse of the same sentence yields a
    different run list depending on how it was typed — and a re-render would then be
    compared against a model that never round-trips to itself.
    """
    out: list[Run] = []
    for run in runs:
        if out and (out[-1].bold, out[-1].italic, out[-1].code, out[-1].link) == (
            run.bold,
            run.italic,
            run.code,
            run.link,
        ):
            out[-1] = Run(
                text=out[-1].text + run.text,
                bold=run.bold,
                italic=run.italic,
                code=run.code,
                link=run.link,
            )
            continue
        out.append(run)
    return out
