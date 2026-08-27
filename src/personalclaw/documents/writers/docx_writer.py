"""DocumentModel → .docx bytes (python-docx).

The one place OOXML wordprocessing vocabulary is allowed. Every `Block.kind` is handled
explicitly — an unhandled kind would silently drop content, so the fallthrough renders
the block's text as a paragraph rather than nothing.

The writer is PURE: it takes a model and returns bytes. It reads no config, opens no
file, touches no store and makes no network call — the only reason it knows a URL at all
is to write it into a relationship.
"""

from __future__ import annotations

import io

from personalclaw.documents.model import Block, DocumentModel
from personalclaw.documents.registry import register_writer

#: The default template has no code style, and a user-supplied template may not define
#: one either; a monospace run conveys the intent without depending on a style existing.
_MONOSPACE = "Courier New"

#: Hyperlink runs get an explicit color + underline rather than the "Hyperlink"
#: character style: that style is not guaranteed to exist in a supplied template, and a
#: dangling style reference renders as ordinary text — the link would look like prose.
_LINK_COLOR = "0563C1"


def render_docx(model: object) -> bytes:
    from docx import Document

    if not isinstance(model, DocumentModel):
        raise TypeError("docx writer expects a DocumentModel")
    doc = Document()
    _apply_page(doc, model.page)
    if model.title:
        doc.add_heading(model.title, level=0)
    for block in model.blocks:
        _add_block(doc, block)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_block(doc, block: Block) -> None:
    kind = block.kind
    if kind == "heading":
        # An empty string adds no run, so the runs path owns the paragraph's content
        # outright; without runs this is byte-for-byte the pre-runs render.
        para = doc.add_heading("" if block.runs else block.text, level=block.level)
        _fill(para, block.runs)
        _apply_style(para, block.style)
    elif kind == "paragraph":
        para = doc.add_paragraph("" if block.runs else block.text)
        _fill(para, block.runs)
        _apply_style(para, block.style)
    elif kind == "bullets":
        for item in block.items:
            _apply_style(doc.add_paragraph(item, style="List Bullet"), block.style)
    elif kind == "numbered":
        for item in block.items:
            _apply_style(doc.add_paragraph(item, style="List Number"), block.style)
    elif kind == "table":
        if block.cells:
            _add_cell_table(doc, block.cells)
        else:
            _add_table(doc, block.rows)
    elif kind == "code":
        # No code style in the default template; a monospace run conveys the intent
        # without depending on a style that may not exist in a user-supplied template.
        para = doc.add_paragraph()
        if block.runs:
            _fill(para, block.runs, monospace=True)
        else:
            run = para.add_run(block.text)
            run.font.name = _MONOSPACE
        _apply_style(para, block.style)
    elif kind == "pagebreak":
        doc.add_page_break()
    elif kind == "image":
        # Images reference an artifact; resolving one to bytes is the CALLER's job (it
        # owns the store). Rendering a visible placeholder is more honest than dropping
        # the block, which would lose the fact that an image belonged here.
        doc.add_paragraph(f"[image: {block.artifact_slug}]")
    else:  # pragma: no cover — Block.__post_init__ rejects unknown kinds
        doc.add_paragraph(block.text)


def _fill(para, runs, *, monospace: bool = False, bold: bool = False) -> None:
    """Append one `w:r` (or `w:hyperlink`) per model Run, in order.

    `monospace` / `bold` are the CONTAINER's contribution — a code block, or a table
    header row. They only ever add to what a run asks for.
    """
    for run in runs:
        _add_run(para, run, monospace=monospace, bold=bold)


def _add_run(para, run, *, monospace: bool = False, bold: bool = False) -> None:
    if run.link:
        _add_hyperlink(para, run, monospace=monospace, bold=bold)
        return
    element = para.add_run(run.text)
    # Only ever turned ON. Writing an explicit "off" toggle would override whatever the
    # paragraph's style says, so a plain run inherits instead of contradicting.
    if run.bold or bold:
        element.bold = True
    if run.italic:
        element.italic = True
    if run.code or monospace:
        element.font.name = _MONOSPACE


def _add_hyperlink(para, run, *, monospace: bool = False, bold: bool = False) -> None:
    """Append `run` as a real, clickable `w:hyperlink`.

    python-docx has no public hyperlink-authoring API, so the relationship and the
    element are built directly — the alternative is rendering the display text and
    dropping the URL, which loses content the user wrote.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rel_id = para.part.relate_to(run.link, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    element = OxmlElement("w:r")
    # CT_RPr is a SEQUENCE, not a choice: rFonts, b, i, color, u must be appended in
    # schema order or Word reports the document as corrupt.
    props = OxmlElement("w:rPr")
    if run.code or monospace:
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), _MONOSPACE)
        fonts.set(qn("w:hAnsi"), _MONOSPACE)
        props.append(fonts)
    if run.bold or bold:
        props.append(OxmlElement("w:b"))
    if run.italic:
        props.append(OxmlElement("w:i"))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _LINK_COLOR)
    props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    element.append(props)
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    # A link with no display text shows its URL: an empty clickable span is invisible.
    text.text = run.text or run.link
    element.append(text)
    link.append(element)
    para._p.append(link)


def _apply_style(para, style) -> None:
    """Apply paragraph-level formatting; `None` means the block declared none.

    Every numeric field reads 0.0 as "writer default" and leaves the template's value
    alone. Reading it as an explicit zero would silently reformat every document that
    simply never mentioned spacing.
    """
    if style is None:
        return
    from docx.shared import Pt

    if style.align:
        _apply_align(para, style.align)
    fmt = para.paragraph_format
    if style.space_before_pt > 0:
        fmt.space_before = Pt(style.space_before_pt)
    if style.space_after_pt > 0:
        fmt.space_after = Pt(style.space_after_pt)
    if style.line_spacing > 0:
        fmt.line_spacing = style.line_spacing
    if style.indent_left_pt > 0:
        fmt.left_indent = Pt(style.indent_left_pt)
    if style.indent_right_pt > 0:
        fmt.right_indent = Pt(style.indent_right_pt)
    # `!= 0` — a NEGATIVE first-line indent is a hanging indent, a real request. The
    # `> 0` test every field above uses would drop it on the floor.
    if style.first_line_indent_pt != 0:
        fmt.first_line_indent = Pt(style.first_line_indent_pt)
    if style.keep_with_next:
        fmt.keep_with_next = True


def _apply_align(para, align: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align)
    # An unknown value leaves the template default rather than guessing an alignment.
    if alignment is not None:
        para.alignment = alignment


def _apply_page(doc, page) -> None:
    """Page geometry; `None` means the model declared none.

    A `0.0` margin is "writer default", never "no margin" — a zero-margin page is
    unprintable, so it can only ever be a field nobody filled in.

    **A named size is applied before orientation**, and it decides the page dimensions
    outright. Swapping the template's own width/height (what this did when size was
    unnameable) can only ever produce a landscape version of the TEMPLATE's paper, so a
    model asking for A4 landscape got landscape Letter and the size was silently dropped.
    """
    if page is None:
        return
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches, Pt

    width_in, height_in = page.size_in()
    for section in doc.sections:
        if width_in and height_in:
            section.page_width, section.page_height = Inches(width_in), Inches(height_in)
        elif page.orientation == "landscape" and section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
        elif page.orientation == "portrait" and section.page_height < section.page_width:
            section.page_width, section.page_height = section.page_height, section.page_width
        # Declared from the resulting geometry, not from the model's string: a reader that
        # trusts `w:orient` over the dimensions must agree with the paper it is printed on.
        if section.page_width > section.page_height:
            section.orientation = WD_ORIENT.LANDSCAPE
        elif section.page_height > section.page_width:
            section.orientation = WD_ORIENT.PORTRAIT
        for edge in ("top", "bottom", "left", "right"):
            points = float(getattr(page, f"margin_{edge}_pt", 0.0))
            if points > 0:
                setattr(section, f"{edge}_margin", Pt(points))
    _apply_header_footer(doc, page)


def _apply_header_footer(doc, page) -> None:
    """The first section's header/footer text and its page-number field.

    Written into the EXISTING first paragraph rather than an added one: python-docx ships
    every header with one empty paragraph, so adding gives a blank line above the text.
    """
    if not (page.header_text or page.footer_text or page.page_numbers):
        return
    section = doc.sections[0]
    for name, text in (("header", page.header_text), ("footer", page.footer_text)):
        if not text:
            continue
        part = getattr(section, name)
        paragraph = part.paragraphs[0] if part.paragraphs else part.add_paragraph()
        paragraph.text = text
    if page.page_numbers:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        _add_page_number_field(paragraph, separator=bool(page.footer_text))


def _add_page_number_field(paragraph, *, separator: bool) -> None:
    """Append a `PAGE` field so the number is computed per page.

    A literal digit would be the same number on every page, which is the whole reason
    `page_numbers` is a flag and not part of `footer_text`.
    """
    from docx.oxml.ns import qn

    if separator:
        paragraph.add_run(" ")
    field = paragraph.add_run()._r.makeelement(qn("w:fldSimple"), {qn("w:instr"): "PAGE"})
    paragraph._p.append(field)


def _add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    # Ragged rows are normalized to the widest row: python-docx needs a fixed column
    # count, and truncating would silently lose cells.
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(width):
            table.cell(r, c).text = str(row[c]) if c < len(row) else ""
    # Row 0 is the header by contract — bold it so the output reads as a real table.
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


def _add_cell_table(doc, rows) -> None:
    """The `Block.cells` table: same geometry rules as `_add_table`, richer cells."""
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(width):
            # Row 0 stays the header by contract; a cell's own `bold` adds to that
            # rather than being able to switch the header back off.
            _fill_cell(table.cell(r, c), row[c] if c < len(row) else None, header=r == 0)


def _fill_cell(cell, source, *, header: bool) -> None:
    if source is None:
        return
    para = cell.paragraphs[0]
    bold = header or source.bold
    if source.runs:
        # Passed DOWN rather than applied afterwards: a hyperlink's run is not a `w:r`
        # child of the paragraph, so `para.runs` cannot reach it.
        _fill(para, source.runs, bold=bold)
    else:
        para.add_run(str(source.text))
        if bold:
            for run in para.runs:
                run.bold = True
    if source.align:
        _apply_align(para, source.align)


register_writer("docx", render_docx)
