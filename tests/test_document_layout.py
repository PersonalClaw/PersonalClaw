"""DFE-6 — layout control, asserted on the file rather than on the model.

Every clause of this atom is about geometry a READER has to honour, so every assertion
here reopens the rendered bytes with python-docx and reads the property back off the
document. A model→model comparison would only prove the writer was handed what the test
authored, which is the exact failure `PageSetup` had before this atom: it carried a single
`margin_in`, the writer dutifully applied it, and the asymmetric geometry every real
template ships was silently dropped — while a model-only test stayed green.

The four clauses, and where each lives:

1. **A4 landscape with 2cm margins round-trips** — `test_a4_landscape_with_2cm_margins_*`.
2. **Each paragraph-layout field round-trips** — `test_every_paragraph_layout_field_*`,
   parametrized so a newly added field cannot be left unmeasured.
3. **A header round-trips; one the model can't hold is reported** — the round-trip half is
   here; the REPORTED half is owned by `test_docx_parser.py`, where the loss vocabulary
   and its `_COVERED_BY` completeness rail live.
4. **The geometry preview** is frontend-side (`documentLayout.test.tsx`); what this file
   owns for it is `PageSetup.size_in()`, the one geometry fact the preview consumes.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.shared import Cm, Mm, Pt

from personalclaw.documents.docx_parser import parse_docx
from personalclaw.documents.model import (
    PAGE_SIZE_IN,
    PAGE_SIZES,
    Block,
    DocumentModel,
    PageSetup,
    ParagraphStyle,
)
from personalclaw.documents.writers.docx_writer import render_docx

#: One twentieth of a point in EMU. `w:pgSz` / `w:pgMar` are stored in twips, so a page
#: dimension read back off a saved file cannot equal its exact EMU value however precisely
#: the writer computes it. Asserting exact equality would fail on a correct writer.
_ONE_TWIP_EMU = 635

#: 2cm, the margin the atom names, in the points `PageSetup` holds.
_TWO_CM_PT = Cm(2).pt


def _reopen(data: bytes):
    return Document(io.BytesIO(data))


def _render(page: PageSetup | None = None, *, style: ParagraphStyle | None = None) -> bytes:
    block = Block(kind="paragraph", text="body")
    block.style = style
    return render_docx(DocumentModel(blocks=[block], page=page))


# ── clause 1: A4 landscape, 2cm margins, read back by python-docx ────────────


def test_a4_landscape_with_2cm_margins_reads_back_correctly():
    """The atom's headline clause, measured on the reopened document."""
    section = _reopen(
        _render(
            PageSetup(
                size="a4",
                orientation="landscape",
                margin_top_pt=_TWO_CM_PT,
                margin_bottom_pt=_TWO_CM_PT,
                margin_left_pt=_TWO_CM_PT,
                margin_right_pt=_TWO_CM_PT,
            )
        )
    ).sections[0]

    # A4 landscape is 297mm wide by 210mm tall — the SWAP is the whole point of landscape.
    assert abs(section.page_width - Mm(297)) <= _ONE_TWIP_EMU
    assert abs(section.page_height - Mm(210)) <= _ONE_TWIP_EMU
    assert section.page_width > section.page_height
    for edge in ("top", "bottom", "left", "right"):
        assert abs(getattr(section, f"{edge}_margin") - Cm(2)) <= _ONE_TWIP_EMU, edge


def test_a4_landscape_is_not_what_the_template_gives_you_anyway():
    """Vacuity floor for the clause above: it must not be asserting the default.

    python-docx's template is Letter PORTRAIT with 1in/1.25in margins, so every number
    the previous test pins genuinely had to be written. Without this, a writer that
    ignored `page` entirely could pass by coincidence on a template that happened to
    match.
    """
    default = _reopen(_render(None)).sections[0]

    assert abs(default.page_width - Mm(297)) > _ONE_TWIP_EMU
    assert default.page_width < default.page_height
    assert abs(default.left_margin - Cm(2)) > _ONE_TWIP_EMU


def test_the_a4_landscape_geometry_survives_a_full_parse_round_trip():
    """Read back as the MODEL too, not only as the file — the editor loads this path."""
    page = PageSetup(
        size="a4",
        orientation="landscape",
        margin_top_pt=_TWO_CM_PT,
        margin_bottom_pt=_TWO_CM_PT,
        margin_left_pt=_TWO_CM_PT,
        margin_right_pt=_TWO_CM_PT,
    )
    parsed, report = parse_docx(_render(page))

    assert parsed.page is not None
    assert (parsed.page.size, parsed.page.orientation) == ("a4", "landscape")
    for edge in ("top", "bottom", "left", "right"):
        assert getattr(parsed.page, f"margin_{edge}_pt") == pytest.approx(_TWO_CM_PT, abs=0.05)
    assert report.kinds() == []


@pytest.mark.parametrize("size", [name for name in PAGE_SIZES if name])
def test_every_named_size_round_trips_to_its_own_name(size):
    """A size that parsed back as a DIFFERENT name would silently reformat the document,
    and one that parsed back as `""` would report a phantom loss on every save."""
    parsed, report = parse_docx(_render(PageSetup(size=size)))

    assert parsed.page is not None and parsed.page.size == size
    assert "page_property" not in report.kinds()


def test_size_in_applies_orientation_and_says_nothing_when_unset():
    """The one geometry fact the frontend preview consumes."""
    assert PageSetup(size="a4").size_in() == PAGE_SIZE_IN["a4"]
    assert PageSetup(size="a4", orientation="landscape").size_in() == (
        PAGE_SIZE_IN["a4"][1],
        PAGE_SIZE_IN["a4"][0],
    )
    # No size named: the writer's template decides, and inventing Letter here would
    # reformat a document that never asked for one.
    assert PageSetup().size_in() == (0.0, 0.0)
    assert PageSetup(orientation="landscape").size_in() == (0.0, 0.0)


def test_an_unknown_page_size_raises_rather_than_defaulting():
    """Same reading as `align`/`orientation`: a typo that quietly became "writer default"
    yields a file that looks plausible while discarding the layout the author asked for."""
    with pytest.raises(ValueError, match="unknown page size"):
        PageSetup(size="A4")  # the closed set is lower-case
    with pytest.raises(ValueError, match="unknown page size"):
        PageSetup(size="a5")


# ── clause 2: every paragraph-layout field round-trips ───────────────────────
#
# Parametrized over the (field, authored value, reader) triples so a field added to
# `ParagraphStyle` without a row here reds `test_every_layout_field_has_a_row`.

_LAYOUT_FIELDS = [
    (
        "align",
        "center",
        lambda fmt, para: para.alignment is not None and "CENTER" in str(para.alignment),
    ),
    ("space_before_pt", 18.0, lambda fmt, para: fmt.space_before == Pt(18)),
    ("space_after_pt", 6.0, lambda fmt, para: fmt.space_after == Pt(6)),
    ("line_spacing", 1.5, lambda fmt, para: fmt.line_spacing == 1.5),
    ("indent_left_pt", 36.0, lambda fmt, para: fmt.left_indent == Pt(36)),
    ("indent_right_pt", 24.0, lambda fmt, para: fmt.right_indent == Pt(24)),
    # NEGATIVE: a hanging indent. The `> 0` reading every sibling numeric uses would drop
    # it, and a positive-only test would call that field "round-tripping".
    ("first_line_indent_pt", -18.0, lambda fmt, para: fmt.first_line_indent == Pt(-18)),
    ("keep_with_next", True, lambda fmt, para: fmt.keep_with_next is True),
]


@pytest.mark.parametrize("name,value,reader", _LAYOUT_FIELDS, ids=[f[0] for f in _LAYOUT_FIELDS])
def test_every_paragraph_layout_field_reads_back_off_the_document(name, value, reader):
    para = _reopen(_render(style=ParagraphStyle(**{name: value}))).paragraphs[0]

    assert reader(para.paragraph_format, para), f"{name}={value!r} did not reach the file"


@pytest.mark.parametrize("name,value,reader", _LAYOUT_FIELDS, ids=[f[0] for f in _LAYOUT_FIELDS])
def test_each_layout_field_is_ABSENT_when_the_model_leaves_it_unset(name, value, reader):
    """The vacuity floor for the rail above, one row at a time.

    Without it, a writer that hard-coded every property would pass the whole round-trip
    table — and would silently reformat every document that never asked for a layout.
    """
    para = _reopen(_render(style=ParagraphStyle())).paragraphs[0]

    assert not reader(para.paragraph_format, para), f"{name} is applied to an unset style"


@pytest.mark.parametrize("name,value,reader", _LAYOUT_FIELDS, ids=[f[0] for f in _LAYOUT_FIELDS])
def test_every_paragraph_layout_field_survives_a_parse_round_trip(name, value, reader):
    """The editor's own circuit: parse → edit → write → parse must not lose the field."""
    parsed, _ = parse_docx(_render(style=ParagraphStyle(**{name: value})))

    style = parsed.blocks[0].style
    assert style is not None, f"{name} produced no ParagraphStyle at all"
    assert getattr(style, name) == value


def test_every_layout_field_has_a_row():
    """A new `ParagraphStyle` field must be measured before it can be added."""
    from dataclasses import fields

    declared = {spec.name for spec in fields(ParagraphStyle)}
    measured = {name for name, _, _ in _LAYOUT_FIELDS}

    assert declared == measured, (
        "every ParagraphStyle field needs a round-trip row; missing "
        f"{sorted(declared - measured)}, stale {sorted(measured - declared)}"
    )


# ── clause 3: a header round-trips ───────────────────────────────────────────


def test_a_header_and_footer_round_trip_through_the_file():
    parsed, report = parse_docx(
        _render(PageSetup(header_text="Quarterly Review", footer_text="Internal"))
    )

    assert parsed.page is not None
    assert parsed.page.header_text == "Quarterly Review"
    assert parsed.page.footer_text == "Internal"
    assert "header_footer" not in report.kinds()


def test_the_header_text_is_really_in_the_headers_own_part():
    """Read off the section's header, not the body — a "header" written into the body
    would satisfy a text search on the document while appearing on page 1 only."""
    section = _reopen(_render(PageSetup(header_text="Top of page"))).sections[0]

    assert [p.text for p in section.header.paragraphs] == ["Top of page"]
    assert "Top of page" not in "".join(
        p.text for p in _reopen(_render(PageSetup(header_text="Top of page"))).paragraphs
    )


def test_an_unset_header_leaves_the_document_without_one():
    """Vacuity floor: the assertion above must be able to fail."""
    section = _reopen(_render(PageSetup())).sections[0]

    assert [p.text for p in section.header.paragraphs if p.text.strip()] == []


def test_page_numbers_are_a_field_not_a_frozen_digit():
    """A literal number would read the same on every page, which is the whole reason
    `page_numbers` is its own flag rather than part of `footer_text`."""
    section = _reopen(_render(PageSetup(page_numbers=True, footer_text="Draft"))).sections[0]
    xml = section.footer._element.xml

    assert "fldSimple" in xml and "PAGE" in xml
    assert "Draft" in xml
    # And it survives a parse.
    parsed, _ = parse_docx(_render(PageSetup(page_numbers=True, footer_text="Draft")))
    assert parsed.page is not None
    assert parsed.page.page_numbers is True
    assert parsed.page.footer_text == "Draft"


def test_no_page_number_field_when_the_flag_is_off():
    """Vacuity floor for the field detection."""
    section = _reopen(_render(PageSetup(footer_text="Draft"))).sections[0]

    assert "fldSimple" not in section.footer._element.xml
    parsed, _ = parse_docx(_render(PageSetup(footer_text="Draft")))
    assert parsed.page is not None and parsed.page.page_numbers is False


def test_a_multi_paragraph_header_is_reported_and_not_half_kept():
    """The atom's "reported, not dropped" clause for the shape a user hits most.

    `PageSetup.header_text` is one string. Flattening two paragraphs into it would claim a
    fidelity the re-render does not have, and keeping only the first would lose the second
    without saying so.
    """
    doc = Document()
    doc.add_paragraph("body")
    header = doc.sections[0].header
    header.paragraphs[0].text = "line one"
    header.add_paragraph("line two")
    buf = io.BytesIO()
    doc.save(buf)

    model, report = parse_docx(buf.getvalue())

    assert "header_footer" in report.kinds()
    detail = report.of_kind("header_footer")[0].detail
    assert "2 non-empty paragraphs" in detail
    assert model.page is not None and model.page.header_text == ""


def test_a_single_paragraph_header_is_NOT_reported():
    """Vacuity floor for the report above — the guard must discriminate, not fire always."""
    doc = Document()
    doc.add_paragraph("body")
    doc.sections[0].header.paragraphs[0].text = "just the one"
    buf = io.BytesIO()
    doc.save(buf)

    model, report = parse_docx(buf.getvalue())

    assert "header_footer" not in report.kinds()
    assert model.page is not None and model.page.header_text == "just the one"


# ── the closed set the frontend mirrors ──────────────────────────────────────


def test_the_frontends_page_size_table_matches_the_models():
    """A cross-language rail, because the preview needs page dimensions in the browser and
    cannot import Python. A drift here is a preview that draws the wrong paper, or a size
    the editor offers and the server refuses with a 400.
    """
    from pathlib import Path

    source = Path(__file__).resolve().parents[1] / "web/src/ui/content/documentPage.ts"
    text = source.read_text(encoding="utf-8")

    for name in PAGE_SIZES:
        if not name:
            continue
        assert f"{name}:" in text, f"{name} is missing from {source.name}"
    # Vacuity floor: a name the model does NOT declare must be absent, or this rail would
    # pass against a file listing every plausible paper size.
    assert "a5:" not in text
