"""DFE-2: run-level .docx fidelity, asserted on the REOPENED document.

Every claim here is a round trip: render to bytes, hand the bytes back to python-docx,
and read the property off the document. Asserting the model instead would prove only
that we called `add_run`, which is exactly the failure this atom exists to rule out.

`Run` / `ParagraphStyle` / `Cell` / `PageSetup` are declared by `documents/model.py`; the
writer only ever READS attributes off them, so the stand-ins below stand in for the real
dataclasses and let this suite pin the writer's contract on its own.
"""

from __future__ import annotations

import io
import zipfile
from dataclasses import MISSING, dataclass, field, fields
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm, Pt

from personalclaw.documents.from_markup import document_from_markdown
from personalclaw.documents.model import (
    BLOCK_KINDS,
    Block,
    Cell,
    DocumentModel,
    PageSetup,
    ParagraphStyle,
    Run,
)
from personalclaw.documents.writers import docx_writer
from personalclaw.documents.writers.docx_writer import render_docx
from personalclaw.knowledge.readers import FileReader

_MONOSPACE = "Courier New"
_URL = "https://example.invalid/docs"

#: One twentieth of a point, in EMU. `w:pgSz` stores twips, so a named page size cannot
#: read back at exact EMU however precisely the writer computes it.
_ONE_TWIP_EMU = 635


@dataclass
class _Run:
    text: str = ""
    bold: bool = False
    italic: bool = False
    code: bool = False
    link: str = ""


@dataclass
class _Style:
    align: str = ""
    space_before_pt: float = 0.0
    space_after_pt: float = 0.0
    line_spacing: float = 0.0
    indent_left_pt: float = 0.0
    indent_right_pt: float = 0.0
    first_line_indent_pt: float = 0.0
    keep_with_next: bool = False


@dataclass
class _Cell:
    runs: list = field(default_factory=list)
    text: str = ""
    bold: bool = False
    align: str = ""


@dataclass
class _Page:
    size: str = ""
    orientation: str = ""
    margin_top_pt: float = 0.0
    margin_bottom_pt: float = 0.0
    margin_left_pt: float = 0.0
    margin_right_pt: float = 0.0
    header_text: str = ""
    footer_text: str = ""
    page_numbers: bool = False

    def size_in(self) -> tuple[float, float]:
        """The writer calls this, so the double has to answer it.

        Duplicating the shipped logic would make the double a second implementation, so it
        delegates to the real one — the point of these doubles is that the writer only
        READS off them, not that they re-derive anything.
        """
        return PageSetup(size=self.size, orientation=self.orientation).size_in()


def _block(kind: str, *, runs=(), cells=(), style=None, **kw) -> Block:
    """A `Block` carrying the inline fields the writer reads."""
    block = Block(kind=kind, **kw)
    block.runs = list(runs)
    block.cells = [list(row) for row in cells]
    block.style = style
    return block


def _model(*blocks: Block, title: str = "", page=None) -> DocumentModel:
    model = DocumentModel(title=title, blocks=list(blocks))
    model.page = page
    return model


def _reopen(data: bytes):
    """The bytes, read back by python-docx — the only surface these tests assert on."""
    return Document(io.BytesIO(data))


def _document_xml(data: bytes) -> bytes:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        return archive.read("word/document.xml")


def _part_names(data: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        return sorted(archive.namelist())


# --------------------------------------------------------------------------- the clause


def test_a_bold_run_reads_back_bold_via_python_docx():
    """DFE-2's clause, round-tripped through the file rather than the model."""
    data = render_docx(
        _model(_block("paragraph", runs=[_Run("plain "), _Run("strong", bold=True)]))
    )

    para = _reopen(data).paragraphs[0]
    assert [run.text for run in para.runs] == ["plain ", "strong"]
    assert para.runs[1].bold is True, "a bold run must read back bold from the document"
    assert para.runs[0].bold is None, "a plain run inherits; it must not declare bold OFF"


def test_an_italic_run_reads_back_italic_via_python_docx():
    data = render_docx(
        _model(_block("paragraph", runs=[_Run("plain "), _Run("leaning", italic=True)]))
    )

    para = _reopen(data).paragraphs[0]
    assert para.runs[1].italic is True
    assert para.runs[0].italic is None


def test_a_bold_run_in_a_heading_reads_back_bold():
    data = render_docx(_model(_block("heading", level=2, runs=[_Run("Big", bold=True)])))

    para = _reopen(data).paragraphs[0]
    assert para.style.name == "Heading 2"
    assert para.runs[0].text == "Big"
    assert para.runs[0].bold is True


def test_a_code_run_reads_back_in_a_monospace_font():
    runs = [_Run("call "), _Run("render_docx", code=True)]
    data = render_docx(_model(_block("paragraph", runs=runs)))

    para = _reopen(data).paragraphs[0]
    assert para.runs[1].font.name == _MONOSPACE
    assert para.runs[0].font.name is None, "only the code run is respelled"


def test_a_link_run_keeps_its_display_text_and_its_url():
    """A dropped URL is lost user content, so both halves are asserted."""
    runs = [_Run("see "), _Run("the docs", link=_URL)]
    data = render_docx(_model(_block("paragraph", runs=runs)))

    para = _reopen(data).paragraphs[0]
    # A hyperlink's run is a child of w:hyperlink, not of w:p, so it is deliberately
    # absent from para.runs — the visible text arrives through para.text.
    assert para.text == "see the docs"
    assert [link.address for link in para.hyperlinks] == [_URL]
    assert [run.text for run in para.hyperlinks[0].runs] == ["the docs"]


def test_a_link_run_with_no_display_text_shows_its_url():
    data = render_docx(_model(_block("paragraph", runs=[_Run(link=_URL)])))

    para = _reopen(data).paragraphs[0]
    assert para.text == _URL, "an empty clickable span would be invisible"
    assert [link.address for link in para.hyperlinks] == [_URL]


def test_the_hyperlink_relationship_is_external_and_points_at_the_url():
    data = render_docx(_model(_block("paragraph", runs=[_Run("the docs", link=_URL)])))

    doc = _reopen(data)
    rel = next(r for r in doc.part.rels.values() if r.reltype.endswith("/hyperlink"))
    assert rel.is_external is True, "an internal relationship would resolve to nothing"
    assert rel.target_ref == _URL


def test_the_hyperlink_run_properties_are_in_ooxml_schema_order():
    """`CT_RPr` is a sequence, not a choice — out-of-order children read as corrupt."""
    runs = [_Run("hot", bold=True, italic=True, code=True, link=_URL)]
    xml = _document_xml(render_docx(_model(_block("paragraph", runs=runs)))).decode("utf-8")

    props = xml.split("<w:rPr>")[1].split("</w:rPr>")[0]
    order = [chunk.split(" ")[0].split("/")[0] for chunk in props.split("<w:")[1:]]
    assert order == ["rFonts", "b", "i", "color", "u"]


def test_a_bold_link_run_reads_back_bold():
    data = render_docx(_model(_block("paragraph", runs=[_Run("hot", bold=True, link=_URL)])))

    para = _reopen(data).paragraphs[0]
    assert para.hyperlinks[0].runs[0].bold is True


# ------------------------------------------------------------------- backwards fidelity


def _legacy_docx(model: DocumentModel) -> bytes:
    """The pre-DFE-2 render, transcribed call-for-call from the writer this atom changes.

    Not a golden file: the SAME model driven through the SAME python-docx calls, so any
    element the runs path adds to a runs-less block turns up as an XML diff.
    """
    doc = Document()
    if model.title:
        doc.add_heading(model.title, level=0)
    for block in model.blocks:
        kind = block.kind
        if kind == "heading":
            doc.add_heading(block.text, level=block.level)
        elif kind == "paragraph":
            doc.add_paragraph(block.text)
        elif kind == "bullets":
            for item in block.items:
                doc.add_paragraph(item, style="List Bullet")
        elif kind == "numbered":
            for item in block.items:
                doc.add_paragraph(item, style="List Number")
        elif kind == "table":
            rows = block.rows
            if rows:
                width = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c in range(width):
                        table.cell(r, c).text = str(row[c]) if c < len(row) else ""
                for cell in table.rows[0].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
        elif kind == "code":
            doc.add_paragraph().add_run(block.text).font.name = _MONOSPACE
        elif kind == "pagebreak":
            doc.add_page_break()
        elif kind == "image":
            doc.add_paragraph(f"[image: {block.artifact_slug}]")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _every_kind_runs_less() -> list[Block]:
    return [
        _block("heading", text="Title Two", level=2),
        _block("paragraph", text="Prose."),
        _block("bullets", items=["one", "two"]),
        _block("numbered", items=["first", "second"]),
        _block("table", rows=[["h1", "h2"], ["a"], ["b", "c"]]),
        _block("code", text="x = 1\n"),
        _block("pagebreak"),
        _block("image", artifact_slug="chart-1"),
    ]


def test_a_runs_less_model_renders_exactly_as_it_did_before():
    """Backwards compatibility, proved on the OUTPUT rather than argued about."""
    model = _model(*_every_kind_runs_less(), title="Doc")

    assert _document_xml(render_docx(model)) == _document_xml(_legacy_docx(model))
    assert _part_names(render_docx(model)) == _part_names(_legacy_docx(model))


def test_a_runs_less_paragraph_carries_no_run_properties():
    """The narrow version of the claim above: nothing is added, not even an empty rPr."""
    para = _reopen(render_docx(_model(_block("paragraph", text="Prose.")))).paragraphs[0]

    assert para.runs[0].bold is None
    assert para.runs[0].italic is None
    assert para.runs[0].font.name is None
    assert b"w:rPr" not in _document_xml(render_docx(_model(_block("paragraph", text="Prose."))))


# ----------------------------------------------------------- 0.0 means "writer default"


def _paragraph_with_style(style) -> object:
    return _reopen(render_docx(_model(_block("paragraph", text="a", style=style)))).paragraphs[0]


def test_zero_spacing_means_writer_default_and_never_an_explicit_zero():
    unset = _paragraph_with_style(_Style())
    tight = _paragraph_with_style(_Style(space_before_pt=2, space_after_pt=3, line_spacing=0.9))

    assert unset.paragraph_format.space_before is None, "0.0 must leave the template alone"
    assert unset.paragraph_format.space_after is None
    assert unset.paragraph_format.line_spacing is None
    assert unset.alignment is None

    assert tight.paragraph_format.space_before == Pt(2)
    assert tight.paragraph_format.space_after == Pt(3)
    assert tight.paragraph_format.line_spacing == 0.9
    # "unset" and "tight" must be distinguishable in the OUTPUT, not just in the model.
    assert tight.paragraph_format.space_before != unset.paragraph_format.space_before
    assert tight.paragraph_format.line_spacing != unset.paragraph_format.line_spacing


def test_a_declared_alignment_reads_back_and_an_unknown_one_does_not_guess():
    assert _paragraph_with_style(_Style(align="center")).alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _paragraph_with_style(_Style(align="right")).alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert _paragraph_with_style(_Style(align="sideways")).alignment is None


def test_a_style_of_none_touches_nothing():
    assert _paragraph_with_style(None).paragraph_format.space_before is None
    assert b"w:pPr" not in _document_xml(render_docx(_model(_block("paragraph", text="a"))))


def _section(page) -> object:
    return _reopen(render_docx(_model(_block("paragraph", text="a"), page=page))).sections[0]


def test_zero_margin_means_writer_default_and_never_a_zero_margin():
    default = _section(None)
    declared = _section(_Page(margin_left_pt=54.0, margin_top_pt=54.0))

    assert _section(_Page()).left_margin == default.left_margin
    assert declared.left_margin == Inches(0.75)
    assert declared.top_margin == Inches(0.75)
    assert declared.left_margin != default.left_margin
    # The edges nobody set keep the template's own values — per-edge means per edge, not
    # "one number spread over four".
    assert declared.right_margin == default.right_margin


def test_landscape_swaps_the_page_and_portrait_is_a_no_op_on_a_portrait_template():
    default = _section(None)
    landscape = _section(_Page(orientation="landscape"))

    assert landscape.page_width == default.page_height
    assert landscape.page_height == default.page_width
    assert _section(_Page(orientation="portrait")).page_width == default.page_width


def test_a_named_size_decides_the_page_rather_than_swapping_the_templates():
    """The bug per-edge margins alone would not have caught: swapping the template's own
    width/height can only ever yield landscape LETTER, whatever size the model asked for."""
    a4 = _section(_Page(size="a4", orientation="landscape"))

    # Within one twip: `w:pgSz` is stored in twentieths of a point, so the exact EMU width
    # cannot survive the file. 130 EMU is a five-thousandth of a millimetre.
    assert abs(a4.page_width - Mm(297)) <= _ONE_TWIP_EMU
    assert abs(a4.page_height - Mm(210)) <= _ONE_TWIP_EMU
    assert a4.page_width != _section(_Page(orientation="landscape")).page_width


# ----------------------------------------------------------------------- richer  tables


def test_a_cells_table_keeps_the_row_zero_header_bold_and_honours_per_cell_bold():
    cells = [
        [_Cell(text="h1"), _Cell(text="h2")],
        [_Cell(text="plain"), _Cell(text="loud", bold=True)],
    ]
    table = _reopen(render_docx(_model(_block("table", cells=cells)))).tables[0]

    assert [c.text for c in table.rows[0].cells] == ["h1", "h2"]
    assert all(run.bold for run in table.cell(0, 0).paragraphs[0].runs), "row 0 is the header"
    assert all(run.bold for run in table.cell(0, 1).paragraphs[0].runs)
    assert table.cell(1, 0).paragraphs[0].runs[0].bold is None
    assert table.cell(1, 1).paragraphs[0].runs[0].bold is True


def test_a_cells_table_renders_run_level_formatting_and_alignment():
    cells = [
        [_Cell(runs=[_Run("H", bold=True)])],
        [_Cell(runs=[_Run("it", italic=True)], align="center")],
    ]
    table = _reopen(render_docx(_model(_block("table", cells=cells)))).tables[0]

    assert table.cell(1, 0).paragraphs[0].runs[0].italic is True
    assert table.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_a_ragged_cells_table_is_normalized_to_its_widest_row():
    cells = [[_Cell(text="a"), _Cell(text="b")], [_Cell(text="c")]]
    table = _reopen(render_docx(_model(_block("table", cells=cells)))).tables[0]

    assert len(table.columns) == 2
    assert table.cell(1, 1).text == "", "a missing cell is empty, not a dropped column"


def test_rows_still_render_when_cells_is_empty():
    table = _reopen(render_docx(_model(_block("table", rows=[["h"], ["v"]])))).tables[0]

    assert [table.cell(0, 0).text, table.cell(1, 0).text] == ["h", "v"]


# ---------------------------------------------------------------- the model's own sweep


def _sample(kind: str) -> tuple[Block, str]:
    """One block per kind plus the marker that proves it reached the document."""
    samples = {
        "heading": (_block("heading", text="H", level=3), "H"),
        "paragraph": (_block("paragraph", text="P"), "P"),
        "bullets": (_block("bullets", items=["B"]), "B"),
        "numbered": (_block("numbered", items=["N"]), "N"),
        "table": (_block("table", rows=[["T"]]), "T"),
        "image": (_block("image", artifact_slug="slug-1"), "[image: slug-1]"),
        "pagebreak": (_block("pagebreak"), 'w:type="page"'),
        "code": (_block("code", text="C"), "C"),
    }
    return samples[kind]


def _visible_text(doc) -> str:
    parts = [para.text for para in doc.paragraphs]
    parts += [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(parts)


def test_every_declared_block_kind_still_renders_something():
    """The model says a writer MUST handle every kind; a skipped kind drops content."""
    assert len(BLOCK_KINDS) == 8, "vacuity floor: this sweep must not shrink silently"

    for kind in BLOCK_KINDS:
        block, marker = _sample(kind)
        data = render_docx(_model(block))
        haystack = _visible_text(_reopen(data)) + _document_xml(data).decode("utf-8")
        assert marker in haystack, f"{kind} rendered nothing"


def test_a_run_carrying_block_still_renders_something_for_every_kind_that_has_runs():
    for kind in ("heading", "paragraph", "code"):
        block = _block(kind, runs=[_Run("marker", bold=True)])
        assert "marker" in _visible_text(_reopen(render_docx(_model(block)))), kind


# ------------------------------------------------------------ through the REAL  reader


def test_a_run_carrying_document_re_reads_through_the_real_reader(tmp_path):
    """The repo's in-process validity proof: the same reader that ingests uploads.

    A `w:hyperlink` keeps its visible text out of `w:p`'s own `w:r` children, so a reader
    that walked only those runs would silently drop the link's words.
    """
    runs = [_Run("plain "), _Run("strong", bold=True), _Run(", see "), _Run("the docs", link=_URL)]
    path = tmp_path / "out.docx"
    path.write_bytes(render_docx(_model(_block("paragraph", runs=runs), title="Doc")))

    text, meta = FileReader().read(str(path))

    assert meta["format"] == "docx"
    assert "plain strong, see the docs" in text, "the link's words must survive the reader"


# ----------------------------------------- the SHIPPED dataclasses, not just the doubles

#: Each double declared at the top of this file, beside the dataclass it stands in for.
#: The doubles pin the writer's contract in isolation, which is worth having — but on their
#: own they pin it to a shape `model.py` is free to walk away from: rename `Run.bold` and the
#: writer stops emitting bold in production while every test above stays green, because the
#: doubles keep the old name. No other test feeds the writer a real `Run`, so this is the
#: link that makes those assertions statements about what we actually ship.
_DOUBLES = ((_Run, Run), (_Style, ParagraphStyle), (_Cell, Cell), (_Page, PageSetup))


def _field_spec(cls) -> dict:
    """Field names mapped to their default VALUES — a rename, a drop and a changed default
    are the three ways a double goes stale, and names alone would only catch two."""
    return {
        spec.name: (spec.default_factory() if spec.default_factory is not MISSING else spec.default)
        for spec in fields(cls)
    }


@pytest.mark.parametrize(
    "double,real", _DOUBLES, ids=["Run", "ParagraphStyle", "Cell", "PageSetup"]
)
def test_each_double_matches_its_shipped_dataclass_field_for_field(double, real) -> None:
    assert _field_spec(double) == _field_spec(real), (
        f"{double.__name__} has drifted from the shipped {real.__name__}; every assertion "
        "made through it is now about a shape this repo does not ship"
    )


def test_the_parity_check_notices_a_drifted_double() -> None:
    """Vacuity floor: parity that cannot fail is not parity."""

    @dataclass
    class _Renamed:  # `bold` under a new name
        text: str = ""
        strong: bool = False
        italic: bool = False
        code: bool = False
        link: str = ""

    @dataclass
    class _Dropped:  # `link` gone
        text: str = ""
        bold: bool = False
        italic: bool = False
        code: bool = False

    @dataclass
    class _Redefaulted:  # a formatting flag that arrives ENABLED
        text: str = ""
        bold: bool = True
        italic: bool = False
        code: bool = False
        link: str = ""

    for stale in (_Renamed, _Dropped, _Redefaulted):
        assert _field_spec(stale) != _field_spec(Run), f"{stale.__name__} read as up to date"
    # ...and the same check calls the real double current, or it would condemn everything.
    assert _field_spec(_Run) == _field_spec(Run)


# ------------------------------- the whole chain: markup → model → bytes → python-docx


def test_markdown_bold_survives_markup_model_bytes_and_read_back() -> None:
    """DFE-2's clause across the ENTIRE chain, which nothing else in the repo joins.

    Every test above starts from a hand-built double, and the markup suite stops at the
    model. Neither can catch a break BETWEEN the parser's `Run` and the writer that consumes
    it, and that seam is the only place the atom's promise actually lives.
    """
    model = document_from_markdown(f"Some **strong** words and a [link]({_URL}).")
    assert type(model.blocks[0].runs[0]) is Run, "the chain must carry the SHIPPED Run"

    para = _reopen(render_docx(model)).paragraphs[0]
    reread = [(run.text, run.bold) for run in para.runs]

    assert ("strong", True) in reread, f"bold did not survive the chain: {reread}"
    assert ("Some ", None) in reread, "a plain run must inherit, not declare bold OFF"
    # The link's words live in `w:hyperlink`, NOT among `w:p`'s own runs — so `para.runs`
    # cannot see them and `para.text` is the surface that proves nothing was dropped. That
    # `.text` includes them at all is precisely what the python-docx >=1.1 floor buys.
    assert "link" not in dict(reread)
    assert para.text == "Some strong words and a link."


# ------------------------------------------------------ soul guardrail 5: writers  pure

#: A writer takes a model and returns bytes. Anything below would make it reach outward.
_IMPURE = (
    "open(",
    "requests",
    "httpx",
    "urllib",
    "socket",
    "subprocess",
    "config_dir",
    "get_config",
    "load_config",
    "os.environ",
    "getenv",
    "Path(",
)


def _impurities(source: str) -> list[str]:
    return sorted(token for token in _IMPURE if token in source)


def test_the_docx_writer_reaches_for_nothing():
    source = Path(docx_writer.__file__).read_text(encoding="utf-8")

    assert _impurities(source) == [], "a writer takes a model and returns bytes"


def test_the_purity_matcher_fires_when_the_impurity_is_present():
    """Vacuity floor: a rail that matches nothing asserts nothing."""
    impure = "import requests\nimport os\nwith open('x') as fh:\n    os.environ['A']\n"

    assert _impurities(impure) == ["open(", "os.environ", "requests"]
