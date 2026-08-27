"""DFE-3: the docx writer's field-coverage CENSUS, measured on the emitted artifact.

This file answers one question for every field of the document model: *does the .docx
writer express it, or does it silently drop it?* That answer is the ground truth a
round-trip proof and a LossReport both need — a parse -> write -> parse equality test is
only meaningful for the fields the writer can express, and for the rest the honest
outcome is a RECORDED LOSS, not a failing round trip.

Three rules this file lives by:

1. **Measured on the artifact, never on the source.** Every verdict below comes from
   rendering real bytes and reading them back with python-docx (or off `word/*.xml`
   where python-docx has no accessor). A test that greps `docx_writer.py` measures the
   source, and the source is exactly what may drift out from under a claim.
2. **Absence is asserted, not assumed — with a vacuity floor.** An empty document
   satisfies every "field X is absent" assertion (a `DocumentModel()` renders zero
   paragraphs and zero tables — see `test_the_vacuity_floor_is_real`). So every absence
   assertion here is paired with a POSITIVE CONTROL in the same test: something the
   writer demonstrably does emit, asserted present in the same bytes.
3. **The census cannot rot.** `WRITER_COVERAGE` is the summary table, and
   `_MEASUREMENTS` re-derives every one of its verdicts from a live render. If the
   writer starts emitting a field this file records as dropped — or stops emitting one
   it records as emitted — `test_the_census_matches_what_the_writer_actually_emits`
   goes red. `test_the_census_covers_every_model_field` makes a NEW model field a
   compile-time-ish failure too: it must be measured before it can be added.

Deliberately NOT re-measured here: the run-level properties DFE-2 already pins in
`tests/test_docx_run_fidelity.py` (bold / italic / code font / hyperlink text + URL +
relationship + rPr schema order, the `ParagraphStyle` numerics, the page margins, the
`Cell` table paths, and the legacy-equivalence golden). Overlap is waste. What was
genuinely unmeasured, and is measured here, is the STRUCTURAL identity of each block
kind, the three container-level fields (`title`, `level`, `items`), and every one of the
writer's silent drops.
"""

from __future__ import annotations

import io
import zipfile
from dataclasses import fields
from typing import Callable

import pytest
from docx import Document

from personalclaw.documents.model import (
    BLOCK_KINDS,
    Block,
    Cell,
    DocumentModel,
    PageSetup,
    ParagraphStyle,
    Run,
)
from personalclaw.documents.writers.docx_writer import render_docx

_MONOSPACE = "Courier New"
_URL = "https://example.invalid/dfe3"

#: One twentieth of a point, in EMU. `w:pgSz` and `w:pgMar` are stored in twips, so a page
#: dimension cannot read back at exact EMU — see `_measure_page_size`.
_ONE_TWIP_EMU = 635

#: The three verdicts. `partial` is not a hedge: it names a field whose value survives
#: in SOME form but is not recoverable as what the model declared (an artifact slug that
#: lands as prose), or which the writer honours only in one direction (a `bold` that can
#: be switched on but never off), or only for some block kinds.
VERDICTS = ("emitted", "dropped", "partial")


# --------------------------------------------------------------------------- plumbing


def _render(*blocks: Block, title: str = "", page: PageSetup | None = None) -> bytes:
    return render_docx(DocumentModel(title=title, blocks=list(blocks), page=page))


def _reopen(data: bytes):
    """Hand the bytes back to python-docx. This is the whole point of the file."""
    return Document(io.BytesIO(data))


def _document_xml(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as bundle:
        return bundle.read("word/document.xml").decode("utf-8")


def _part_names(data: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(data)) as bundle:
        return sorted(bundle.namelist())


def _styles(data: bytes) -> list[tuple[str, str]]:
    return [(para.style.name, para.text) for para in _reopen(data).paragraphs]


def _table_cell_runs(cells: list[list[Cell]], row: int, col: int) -> list:
    table = _reopen(_render(Block("table", cells=cells))).tables[0]
    return list(table.cell(row, col).paragraphs[0].runs)


# ------------------------------------------------------- the vacuity floor, made real


def test_the_vacuity_floor_is_real():
    """Why every absence assertion below carries a positive control.

    An empty model renders a structurally valid .docx with NO paragraphs, NO tables and
    no paragraph properties at all. So "field X is not in the output" is satisfied by a
    document that contains nothing whatsoever — the assertion proves nothing on its own.
    """
    empty = _render()

    assert _reopen(empty).paragraphs == []
    assert _reopen(empty).tables == []
    assert "w:pPr" not in _document_xml(empty)
    # ...and yet it IS a real docx: the absence is genuinely vacuous, not a crash.
    assert "word/document.xml" in _part_names(empty)


# ------------------------------------------------- POSITIVE: every block kind's shape
#
# DFE-2 asserts each kind "renders something" (its marker appears in the visible text
# or the raw XML). That rules out a dropped block but not a MISRENDERED one: a table
# emitted as prose passes it. These rows pin what each kind actually becomes.


def test_a_heading_block_becomes_a_heading_styled_paragraph_at_its_level():
    """`Block.level` is emitted as the `Heading N` paragraph style, 1-6, clamped."""
    blocks = [Block("heading", text=f"h{n}", level=n) for n in range(1, 7)]

    assert _styles(_render(*blocks)) == [(f"Heading {n}", f"h{n}") for n in range(1, 7)]


def test_an_out_of_range_heading_level_is_clamped_in_the_emitted_document():
    """The model clamps to 1-6; the OUTPUT is where that has to be visible."""
    assert _styles(_render(Block("heading", text="lo", level=0))) == [("Heading 1", "lo")]
    assert _styles(_render(Block("heading", text="hi", level=99))) == [("Heading 6", "hi")]


def test_a_paragraph_block_becomes_a_normal_paragraph():
    assert _styles(_render(Block("paragraph", text="Prose."))) == [("Normal", "Prose.")]


def test_bullets_and_numbered_items_become_one_list_paragraph_each_in_order():
    """`Block.items` is emitted: one paragraph per item, list style, order preserved."""
    data = _render(
        Block("bullets", items=["a", "b"]),
        Block("numbered", items=["one", "two"]),
    )

    assert _styles(data) == [
        ("List Bullet", "a"),
        ("List Bullet", "b"),
        ("List Number", "one"),
        ("List Number", "two"),
    ]


def test_a_table_block_becomes_a_real_table_element_not_prose():
    """The misrender DFE-2's "renders something" sweep cannot see."""
    data = _render(Block("table", rows=[["h1", "h2"], ["a", "b"]]))
    table = _reopen(data).tables[0]

    assert "<w:tbl>" in _document_xml(data)
    assert table.style.name == "Table Grid"
    assert [c.text for c in table.rows[0].cells] == ["h1", "h2"]
    assert [c.text for c in table.rows[1].cells] == ["a", "b"]
    # A table's content must NOT also leak out as body paragraphs.
    assert [para.text for para in _reopen(data).paragraphs] == []


def test_a_pagebreak_block_becomes_a_real_page_break_run():
    """python-docx exposes no page-break accessor, so this reads the XML directly."""
    data = _render(Block("pagebreak"))

    assert '<w:br w:type="page"/>' in _document_xml(data)
    assert _styles(data) == [("Normal", "")]


def test_a_code_block_becomes_a_monospace_paragraph_and_carries_no_code_style():
    """PARTIAL by design: `code` is a font, not a style — recorded, with its reason.

    The writer documents why (`_MONOSPACE`): the default template defines no code style
    and a user-supplied template may not either, so a dangling style reference would
    render as ordinary prose. The consequence for a LossReport is that a reader cannot
    tell a `code` block from a `paragraph` of monospace runs — the KIND is lost even
    though the text and the monospace intent survive.
    """
    data = _render(Block("code", text="x = 1"))
    para = _reopen(data).paragraphs[0]

    assert para.style.name == "Normal", "recorded loss: no code paragraph style exists"
    assert [run.font.name for run in para.runs] == [_MONOSPACE]  # positive control
    assert para.text == "x = 1"


def test_the_structural_sweep_covers_every_declared_block_kind():
    """A kind with no coverage row is exactly how a silent drop survives."""
    assert len(BLOCK_KINDS) == 8, "vacuity floor: this census must not shrink silently"
    assert set(BLOCK_KINDS) == set(_KIND_RECOVERY), "every kind needs a recovery probe"

    recovered = {kind: probe() for kind, probe in _KIND_RECOVERY.items()}

    assert recovered == {
        "heading": True,
        "paragraph": True,
        "bullets": True,
        "numbered": True,
        "table": True,
        "pagebreak": True,
        # Recorded losses: a `code` block is a Normal paragraph (font only) and an
        # `image` block is a text paragraph. Neither kind is recoverable as itself.
        "code": False,
        "image": False,
    }


# ---------------------------------------------------- POSITIVE: the container fields


def test_the_document_title_becomes_a_title_styled_paragraph():
    """`DocumentModel.title` — emitted as the level-0 heading, i.e. the `Title` style."""
    data = _render(Block("paragraph", text="body"), title="My Report")

    assert _styles(data) == [("Title", "My Report"), ("Normal", "body")]


def test_an_empty_title_adds_no_paragraph_at_all():
    """The other half of the claim: `""` means "no title", not "an empty one"."""
    assert _styles(_render(Block("paragraph", text="body"))) == [("Normal", "body")]


def test_the_block_order_is_the_emitted_order():
    """`DocumentModel.blocks` — sequence is content; a reordered document is a wrong one."""
    blocks = [Block("paragraph", text=f"p{n}") for n in range(4)]

    assert [text for _, text in _styles(_render(*blocks))] == ["p0", "p1", "p2", "p3"]


def test_a_paragraph_style_on_a_list_block_reaches_every_item_paragraph():
    """`Block.style` on a multi-paragraph kind applies per item, not just to the first."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = _render(Block("bullets", items=["a", "b"], style=ParagraphStyle(align="right")))

    assert [para.alignment for para in _reopen(data).paragraphs] == [
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT,
    ]


# ------------------------------------------------------------------ RECORDED LOSSES
#
# Each test below names the field, asserts the emitted document does not carry it, and
# pairs that with a positive control in the SAME bytes so the absence is not vacuous.
# These are the LossReport's vocabulary: a round-trip proof must expect a loss here,
# not an equality.


def test_an_image_block_carries_no_image_only_a_text_placeholder():
    """RECORDED LOSS: `Block.artifact_slug` survives as PROSE, never as an image.

    An image block references an artifact rather than carrying bytes, and resolving one
    is the caller's job (it owns the store). The writer's honest fallback is a visible
    placeholder — which means the emitted .docx has no image part, no drawing, and no
    image relationship. A round trip recovers the slug only by parsing the placeholder
    text back out, and recovers `kind == "image"` not at all.
    """
    data = _render(Block("image", artifact_slug="chart-1"))
    doc = _reopen(data)
    xml = _document_xml(data)

    # Positive control: the slug IS present, as text. The absences below are not vacuous.
    assert _styles(data) == [("Normal", "[image: chart-1]")]

    assert "w:drawing" not in xml, "recorded loss: no drawing element"
    assert "a:blip" not in xml, "recorded loss: no image reference"
    assert [name for name in _part_names(data) if "media" in name] == []
    assert len(doc.inline_shapes) == 0, "recorded loss: no inline shape"
    assert "image" not in {rel.reltype.rsplit("/", 1)[-1] for rel in doc.part.rels.values()}


def test_an_explicit_block_text_beside_runs_is_dropped():
    """RECORDED LOSS: `Block.text` supplied ALONGSIDE `Block.runs` never reaches the file.

    `Block.__post_init__` treats an explicit `text` beside `runs` as a deliberate
    override — "a plain-text summary of a link-bearing run" — and keeps it untouched
    rather than recomputing it. The writer then renders the RUNS and discards that
    override (`"" if block.runs else block.text`). So the model can express something
    the .docx format, as this writer drives it, cannot carry.
    """
    block = Block("paragraph", text="PLAIN-OVERRIDE", runs=[Run(text="RICH-RUN")])
    assert block.text == "PLAIN-OVERRIDE", "premise: the model preserves the override"

    data = _render(block)

    assert _styles(data) == [("Normal", "RICH-RUN")]  # positive control
    assert "PLAIN-OVERRIDE" not in _document_xml(data)


def test_an_explicit_heading_text_beside_runs_is_dropped_too():
    """The same loss on the other runs-bearing kind, so it is recorded as a rule."""
    data = _render(Block("heading", text="H-OVERRIDE", level=2, runs=[Run(text="H-RUN")]))

    assert _styles(data) == [("Heading 2", "H-RUN")]  # positive control
    assert "H-OVERRIDE" not in _document_xml(data)


def test_explicit_table_rows_beside_cells_are_dropped():
    """RECORDED LOSS: `Block.rows` supplied ALONGSIDE `Block.cells` never reaches the file.

    Same shape as the `text`/`runs` override: the rich path wins outright and the plain
    one is discarded, so a hand-written plain rendering of a rich table is unexpressible.
    """
    block = Block("table", rows=[["ROWS-ONLY"]], cells=[[Cell(text="CELLS-WIN")]])
    assert block.rows == [["ROWS-ONLY"]], "premise: the model preserves the override"

    data = _render(block)

    assert _reopen(data).tables[0].cell(0, 0).text == "CELLS-WIN"  # positive control
    assert "ROWS-ONLY" not in _document_xml(data)


def test_an_explicit_cell_text_beside_cell_runs_is_dropped():
    """RECORDED LOSS: `Cell.text` supplied ALONGSIDE `Cell.runs` never reaches the file."""
    cell = Cell(text="CELL-OVERRIDE", runs=[Run(text="CELL-RUN")])
    assert cell.text == "CELL-OVERRIDE", "premise: the model preserves the override"

    data = _render(Block("table", cells=[[cell], [Cell(text="body")]]))

    assert _reopen(data).tables[0].cell(0, 0).text == "CELL-RUN"  # positive control
    assert "CELL-OVERRIDE" not in _document_xml(data)


def test_a_run_cannot_switch_off_the_bold_its_container_imposes():
    """RECORDED LOSS: `Run.bold=False` / `Cell.bold=False` are unexpressible in row 0.

    The writer only ever turns formatting ON — writing an explicit "off" toggle would
    override whatever the paragraph's style says, so a plain run inherits instead of
    contradicting. Combined with "row 0 is the header by contract", a header run that
    asks for `bold=False` is emitted BOLD. The model can say it; the file cannot.
    """
    cells = [
        [Cell(runs=[Run(text="hdr", bold=False)])],
        [Cell(runs=[Run(text="body", bold=False)])],
    ]

    header = _table_cell_runs(cells, 0, 0)
    body = _table_cell_runs(cells, 1, 0)

    # Positive control: outside the header the same run IS left alone, so the header's
    # forced `True` is a real override rather than a property of every run.
    assert [run.bold for run in body] == [None]
    assert [run.bold for run in header] == [True], "recorded loss: bold=False ignored"


def test_a_code_blocks_monospace_overrides_a_runs_own_code_flag():
    """RECORDED LOSS: `Run.code=False` is unexpressible inside a `code` block.

    Same one-directional convention as bold: the container's monospace only ever adds.
    The positive control is the same run in a `paragraph` block, where it is untouched.
    """
    in_code = _reopen(_render(Block("code", runs=[Run(text="p", code=False)]))).paragraphs[0]
    in_prose = _reopen(_render(Block("paragraph", runs=[Run(text="p", code=False)]))).paragraphs[0]

    assert [run.font.name for run in in_prose.runs] == [None]  # positive control
    assert [run.font.name for run in in_code.runs] == [_MONOSPACE]


@pytest.mark.parametrize(
    "kind, extra",
    [
        ("table", {"rows": [["cell"]]}),
        ("image", {"artifact_slug": "slug"}),
        ("pagebreak", {}),
    ],
)
def test_a_paragraph_style_on_a_table_image_or_pagebreak_block_is_dropped(kind, extra):
    """RECORDED LOSS: `Block.style` is honoured on 5 of the 8 kinds and dropped on 3.

    `_apply_style` is called for heading / paragraph / bullets / numbered / code and for
    nothing else, so a centred, generously spaced table or placeholder is silently
    rendered with the template's defaults. Not a wish for a feature — a fact a
    round-trip proof must expect, since the model round-trips a value the file lacks.
    """
    style = ParagraphStyle(align="center", space_before_pt=12, space_after_pt=13, line_spacing=2.0)
    dropped = _document_xml(_render(Block(kind, style=style, **extra)))

    # Positive control: the SAME style on a paragraph block does reach the file, so
    # these absences measure the kind dispatch and not a broken style object.
    honoured = _document_xml(_render(Block("paragraph", text="a", style=style)))
    assert "w:jc" in honoured and "w:spacing" in honoured

    assert "w:jc" not in dropped, f"recorded loss: {kind} drops style.align"
    assert "w:spacing" not in dropped, f"recorded loss: {kind} drops style spacing"


def test_a_portrait_page_is_emitted_as_geometry_only_never_as_an_explicit_orientation():
    """PARTIAL: `PageSetup.orientation` writes `w:orient` for landscape but not portrait.

    The writer only swaps when the template disagrees, so `orientation="portrait"` on a
    portrait template is a no-op: the emitted section carries the right GEOMETRY but no
    `w:orient` attribute, and is byte-identical to a document that declared nothing. A
    reader recovers "portrait" by comparing width to height, never by reading a field.
    """
    landscape = _render(Block("paragraph", text="a"), page=PageSetup(orientation="landscape"))
    portrait = _render(Block("paragraph", text="a"), page=PageSetup(orientation="portrait"))
    silent = _render(Block("paragraph", text="a"))

    assert "w:orient" in _document_xml(landscape)  # positive control
    assert "w:orient" not in _document_xml(portrait)
    assert _document_xml(portrait) == _document_xml(silent)
    # The geometry still expresses the intent, which is why this is partial, not lost.
    section = _reopen(portrait).sections[0]
    assert section.page_height > section.page_width


@pytest.mark.parametrize(
    "field_name, extra, marker",
    [
        ("items", {"items": ["ORPHAN-ITEM"]}, "ORPHAN-ITEM"),
        ("rows", {"rows": [["ORPHAN-ROW"]]}, "ORPHAN-ROW"),
        ("cells", {"cells": [[Cell(text="ORPHAN-CELL")]]}, "ORPHAN-CELL"),
        ("artifact_slug", {"artifact_slug": "ORPHAN-SLUG"}, "ORPHAN-SLUG"),
    ],
)
def test_a_field_belonging_to_another_kind_is_dropped_without_complaint(field_name, extra, marker):
    """RECORDED LOSS: `Block` is one flat dataclass, so every kind ignores most of it.

    Setting `items` on a paragraph, or `rows` on a heading, is silently discarded. This
    is the writer behaving correctly — but a round-trip proof that compares whole models
    still has to expect it, because the model can hold what the file will not carry.
    """
    data = _render(Block("paragraph", text="KEPT", **extra))

    assert _styles(data) == [("Normal", "KEPT")]  # positive control
    assert _reopen(data).tables == []
    assert marker not in _document_xml(data), f"recorded loss: paragraph drops {field_name}"


# =========================================================================== THE CENSUS
#
# One structured collection, every model field, asserted against the live measurements
# below so it cannot rot. This is the table the next session reads.

WRITER_COVERAGE: dict[str, str] = {
    # --- Run -----------------------------------------------------------------------
    "Run.text": "emitted",
    # ON only: a container (a header row) can force it on and the run cannot refuse.
    "Run.bold": "partial",
    "Run.italic": "emitted",
    # A font substitution, not a semantic flag, and a `code` block forces it on.
    "Run.code": "partial",
    "Run.link": "emitted",
    # --- ParagraphStyle (measured on a `paragraph` block; the per-kind drop is
    #     recorded against `Block.style`, which is where the dispatch lives) ---------
    "ParagraphStyle.align": "emitted",
    "ParagraphStyle.space_before_pt": "emitted",
    "ParagraphStyle.space_after_pt": "emitted",
    "ParagraphStyle.line_spacing": "emitted",
    "ParagraphStyle.indent_left_pt": "emitted",
    "ParagraphStyle.indent_right_pt": "emitted",
    # Measured with a NEGATIVE value on purpose: a hanging indent is the only place this
    # model treats a non-zero negative as a real request.
    "ParagraphStyle.first_line_indent_pt": "emitted",
    "ParagraphStyle.keep_with_next": "emitted",
    # --- PageSetup -----------------------------------------------------------------
    # `landscape` writes `w:orient`; `portrait` is geometry only.
    "PageSetup.orientation": "partial",
    "PageSetup.size": "emitted",
    "PageSetup.margin_top_pt": "emitted",
    "PageSetup.margin_bottom_pt": "emitted",
    "PageSetup.margin_left_pt": "emitted",
    "PageSetup.margin_right_pt": "emitted",
    "PageSetup.header_text": "emitted",
    "PageSetup.footer_text": "emitted",
    "PageSetup.page_numbers": "emitted",
    # --- Cell ----------------------------------------------------------------------
    "Cell.runs": "emitted",
    # Dropped when `runs` is also supplied.
    "Cell.text": "partial",
    # ON only, exactly like `Run.bold`.
    "Cell.bold": "partial",
    "Cell.align": "emitted",
    # --- Block ---------------------------------------------------------------------
    # 6 of 8 kinds are recoverable from the emitted structure; `code` and `image` are not.
    "Block.kind": "partial",
    # Dropped when `runs` is also supplied; ignored outright by 5 of the 8 kinds.
    "Block.text": "partial",
    "Block.level": "emitted",
    "Block.items": "emitted",
    # Dropped when `cells` is also supplied.
    "Block.rows": "partial",
    # Survives as placeholder PROSE: no image part, no drawing, no relationship.
    "Block.artifact_slug": "partial",
    "Block.runs": "emitted",
    "Block.cells": "emitted",
    # Honoured on heading/paragraph/bullets/numbered/code; dropped on table/image/pagebreak.
    "Block.style": "partial",
    # --- DocumentModel -------------------------------------------------------------
    "DocumentModel.title": "emitted",
    "DocumentModel.blocks": "emitted",
    "DocumentModel.page": "emitted",
}

#: How many rows of each verdict the census holds. A vacuity floor for the table
#: itself: without it the census could quietly collapse to "everything is emitted".
#:
#: `dropped: 0` is a FINDING, not an oversight: no model field is unconditionally
#: dropped by this writer. Every loss recorded above is CONDITIONAL — on a sibling
#: field being supplied too (`Block.text` beside `runs`), on the block kind
#: (`Block.style` on a table), or on direction (`bold` can be switched on, never off).
#: That is why a LossReport has to be computed per block, not read off a static list.
CENSUS_SHAPE = {"emitted": 28, "partial": 10, "dropped": 0}


# ------------------------------------------------------- the live measurements
#
# Each returns a verdict from real bytes. They are what makes the table above an
# assertion rather than a comment.


def _kind_recovers_heading() -> bool:
    return _styles(_render(Block("heading", text="h", level=3))) == [("Heading 3", "h")]


def _kind_recovers_paragraph() -> bool:
    return _styles(_render(Block("paragraph", text="p"))) == [("Normal", "p")]


def _kind_recovers_bullets() -> bool:
    return _styles(_render(Block("bullets", items=["i"]))) == [("List Bullet", "i")]


def _kind_recovers_numbered() -> bool:
    return _styles(_render(Block("numbered", items=["i"]))) == [("List Number", "i")]


def _kind_recovers_table() -> bool:
    return len(_reopen(_render(Block("table", rows=[["c"]]))).tables) == 1


def _kind_recovers_pagebreak() -> bool:
    return '<w:br w:type="page"/>' in _document_xml(_render(Block("pagebreak")))


def _kind_recovers_code() -> bool:
    """A `code` block is a Normal paragraph: nothing in the file says "code"."""
    return _styles(_render(Block("code", text="x"))) != [("Normal", "x")]


def _kind_recovers_image() -> bool:
    """An `image` block is a text paragraph: no drawing, no media part, no relationship."""
    data = _render(Block("image", artifact_slug="s"))
    return "w:drawing" in _document_xml(data) or any("media" in n for n in _part_names(data))


#: kind -> "is this kind recoverable AS ITSELF from the emitted document?"
_KIND_RECOVERY: dict[str, Callable[[], bool]] = {
    "heading": _kind_recovers_heading,
    "paragraph": _kind_recovers_paragraph,
    "bullets": _kind_recovers_bullets,
    "numbered": _kind_recovers_numbered,
    "table": _kind_recovers_table,
    "pagebreak": _kind_recovers_pagebreak,
    "code": _kind_recovers_code,
    "image": _kind_recovers_image,
}


def _verdict(*, present: bool, faithful: bool) -> str:
    """`present` = the value reaches the file at all; `faithful` = as what it declared."""
    if not present:
        return "dropped"
    return "emitted" if faithful else "partial"


def _measure_run_text() -> str:
    para = _reopen(_render(Block("paragraph", runs=[Run(text="wording")]))).paragraphs[0]
    return _verdict(present="wording" in para.text, faithful=para.text == "wording")


def _measure_run_bold() -> str:
    on = _reopen(_render(Block("paragraph", runs=[Run("x", bold=True)]))).paragraphs[0].runs[0]
    refused = [[Cell(runs=[Run("h", bold=False)])], [Cell(text="b")]]
    forced = _table_cell_runs(refused, 0, 0)[0].bold is True
    return _verdict(present=on.bold is True, faithful=not forced)


def _measure_run_italic() -> str:
    runs = _reopen(_render(Block("paragraph", runs=[Run("x", italic=True)]))).paragraphs[0].runs
    plain = _reopen(_render(Block("paragraph", runs=[Run("x")]))).paragraphs[0].runs
    return _verdict(present=runs[0].italic is True, faithful=plain[0].italic is None)


def _measure_run_code() -> str:
    on = _reopen(_render(Block("paragraph", runs=[Run("x", code=True)]))).paragraphs[0].runs[0]
    in_code = _reopen(_render(Block("code", runs=[Run("x", code=False)]))).paragraphs[0].runs[0]
    return _verdict(present=on.font.name == _MONOSPACE, faithful=in_code.font.name is None)


def _measure_run_link() -> str:
    data = _render(Block("paragraph", runs=[Run(text="docs", link=_URL)]))
    targets = {rel.target_ref for rel in _reopen(data).part.rels.values()}
    return _verdict(present=_URL in targets, faithful="docs" in _document_xml(data))


def _styled_paragraph(style: ParagraphStyle):
    return _reopen(_render(Block("paragraph", text="a", style=style))).paragraphs[0]


def _measure_style_align() -> str:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    got = _styled_paragraph(ParagraphStyle(align="center")).alignment
    return _verdict(present=got is not None, faithful=got == WD_ALIGN_PARAGRAPH.CENTER)


def _measure_style_space_before() -> str:
    from docx.shared import Pt

    got = _styled_paragraph(ParagraphStyle(space_before_pt=7)).paragraph_format.space_before
    return _verdict(present=got is not None, faithful=got == Pt(7))


def _measure_style_space_after() -> str:
    from docx.shared import Pt

    got = _styled_paragraph(ParagraphStyle(space_after_pt=9)).paragraph_format.space_after
    return _verdict(present=got is not None, faithful=got == Pt(9))


def _measure_style_line_spacing() -> str:
    got = _styled_paragraph(ParagraphStyle(line_spacing=1.5)).paragraph_format.line_spacing
    return _verdict(present=got is not None, faithful=got == 1.5)


def _measure_page_orientation() -> str:
    landscape = "w:orient" in _document_xml(
        _render(Block("paragraph", text="a"), page=PageSetup(orientation="landscape"))
    )
    portrait = "w:orient" in _document_xml(
        _render(Block("paragraph", text="a"), page=PageSetup(orientation="portrait"))
    )
    return _verdict(present=landscape, faithful=portrait)


def _measure_style_indent_left() -> str:
    from docx.shared import Pt

    got = _styled_paragraph(ParagraphStyle(indent_left_pt=24)).paragraph_format.left_indent
    return _verdict(present=got is not None, faithful=got == Pt(24))


def _measure_style_indent_right() -> str:
    from docx.shared import Pt

    got = _styled_paragraph(ParagraphStyle(indent_right_pt=12)).paragraph_format.right_indent
    return _verdict(present=got is not None, faithful=got == Pt(12))


def _measure_style_first_line_indent() -> str:
    """NEGATIVE — a hanging indent. A `> 0` guard in the writer drops this silently, so
    measuring the positive case would report "emitted" for a half-working field."""
    from docx.shared import Pt

    style = ParagraphStyle(first_line_indent_pt=-18)
    got = _styled_paragraph(style).paragraph_format.first_line_indent
    return _verdict(present=got is not None, faithful=got == Pt(-18))


def _measure_style_keep_with_next() -> str:
    got = _styled_paragraph(ParagraphStyle(keep_with_next=True)).paragraph_format.keep_with_next
    return _verdict(present=got is not None, faithful=got is True)


def _measure_page_size() -> str:
    """Within one twip, not exact. `w:pgSz` is stored in twentieths of a point, so A4's
    7560000 EMU is written as 11906 twips and reads back 310 EMU high. An exact comparison
    here would report every named size as `partial` for a format rounding this code does
    not control."""
    from docx.shared import Mm

    section = _reopen(_render(Block("paragraph", text="a"), page=PageSetup(size="a4"))).sections[0]
    return _verdict(
        present=section.page_width is not None,
        faithful=abs(section.page_width - Mm(210)) <= _ONE_TWIP_EMU,
    )


def _measure_page_header_text() -> str:
    section = _reopen(
        _render(Block("paragraph", text="a"), page=PageSetup(header_text="Top"))
    ).sections[0]
    texts = [p.text for p in section.header.paragraphs]
    return _verdict(present="Top" in " ".join(texts), faithful=texts == ["Top"])


def _measure_page_footer_text() -> str:
    section = _reopen(
        _render(Block("paragraph", text="a"), page=PageSetup(footer_text="Bottom"))
    ).sections[0]
    texts = [p.text for p in section.footer.paragraphs]
    return _verdict(present="Bottom" in " ".join(texts), faithful=texts == ["Bottom"])


def _measure_page_numbers() -> str:
    """A real `PAGE` field, not a literal digit — the number must differ per page."""
    section = _reopen(
        _render(Block("paragraph", text="a"), page=PageSetup(page_numbers=True))
    ).sections[0]
    xml = section.footer._element.xml
    return _verdict(present="PAGE" in xml, faithful="fldSimple" in xml)


def _measure_page_margin_edge(edge: str) -> Callable[[], str]:
    def measure() -> str:
        from docx.shared import Pt

        section = _reopen(
            _render(Block("paragraph", text="a"), page=PageSetup(**{f"margin_{edge}_pt": 90.0}))
        ).sections[0]
        got = getattr(section, f"{edge}_margin")
        return _verdict(present=got is not None, faithful=got == Pt(90))

    return measure


def _measure_cell_runs() -> str:
    runs = _table_cell_runs([[Cell(text="h")], [Cell(runs=[Run("rich", italic=True)])]], 1, 0)
    return _verdict(
        present=[run.text for run in runs] == ["rich"],
        faithful=runs[0].italic is True,
    )


def _measure_cell_text() -> str:
    plain = _table_cell_runs([[Cell(text="h")], [Cell(text="only")]], 1, 0)
    beside = _render(Block("table", cells=[[Cell(text="h")], [Cell(text="LOST", runs=[Run("r")])]]))
    return _verdict(
        present=[run.text for run in plain] == ["only"],
        faithful="LOST" in _document_xml(beside),
    )


def _measure_cell_bold() -> str:
    body = _table_cell_runs([[Cell(text="h")], [Cell(text="loud", bold=True)]], 1, 0)
    header = _table_cell_runs([[Cell(text="quiet", bold=False)], [Cell(text="b")]], 0, 0)
    return _verdict(present=body[0].bold is True, faithful=header[0].bold is not True)


def _measure_cell_align() -> str:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cells = [[Cell(text="h")], [Cell(text="mid", align="center")]]
    table = _reopen(_render(Block("table", cells=cells))).tables[0]
    got = table.cell(1, 0).paragraphs[0].alignment
    return _verdict(present=got is not None, faithful=got == WD_ALIGN_PARAGRAPH.CENTER)


def _measure_block_kind() -> str:
    recovered = [probe() for probe in _KIND_RECOVERY.values()]
    return _verdict(present=any(recovered), faithful=all(recovered))


def _measure_block_text() -> str:
    plain = _styles(_render(Block("paragraph", text="body")))
    beside = _document_xml(_render(Block("paragraph", text="LOST", runs=[Run("rich")])))
    return _verdict(present=plain == [("Normal", "body")], faithful="LOST" in beside)


def _measure_block_level() -> str:
    got = [
        style
        for style, _ in _styles(_render(*(Block("heading", text="h", level=n) for n in (2, 5))))
    ]
    return _verdict(
        present=got != ["Heading 1", "Heading 1"], faithful=got == ["Heading 2", "Heading 5"]
    )


def _measure_block_items() -> str:
    got = _styles(_render(Block("bullets", items=["a", "b"])))
    return _verdict(
        present=[text for _, text in got] == ["a", "b"],
        faithful=got == [("List Bullet", "a"), ("List Bullet", "b")],
    )


def _measure_block_rows() -> str:
    plain = _reopen(_render(Block("table", rows=[["h"], ["only"]]))).tables[0]
    beside = _render(Block("table", rows=[["LOST"]], cells=[[Cell(text="c")]]))
    return _verdict(
        present=plain.cell(1, 0).text == "only",
        faithful="LOST" in _document_xml(beside),
    )


def _measure_block_artifact_slug() -> str:
    data = _render(Block("image", artifact_slug="chart-1"))
    return _verdict(
        present="chart-1" in _document_xml(data),
        faithful=len(_reopen(data).inline_shapes) > 0,
    )


def _measure_block_runs() -> str:
    runs = _reopen(_render(Block("paragraph", runs=[Run("a"), Run("b", bold=True)]))).paragraphs[0]
    return _verdict(
        present=[run.text for run in runs.runs] == ["a", "b"],
        faithful=runs.runs[1].bold is True,
    )


def _measure_block_cells() -> str:
    cells = [[Cell(text="h1"), Cell(text="h2")], [Cell(text="a"), Cell(text="b")]]
    table = _reopen(_render(Block("table", cells=cells))).tables[0]
    return _verdict(
        present=len(table.rows) == 2,
        faithful=[c.text for c in table.rows[1].cells] == ["a", "b"],
    )


def _measure_block_style() -> str:
    style = ParagraphStyle(align="center")
    honoured = [
        "w:jc" in _document_xml(_render(Block(kind, **extra, style=style)))
        for kind, extra in (
            ("paragraph", {"text": "a"}),
            ("heading", {"text": "a"}),
            ("bullets", {"items": ["a"]}),
            ("numbered", {"items": ["a"]}),
            ("code", {"text": "a"}),
            ("table", {"rows": [["a"]]}),
            ("image", {"artifact_slug": "a"}),
            ("pagebreak", {}),
        )
    ]
    return _verdict(present=any(honoured), faithful=all(honoured))


def _measure_document_title() -> str:
    got = _styles(_render(Block("paragraph", text="b"), title="Report"))
    return _verdict(present=any(text == "Report" for _, text in got), faithful=got[0][0] == "Title")


def _measure_document_blocks() -> str:
    got = [
        text for _, text in _styles(_render(*(Block("paragraph", text=f"p{n}") for n in range(3))))
    ]
    return _verdict(present=set(got) == {"p0", "p1", "p2"}, faithful=got == ["p0", "p1", "p2"])


def _measure_document_page() -> str:
    from docx.shared import Pt

    section = _reopen(
        _render(Block("paragraph", text="a"), page=PageSetup(margin_left_pt=90.0))
    ).sections[0]
    return _verdict(present=section.left_margin is not None, faithful=section.left_margin == Pt(90))


_MEASUREMENTS: dict[str, Callable[[], str]] = {
    "Run.text": _measure_run_text,
    "Run.bold": _measure_run_bold,
    "Run.italic": _measure_run_italic,
    "Run.code": _measure_run_code,
    "Run.link": _measure_run_link,
    "ParagraphStyle.align": _measure_style_align,
    "ParagraphStyle.space_before_pt": _measure_style_space_before,
    "ParagraphStyle.space_after_pt": _measure_style_space_after,
    "ParagraphStyle.line_spacing": _measure_style_line_spacing,
    "ParagraphStyle.indent_left_pt": _measure_style_indent_left,
    "ParagraphStyle.indent_right_pt": _measure_style_indent_right,
    "ParagraphStyle.first_line_indent_pt": _measure_style_first_line_indent,
    "ParagraphStyle.keep_with_next": _measure_style_keep_with_next,
    "PageSetup.orientation": _measure_page_orientation,
    "PageSetup.size": _measure_page_size,
    "PageSetup.margin_top_pt": _measure_page_margin_edge("top"),
    "PageSetup.margin_bottom_pt": _measure_page_margin_edge("bottom"),
    "PageSetup.margin_left_pt": _measure_page_margin_edge("left"),
    "PageSetup.margin_right_pt": _measure_page_margin_edge("right"),
    "PageSetup.header_text": _measure_page_header_text,
    "PageSetup.footer_text": _measure_page_footer_text,
    "PageSetup.page_numbers": _measure_page_numbers,
    "Cell.runs": _measure_cell_runs,
    "Cell.text": _measure_cell_text,
    "Cell.bold": _measure_cell_bold,
    "Cell.align": _measure_cell_align,
    "Block.kind": _measure_block_kind,
    "Block.text": _measure_block_text,
    "Block.level": _measure_block_level,
    "Block.items": _measure_block_items,
    "Block.rows": _measure_block_rows,
    "Block.artifact_slug": _measure_block_artifact_slug,
    "Block.runs": _measure_block_runs,
    "Block.cells": _measure_block_cells,
    "Block.style": _measure_block_style,
    "DocumentModel.title": _measure_document_title,
    "DocumentModel.blocks": _measure_document_blocks,
    "DocumentModel.page": _measure_document_page,
}


# ------------------------------------------------------------- the census, asserted


def test_the_census_covers_every_model_field():
    """A new model field must be MEASURED before it can be added. No silent drops."""
    declared = {
        f"{model.__name__}.{spec.name}"
        for model in (Run, ParagraphStyle, PageSetup, Cell, Block, DocumentModel)
        for spec in fields(model)
    }

    assert declared == set(WRITER_COVERAGE), "every model field needs a census verdict"
    assert declared == set(_MEASUREMENTS), "every census verdict needs a live measurement"


def test_the_census_verdicts_are_from_the_closed_vocabulary():
    assert set(WRITER_COVERAGE.values()) <= set(VERDICTS)


def test_the_census_shape_is_pinned_so_it_cannot_collapse_to_all_emitted():
    """Vacuity floor for the table itself."""
    counted = {verdict: list(WRITER_COVERAGE.values()).count(verdict) for verdict in VERDICTS}

    assert counted == CENSUS_SHAPE
    assert sum(CENSUS_SHAPE.values()) == len(WRITER_COVERAGE) == 38


@pytest.mark.parametrize("field_name", sorted(WRITER_COVERAGE))
def test_the_census_matches_what_the_writer_actually_emits(field_name):
    """The row-by-row proof: every verdict re-derived from freshly rendered bytes.

    This is the test that makes the census un-rottable. Teach the writer to express a
    field recorded as partial — or break one recorded as emitted — and the row goes red
    with the field's name in it.
    """
    measured = _MEASUREMENTS[field_name]()

    assert measured == WRITER_COVERAGE[field_name], (
        f"{field_name}: the writer now emits {measured!r}, "
        f"the census says {WRITER_COVERAGE[field_name]!r} — update the census deliberately"
    )
