"""DFE-3: the .docx → model parser, asserted on real OOXML.

Every fixture here is a genuine .docx package built in memory — either by this repo's
writer (so the writer's own conventions are the thing being inverted) or by python-docx
plus raw OOXML injected into the body. Nothing asserts against a hand-built model, because
a parser test that never opens a file proves only that the test author agreed with
themselves.

The two claims this suite exists to hold:

1. **Order survives.** `test_interleaved_*` builds a paragraph/table/paragraph/table
   document and pins the block order, together with a companion test showing that the
   obvious wrong implementation (python-docx's two separate sequences) would produce a
   DIFFERENT order — without that companion the order assertion could pass vacuously.
2. **Every unrepresentable construct is named.** `_COVERED_BY` maps every member of
   `LOSS_KINDS` to the test that exercises it, and `test_every_loss_kind_has_a_test`
   reds when a kind is added without one.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches

from personalclaw.documents.docx_parser import (
    LOSS_KINDS,
    LossItem,
    LossReport,
    parse_docx,
)
from personalclaw.documents.model import (
    Block,
    DocumentModel,
    PageSetup,
    ParagraphStyle,
    Run,
)
from personalclaw.documents.writers.docx_writer import render_docx

_W = nsdecls("w")
#: VML is not in python-docx's namespace map, so a legacy shape's declaration is spelled
#: out. Word still emits VML text boxes for compatibility-mode documents.
_V = 'xmlns:v="urn:schemas-microsoft-com:vml"'
_URL = "https://example.invalid/docs"


def _bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _paragraph_doc(*fragments: str, text: str = "body", style: str | None = None) -> bytes:
    """A one-paragraph .docx with raw OOXML appended to that paragraph."""
    doc = Document()
    para = doc.add_paragraph(text)
    if style:
        para.style = doc.styles[style]
    for fragment in fragments:
        para._p.append(parse_xml(fragment))
    return _bytes(doc)


def _pPr_doc(fragment: str, *, text: str = "body") -> bytes:
    """A one-paragraph .docx with raw OOXML appended to that paragraph's `w:pPr`."""
    doc = Document()
    para = doc.add_paragraph(text)
    para._p.get_or_add_pPr().append(parse_xml(fragment))
    return _bytes(doc)


def _numbered_doc(*items: str, num_id: str = "1", level: str = "0") -> bytes:
    """A Word-shaped list: the "List Paragraph" style plus a `w:numPr` reference.

    This is how Word writes a list. The repo's writer instead uses the "List Bullet" /
    "List Number" styles, so a parser that only checks style names reads a Word list as a
    run of plain paragraphs — which is why both paths are tested.
    """
    doc = Document()
    for item in items:
        para = doc.add_paragraph(item, style="List Paragraph")
        para._p.get_or_add_pPr().append(
            parse_xml(
                f'<w:numPr {_W}><w:ilvl w:val="{level}"/><w:numId w:val="{num_id}"/></w:numPr>'
            )
        )
    return _bytes(doc)


def _kinds(report: LossReport) -> list[str]:
    """Loss kinds excluding `page_property`.

    python-docx's default template has 1.00in top/bottom and 1.25in left/right margins,
    which `PageSetup.margin_in` genuinely cannot hold — so EVERY document built from that
    template reports one honest `page_property` item. Filtering it here keeps each test
    about its own construct; `test_default_template_margins_are_reported` owns that item.
    """
    return [kind for kind in report.kinds() if kind != "page_property"]


# --------------------------------------------------------------------------------------
# document order — the claim that a two-sequence parser cannot make
# --------------------------------------------------------------------------------------


def _interleaved_bytes() -> bytes:
    return render_docx(
        DocumentModel(
            blocks=[
                Block(kind="paragraph", text="first"),
                Block(kind="table", rows=[["a"], ["b"]]),
                Block(kind="paragraph", text="second"),
                Block(kind="table", rows=[["c"], ["d"]]),
                Block(kind="paragraph", text="third"),
            ]
        )
    )


def test_interleaved_paragraph_and_table_order_is_preserved():
    model, report = parse_docx(_interleaved_bytes())

    assert [block.kind for block in model.blocks] == [
        "paragraph",
        "table",
        "paragraph",
        "table",
        "paragraph",
    ]
    assert [block.text for block in model.blocks if block.kind == "paragraph"] == [
        "first",
        "second",
        "third",
    ]
    assert [block.rows for block in model.blocks if block.kind == "table"] == [
        [["a"], ["b"]],
        [["c"], ["d"]],
    ]
    assert _kinds(report) == []


def test_the_two_sequence_reading_would_reorder_that_fixture():
    """The companion that keeps the order assertion from passing vacuously.

    If `doc.paragraphs` + `doc.tables` happened to yield document order on this fixture,
    the test above would hold for a parser with the interleaving bug. It does not: the
    naive reading groups all paragraphs before all tables, and this test pins that
    difference so the fixture can never quietly stop discriminating.
    """
    doc = Document(io.BytesIO(_interleaved_bytes()))

    naive = ["paragraph"] * len(doc.paragraphs) + ["table"] * len(doc.tables)
    walked = [
        {"p": "paragraph", "tbl": "table"}[child.tag.rsplit("}", 1)[-1]]
        for child in doc.element.body.iterchildren()
        if child.tag.rsplit("}", 1)[-1] in ("p", "tbl")
    ]

    assert walked == ["paragraph", "table", "paragraph", "table", "paragraph"]
    assert naive != walked


# --------------------------------------------------------------------------------------
# block kinds
# --------------------------------------------------------------------------------------


def test_title_paragraph_becomes_the_model_title():
    model, _ = parse_docx(render_docx(DocumentModel(title="Report", blocks=[])))

    assert model.title == "Report"
    assert model.blocks == []


def test_heading_style_carries_its_level():
    model, _ = parse_docx(
        render_docx(
            DocumentModel(
                blocks=[Block(kind="heading", text=f"h{n}", level=n) for n in range(1, 7)]
            )
        )
    )

    assert [(block.kind, block.level, block.text) for block in model.blocks] == [
        ("heading", n, f"h{n}") for n in range(1, 7)
    ]


def test_consecutive_list_paragraphs_group_into_one_block():
    model, report = parse_docx(
        render_docx(
            DocumentModel(
                blocks=[
                    Block(kind="bullets", items=["a", "b"]),
                    Block(kind="numbered", items=["1", "2"]),
                    Block(kind="bullets", items=["c"]),
                ]
            )
        )
    )

    assert [(block.kind, block.items) for block in model.blocks] == [
        ("bullets", ["a", "b"]),
        ("numbered", ["1", "2"]),
        ("bullets", ["c"]),
    ]
    assert _kinds(report) == []


def test_a_word_shaped_bulleted_list_is_read_from_its_numbering():
    """numId 1 resolves to a `bullet` numFmt in python-docx's default numbering part."""
    model, report = parse_docx(_numbered_doc("a", "b", num_id="1"))

    assert [(block.kind, block.items) for block in model.blocks] == [("bullets", ["a", "b"])]
    assert _kinds(report) == []


def test_a_word_shaped_numbered_list_is_read_from_its_numbering():
    """numId 5 resolves to a `decimal` numFmt in the same part."""
    model, _ = parse_docx(_numbered_doc("one", "two", num_id="5"))

    assert [(block.kind, block.items) for block in model.blocks] == [("numbered", ["one", "two"])]


def test_an_unresolvable_numbering_reference_reads_as_numbered():
    """An unresolvable numId falls back to `numbered`, not `bullets`.

    Guessing "bullet" for a definition we cannot read would silently strip the numbers off
    a numbered list, which is the more visible of the two wrong answers.
    """
    model, _ = parse_docx(_numbered_doc("x", num_id="999"))

    assert [block.kind for block in model.blocks] == ["numbered"]


def test_an_all_monospace_paragraph_reads_as_a_code_block():
    model, report = parse_docx(render_docx(DocumentModel(blocks=[Block(kind="code", text="f(1)")])))

    assert [(block.kind, block.text, block.runs) for block in model.blocks] == [
        ("code", "f(1)", [])
    ]
    assert _kinds(report) == []


def test_a_partly_monospace_paragraph_stays_a_paragraph_with_a_code_run():
    model, _ = parse_docx(
        render_docx(
            DocumentModel(
                blocks=[
                    Block(
                        kind="paragraph",
                        runs=[Run(text="call "), Run(text="f(1)", code=True), Run(text=" now")],
                    )
                ]
            )
        )
    )

    (block,) = model.blocks
    assert block.kind == "paragraph"
    assert [(run.text, run.code) for run in block.runs] == [
        ("call ", False),
        ("f(1)", True),
        (" now", False),
    ]


def test_the_image_placeholder_paragraph_reads_back_as_an_image_block():
    model, report = parse_docx(
        render_docx(DocumentModel(blocks=[Block(kind="image", artifact_slug="hero-shot")]))
    )

    assert [(block.kind, block.artifact_slug) for block in model.blocks] == [("image", "hero-shot")]
    assert _kinds(report) == []


def test_a_page_break_paragraph_reads_back_as_a_pagebreak_block():
    model, report = parse_docx(
        render_docx(
            DocumentModel(
                blocks=[
                    Block(kind="paragraph", text="before"),
                    Block(kind="pagebreak"),
                    Block(kind="paragraph", text="after"),
                ]
            )
        )
    )

    assert [block.kind for block in model.blocks] == ["paragraph", "pagebreak", "paragraph"]
    assert _kinds(report) == []


def test_an_empty_paragraph_is_kept_not_dropped():
    """A blank paragraph is spacing the author asked for, not noise to tidy away."""
    model, _ = parse_docx(
        render_docx(
            DocumentModel(
                blocks=[
                    Block(kind="paragraph", text="a"),
                    Block(kind="paragraph", text=""),
                    Block(kind="paragraph", text="b"),
                ]
            )
        )
    )

    assert [block.text for block in model.blocks] == ["a", "", "b"]


# --------------------------------------------------------------------------------------
# runs, styles, tables, page
# --------------------------------------------------------------------------------------


def test_run_formatting_and_hyperlink_position_survive():
    model, report = parse_docx(
        render_docx(
            DocumentModel(
                blocks=[
                    Block(
                        kind="paragraph",
                        runs=[
                            Run(text="b", bold=True),
                            Run(text="i", italic=True),
                            Run(text="c", code=True),
                            Run(text="link", link=_URL),
                            Run(text="tail"),
                        ],
                    )
                ]
            )
        )
    )

    (block,) = model.blocks
    assert [(r.text, r.bold, r.italic, r.code, r.link) for r in block.runs] == [
        ("b", True, False, False, ""),
        ("i", False, True, False, ""),
        ("c", False, False, True, ""),
        ("link", False, False, False, _URL),
        ("tail", False, False, False, ""),
    ]
    # The link's colour + underline are the writer's own decoration and are consumed by
    # `link`; reporting them would make every link this repo writes look lossy.
    assert _kinds(report) == []


def test_adjacent_runs_with_identical_formatting_are_merged():
    """Word splits one formatted phrase across many `w:r` elements."""
    doc = Document()
    para = doc.add_paragraph()
    for chunk in ("Hel", "lo ", "wor", "ld"):
        para.add_run(chunk).bold = True

    model, _ = parse_docx(_bytes(doc))

    (block,) = model.blocks
    assert [(run.text, run.bold) for run in block.runs] == [("Hello world", True)]


def test_a_textless_run_is_dropped_without_a_loss_item():
    """`w:r` with properties and no text is cursor state, not content."""
    model, report = parse_docx(
        _paragraph_doc(f"<w:r {_W}><w:rPr><w:b/></w:rPr></w:r>", text="only")
    )

    (block,) = model.blocks
    assert block.text == "only"
    assert _kinds(report) == []


def test_paragraph_style_alignment_and_spacing_are_read():
    model, report = parse_docx(
        render_docx(
            DocumentModel(
                blocks=[
                    Block(
                        kind="paragraph",
                        text="styled",
                        style=ParagraphStyle(
                            align="center",
                            space_before_pt=6.0,
                            space_after_pt=3.0,
                            line_spacing=1.5,
                        ),
                    )
                ]
            )
        )
    )

    (block,) = model.blocks
    assert block.style is not None
    assert (
        block.style.align,
        block.style.space_before_pt,
        block.style.space_after_pt,
        block.style.line_spacing,
    ) == ("center", 6.0, 3.0, 1.5)
    assert _kinds(report) == []


def test_an_unstyled_paragraph_has_no_style_object():
    """None keeps the writer's "the template decides" reading."""
    model, _ = parse_docx(render_docx(DocumentModel(blocks=[Block(kind="paragraph", text="x")])))

    assert model.blocks[0].style is None


def test_table_cells_are_read_with_the_header_bold_on_the_runs():
    model, report = parse_docx(
        render_docx(DocumentModel(blocks=[Block(kind="table", rows=[["h1", "h2"], ["a", "b"]])]))
    )

    (block,) = model.blocks
    assert block.rows == [["h1", "h2"], ["a", "b"]]
    header, body = block.cells
    # Bold lives on the RUNS, never on `Cell.bold`: the writer bolds every header cell by
    # contract, so a cell-level flag read back from row 0 could not be told apart from
    # that convention.
    assert [(cell.text, cell.bold) for cell in header] == [("h1", False), ("h2", False)]
    assert all(run.bold for cell in header for run in cell.runs)
    assert [(cell.text, cell.runs) for cell in body] == [("a", []), ("b", [])]
    assert _kinds(report) == []


def test_page_setup_orientation_and_margins_are_read():
    model, report = parse_docx(
        render_docx(
            DocumentModel(
                page=PageSetup(orientation="landscape", margin_in=0.75),
                blocks=[Block(kind="paragraph", text="x")],
            )
        )
    )

    assert model.page == PageSetup(orientation="landscape", margin_in=0.75)
    assert _kinds(report) == []


def test_page_setup_is_always_populated():
    """The file declares a geometry, so "the writer's default" would discard a fact."""
    model, _ = parse_docx(_paragraph_doc())

    assert model.page is not None
    assert model.page.orientation == "portrait"


# --------------------------------------------------------------------------------------
# LossReport — the record of what the model cannot hold
# --------------------------------------------------------------------------------------


def test_loss_item_rejects_an_unknown_kind():
    with pytest.raises(ValueError, match="unknown loss kind"):
        LossItem(kind="not-a-kind", detail="x")


def test_loss_report_shape():
    report = LossReport()
    assert report.lossless is True
    assert report.summary() == "no losses"

    report.add("comment", "one", block_index=2, paragraph_ordinal=5)
    report.add("footnote", "two", paragraph_ordinal=5)
    report.add("comment", "three")

    assert report.lossless is False
    # `LOSS_KINDS` order, not insertion order, so a caller's branching is stable.
    assert report.kinds() == ["footnote", "comment"]
    assert [item.detail for item in report.of_kind("comment")] == ["one", "three"]
    assert report.summary() == "footnote×1, comment×2"
    assert report.items[0].where == "block 2, paragraph 5"
    assert report.items[2].where == "document"
    assert str(report.items[1]) == "footnote at paragraph 5: two"

    with pytest.raises(ValueError, match="unknown loss kind"):
        report.of_kind("nope")


def test_a_loss_item_names_the_block_and_paragraph_it_came_from():
    doc = Document()
    doc.add_paragraph("intro")
    doc.add_paragraph("second")
    doc.paragraphs[1]._p.append(parse_xml(f'<w:r {_W}><w:footnoteReference w:id="2"/></w:r>'))

    _, report = parse_docx(_bytes(doc))

    (item,) = report.of_kind("footnote")
    assert (item.block_index, item.paragraph_ordinal) == (1, 1)
    assert item.where == "block 1, paragraph 1"


def test_a_paragraph_ordinal_counts_paragraphs_a_block_index_cannot_reach():
    """Three bullets collapse into ONE block; the ordinal still names which bullet."""
    doc = Document()
    for text in ("a", "b", "c"):
        doc.add_paragraph(text, style="List Bullet")
    doc.paragraphs[2]._p.append(parse_xml(f'<w:commentReference {_W} w:id="1"/>'))

    _, report = parse_docx(_bytes(doc))

    (item,) = report.of_kind("comment")
    assert (item.block_index, item.paragraph_ordinal) == (0, 2)


# --------------------------------------------------------------------------------------
# one loss family per construct
# --------------------------------------------------------------------------------------


def test_footnote_and_endnote_references_are_reported():
    _, report = parse_docx(
        _paragraph_doc(
            f'<w:r {_W}><w:footnoteReference w:id="2"/></w:r>',
            f'<w:r {_W}><w:endnoteReference w:id="3"/></w:r>',
        )
    )

    assert _kinds(report) == ["footnote", "endnote"]


def test_a_comment_reference_is_reported():
    _, report = parse_docx(_paragraph_doc(f'<w:commentReference {_W} w:id="1"/>'))

    assert _kinds(report) == ["comment"]


def test_a_text_box_is_reported_once_and_not_also_as_an_image():
    """A text box lives inside a `w:pict`; the wrapper must not double-report."""
    _, report = parse_docx(
        _paragraph_doc(
            f"<w:r {_W}><w:pict><v:shape {_V}><v:textbox><w:txbxContent>"
            "<w:p><w:r><w:t>floating</w:t></w:r></w:p>"
            "</w:txbxContent></v:textbox></v:shape></w:pict></w:r>"
        )
    )

    assert _kinds(report) == ["text_box"]


def test_an_embedded_image_is_reported():
    """An `image` block REFERENCES an artifact, so pixel data has nowhere to go."""
    _, report = parse_docx(
        _paragraph_doc(f'<w:r {_W}><w:drawing><wp:inline {nsdecls("wp")}/></w:drawing></w:r>')
    )

    assert _kinds(report) == ["embedded_image"]


def test_an_embedded_ole_object_is_reported():
    _, report = parse_docx(
        _paragraph_doc(
            f"<w:r {_W}><w:object>"
            '<o:OLEObject xmlns:o="urn:schemas-microsoft-com:office:office"/>'
            "</w:object></w:r>"
        )
    )

    assert _kinds(report) == ["embedded_object"]


def test_a_tracked_insertion_is_reported():
    _, report = parse_docx(
        _paragraph_doc(f'<w:ins {_W} w:id="9" w:author="a"><w:r><w:t>new</w:t></w:r></w:ins>')
    )

    assert _kinds(report) == ["tracked_change"]


def test_a_field_is_reported():
    _, report = parse_docx(
        _paragraph_doc(
            f'<w:r {_W}><w:fldChar w:fldCharType="begin"/></w:r>',
            f"<w:r {_W}><w:instrText>PAGE</w:instrText></w:r>",
        )
    )

    assert _kinds(report) == ["field"]


def test_a_bookmark_is_reported():
    _, report = parse_docx(_paragraph_doc(f'<w:bookmarkStart {_W} w:id="0" w:name="anchor"/>'))

    assert _kinds(report) == ["bookmark"]


def test_an_equation_is_reported():
    _, report = parse_docx(_paragraph_doc(f'<m:oMath {nsdecls("m")}/>'))

    assert _kinds(report) == ["math"]


def test_a_block_level_content_control_is_reported_and_its_content_kept():
    doc = Document()
    doc.element.body.insert(
        0,
        parse_xml(
            f"<w:sdt {_W}><w:sdtContent><w:p><w:r><w:t>inside</w:t></w:r></w:p>"
            "</w:sdtContent></w:sdt>"
        ),
    )

    model, report = parse_docx(_bytes(doc))

    assert _kinds(report) == ["content_control"]
    # The wrapper is unrepresentable; its TEXT is not, and dropping it would be the one
    # failure a loss report must never cover for.
    assert [block.text for block in model.blocks] == ["inside"]


def test_an_internal_anchor_link_is_reported_and_its_text_kept_unlinked():
    model, report = parse_docx(
        _paragraph_doc(f'<w:hyperlink {_W} w:anchor="top"><w:r><w:t>up</w:t></w:r></w:hyperlink>')
    )

    assert _kinds(report) == ["internal_link"]
    assert model.blocks[0].text == "bodyup"
    assert all(not run.link for run in model.blocks[0].runs)


@pytest.mark.parametrize(
    "fragment,expected",
    [
        ('<w:color w:val="FF0000"/>', "text colour"),
        ('<w:sz w:val="48"/>', "font size"),
        ("<w:strike/>", "strikethrough"),
        ('<w:highlight w:val="yellow"/>', "highlight"),
        ('<w:vertAlign w:val="superscript"/>', "superscript/subscript"),
        ("<w:smallCaps/>", "small caps"),
        ('<w:u w:val="single"/>', "underline"),
        ('<w:rStyle w:val="Emphasis"/>', "character style"),
        ('<w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>', "font 'Arial'"),
    ],
)
def test_a_character_property_outside_the_model_is_reported(fragment, expected):
    _, report = parse_docx(_paragraph_doc(f"<w:r {_W}><w:rPr>{fragment}</w:rPr><w:t>x</w:t></w:r>"))

    assert _kinds(report) == ["run_property"]
    (item,) = report.of_kind("run_property")
    assert expected in item.detail


def test_an_explicitly_disabled_toggle_is_reported():
    """The model's False means "inherit" — the writer only ever turns a property ON."""
    _, report = parse_docx(
        _paragraph_doc(f'<w:r {_W}><w:rPr><w:b w:val="0"/></w:rPr><w:t>x</w:t></w:r>')
    )

    assert _kinds(report) == ["explicit_off_toggle"]


def test_the_writers_own_link_colour_and_underline_are_not_reported_outside_a_link():
    """The exemption is scoped to a hyperlink; the same values in prose are a real loss."""
    _, report = parse_docx(
        _paragraph_doc(
            f"<w:r {_W}><w:rPr>"
            '<w:color w:val="0563C1"/><w:u w:val="single"/>'
            "</w:rPr><w:t>fake link</w:t></w:r>"
        )
    )

    assert [item.detail.split(";")[0] for item in report.of_kind("run_property")] == [
        "text colour (color=0563C1)",
        "underline (u=single)",
    ]


def test_an_unmodelled_paragraph_property_is_reported():
    _, report = parse_docx(_pPr_doc(f'<w:ind {_W} w:left="720"/>'))

    assert _kinds(report) == ["paragraph_property"]
    assert "indentation" in report.of_kind("paragraph_property")[0].detail


def test_an_explicit_zero_spacing_is_reported():
    """0.0 means "unset" in `ParagraphStyle`, so a deliberate zero is unrepresentable."""
    _, report = parse_docx(_pPr_doc(f'<w:spacing {_W} w:after="0"/>'))

    assert _kinds(report) == ["paragraph_property"]
    assert "explicit zero after" in report.of_kind("paragraph_property")[0].detail


def test_an_absolute_line_spacing_rule_is_reported():
    _, report = parse_docx(_pPr_doc(f'<w:spacing {_W} w:line="360" w:lineRule="exact"/>'))

    assert _kinds(report) == ["line_spacing_exact"]


def test_a_named_paragraph_style_the_model_cannot_hold_is_reported():
    model, report = parse_docx(_paragraph_doc(style="Quote"))

    assert _kinds(report) == ["paragraph_style"]
    assert "'Quote'" in report.of_kind("paragraph_style")[0].detail
    assert [block.kind for block in model.blocks] == ["paragraph"]


def test_a_second_title_styled_paragraph_is_demoted_to_a_heading_and_reported():
    doc = Document()
    doc.add_heading("real title", level=0)
    doc.add_heading("impostor", level=0)

    model, report = parse_docx(_bytes(doc))

    assert model.title == "real title"
    assert [(block.kind, block.level, block.text) for block in model.blocks] == [
        ("heading", 1, "impostor")
    ]
    assert _kinds(report) == ["paragraph_style"]


def test_a_heading_deeper_than_the_model_allows_is_reported():
    model, report = parse_docx(_paragraph_doc(style="Heading 7"))

    assert _kinds(report) == ["heading_level_clamped"]
    assert [(block.kind, block.level) for block in model.blocks] == [("heading", 6)]


def test_a_formatted_list_item_is_reported():
    """`Block.items` is a list of plain strings — a bold word in a bullet cannot survive."""
    doc = Document()
    para = doc.add_paragraph(style="List Bullet")
    para.add_run("shouted").bold = True

    model, report = parse_docx(_bytes(doc))

    assert _kinds(report) == ["list_item_formatting"]
    assert [(block.kind, block.items) for block in model.blocks] == [("bullets", ["shouted"])]


def test_a_nested_list_level_is_reported():
    _, report = parse_docx(_numbered_doc("deep", num_id="1", level="2"))

    assert _kinds(report) == ["nested_list_level"]


def test_table_hazards_each_add_their_own_item():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).add_paragraph("second paragraph")
    inner = table.cell(1, 1).add_table(rows=1, cols=1)
    inner.style = "Table Grid"

    model, report = parse_docx(_bytes(doc))

    assert set(_kinds(report)) == {"nested_table", "merged_cells", "multi_paragraph_cell"}
    # Short rows are padded to the widest, exactly as the writer normalizes on the way
    # out — otherwise a merged table would parse ragged and re-parse padded, and the
    # round trip would disagree with itself.
    (block,) = model.blocks
    assert [len(row) for row in block.cells] == [2, 2]


def test_a_table_style_other_than_the_writers_is_reported():
    doc = Document()
    doc.add_table(rows=1, cols=1).style = "Light Shading Accent 1"

    _, report = parse_docx(_bytes(doc))

    assert _kinds(report) == ["table_style"]


def test_an_inline_page_break_is_reported():
    doc = Document()
    para = doc.add_paragraph("text and a break")
    para._p.append(parse_xml(f'<w:r {_W}><w:br w:type="page"/></w:r>'))

    model, report = parse_docx(_bytes(doc))

    assert _kinds(report) == ["inline_page_break"]
    assert [block.kind for block in model.blocks] == ["paragraph"]


@pytest.mark.parametrize("character,kind", [("\n", "line_break"), ("\t", "tab")])
def test_a_break_inside_a_hyperlink_is_reported(character, kind):
    """MEASURED asymmetry, not an assumption.

    python-docx translates `w:br` ↔ "\\n" and `w:tab` ↔ "\\t" on an ordinary run, so both
    survive the model's plain text. The writer's hyperlink path builds its `w:t` directly
    and cannot re-emit them, which is the only place they are a loss.
    """
    plain = DocumentModel(blocks=[Block(kind="paragraph", text=f"a{character}b")])
    linked = DocumentModel(
        blocks=[Block(kind="paragraph", runs=[Run(text=f"a{character}b", link=_URL)])]
    )

    plain_model, plain_report = parse_docx(render_docx(plain))
    _, linked_report = parse_docx(render_docx(linked))

    assert _kinds(plain_report) == []
    assert plain_model.blocks[0].text == f"a{character}b"
    assert _kinds(linked_report) == [kind]


def test_multiple_sections_and_a_header_are_reported():
    doc = Document()
    doc.add_paragraph("one")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("two")
    doc.sections[0].header.paragraphs[0].text = "confidential"

    _, report = parse_docx(_bytes(doc))

    assert "multi_section" in report.kinds()
    assert "header_footer" in report.kinds()
    assert "confidential" in report.of_kind("header_footer")[0].detail


def test_default_template_margins_are_reported():
    """python-docx's default template really is 1.00in/1.25in — an honest loss.

    `PageSetup.margin_in` is one number, so the parser keeps the left margin and says so.
    This item is filtered out of every other test by `_kinds`, which is why it needs an
    owner here.
    """
    model, report = parse_docx(_paragraph_doc())

    assert report.kinds() == ["page_property"]
    assert "margins differ" in report.of_kind("page_property")[0].detail
    assert model.page == PageSetup(orientation="portrait", margin_in=1.25)


def test_a_page_size_the_model_cannot_hold_is_reported():
    """`PageSetup` carries orientation, not size: an A5 page re-renders as Letter."""
    doc = Document()
    doc.sections[0].page_width = Inches(5.83)
    doc.sections[0].page_height = Inches(8.27)
    for name in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(doc.sections[0], name, Inches(1))
    doc.add_paragraph("a5")

    _, report = parse_docx(_bytes(doc))

    assert report.kinds() == ["page_property"]
    assert "page size 5.83x8.27in" in report.of_kind("page_property")[0].detail


# --------------------------------------------------------------------------------------
# constructs Word emits — read the docstring before believing this is a Word file
# --------------------------------------------------------------------------------------


def test_a_word_construct_fixture_reports_every_loss_honestly():
    """A composite of constructs WORD EMITS, injected as raw XML — NOT a file Word saved.

    ⚠️  READ THIS BEFORE CITING IT. The atom's ``done_when`` asks for "a Word-authored
    fixture reporting its losses honestly". No Microsoft Word exists in this environment
    and no owner-supplied `.docx` is committed to the repo, so this fixture is the closest
    honest substitute: the OOXML constructs Word writes (footnote reference, comment
    reference, nested table, VML text box, tracked insertion, field, bookmark) assembled
    with python-docx and raw XML injection.

    What that buys and what it does not:
      * It DOES prove the parser reports each construct Word would put in the file.
      * It does NOT prove the parser survives a real Word package's incidental baggage —
        `w:rsid*` revision ids, `mc:AlternateContent` fallbacks, `w:proofErr`, theme
        fonts, a `settings.xml` this repo never writes, or Word's own numbering shapes.

    Closing the literal clause needs an owner-supplied `.docx` saved by Word, committed as
    a fixture, and asserted here. Until then this test is deliberately named for what it
    is, so nobody reads the suite as evidence the clause is met.
    """
    doc = Document()
    doc.add_heading("Quarterly Review", level=0)
    body = doc.add_paragraph("Revenue rose sharply")
    for fragment in (
        f'<w:r {_W}><w:footnoteReference w:id="2"/></w:r>',
        f'<w:commentReference {_W} w:id="1"/>',
        f'<w:ins {_W} w:id="9" w:author="reviewer"><w:r><w:t> (revised)</w:t></w:r></w:ins>',
        f'<w:bookmarkStart {_W} w:id="0" w:name="revenue"/>',
        f'<w:r {_W}><w:fldChar w:fldCharType="begin"/></w:r>',
        f"<w:r {_W}><w:instrText>PAGE</w:instrText></w:r>",
        f"<w:r {_W}><w:pict><v:shape {_V}><v:textbox><w:txbxContent>"
        "<w:p><w:r><w:t>pull quote</w:t></w:r></w:p>"
        "</w:txbxContent></v:textbox></v:shape></w:pict></w:r>",
    ):
        body._p.append(parse_xml(fragment))
    outer = doc.add_table(rows=1, cols=1)
    outer.style = "Table Grid"
    outer.cell(0, 0).add_table(rows=1, cols=1).style = "Table Grid"

    model, report = parse_docx(_bytes(doc))

    assert model.title == "Quarterly Review"
    assert [block.kind for block in model.blocks] == ["paragraph", "table"]
    assert model.blocks[0].text.startswith("Revenue rose sharply")
    assert set(_kinds(report)) == {
        "footnote",
        "comment",
        "tracked_change",
        "bookmark",
        "field",
        "text_box",
        "nested_table",
        # python-docx appends a trailing empty paragraph after a table inside a cell (a
        # cell must not end with a table), so the nesting brings a second paragraph with
        # it — reported, correctly, as its own loss.
        "multi_paragraph_cell",
    }
    # Honesty is the point: the report NAMES each one rather than the parse looking clean.
    assert not report.lossless
    assert all(item.detail for item in report.items)


# --------------------------------------------------------------------------------------
# the coverage rail
# --------------------------------------------------------------------------------------

#: Every `LOSS_KINDS` member and the test that exercises it. A new kind with no test is
#: a kind nobody has shown fires, which is the exact defect this atom closes — so the
#: mapping is asserted complete rather than left as a convention.
_COVERED_BY = {
    "footnote": "test_footnote_and_endnote_references_are_reported",
    "endnote": "test_footnote_and_endnote_references_are_reported",
    "comment": "test_a_comment_reference_is_reported",
    "nested_table": "test_table_hazards_each_add_their_own_item",
    "text_box": "test_a_text_box_is_reported_once_and_not_also_as_an_image",
    "merged_cells": "test_table_hazards_each_add_their_own_item",
    "multi_section": "test_multiple_sections_and_a_header_are_reported",
    "header_footer": "test_multiple_sections_and_a_header_are_reported",
    "embedded_image": "test_an_embedded_image_is_reported",
    "embedded_object": "test_an_embedded_ole_object_is_reported",
    "tracked_change": "test_a_tracked_insertion_is_reported",
    "field": "test_a_field_is_reported",
    "bookmark": "test_a_bookmark_is_reported",
    "math": "test_an_equation_is_reported",
    "content_control": "test_a_block_level_content_control_is_reported_and_its_content_kept",
    "internal_link": "test_an_internal_anchor_link_is_reported_and_its_text_kept_unlinked",
    "run_property": "test_a_character_property_outside_the_model_is_reported",
    "explicit_off_toggle": "test_an_explicitly_disabled_toggle_is_reported",
    "paragraph_property": "test_an_unmodelled_paragraph_property_is_reported",
    "paragraph_style": "test_a_named_paragraph_style_the_model_cannot_hold_is_reported",
    "table_style": "test_a_table_style_other_than_the_writers_is_reported",
    "line_break": "test_a_break_inside_a_hyperlink_is_reported",
    "tab": "test_a_break_inside_a_hyperlink_is_reported",
    "line_spacing_exact": "test_an_absolute_line_spacing_rule_is_reported",
    "page_property": "test_default_template_margins_are_reported",
    "list_item_formatting": "test_a_formatted_list_item_is_reported",
    "nested_list_level": "test_a_nested_list_level_is_reported",
    "heading_level_clamped": "test_a_heading_deeper_than_the_model_allows_is_reported",
    "multi_paragraph_cell": "test_table_hazards_each_add_their_own_item",
    "inline_page_break": "test_an_inline_page_break_is_reported",
}


def test_every_loss_kind_has_a_test():
    assert sorted(_COVERED_BY) == sorted(LOSS_KINDS)
    missing = [name for name in set(_COVERED_BY.values()) if name not in globals()]
    assert missing == [], f"_COVERED_BY names tests that do not exist: {missing}"
