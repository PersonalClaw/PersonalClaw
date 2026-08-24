"""DFE-3's V1 gate clause: a fixture **Word** authored, reporting its losses honestly.

`tests/fixtures/word_authored.docx` was written by Microsoft Word (Word for Mac 16.111)
driven over AppleScript. It was not written by python-docx and its constructs were not
hand-injected as raw OOXML. That distinction is the entire point of the clause: a synthetic
fixture can only contain what its author remembered to inject, while a file Word saved
carries whatever Word puts in a document whether anybody thought of it or not — the
`w14:paraId`/`w14:textId` pair on every paragraph, the `w15`/`w16du` namespace family,
Word's own `numbering.xml` / `footnotes.xml` / `endnotes.xml` / `theme1.xml` parts, and
`w:sectPr` values from Word's template rather than python-docx's.

How it was produced, so it can be regenerated (all of step 2 through Word's own UI model):

1. `render_docx` wrote the base document: a title, a heading, a paragraph of formatted
   runs, two tables interleaved with paragraphs, a bulleted list and a numbered list.
2. Word opened that file and added a footnote, a tracked insertion inside the first
   table's first cell, and two more items typed at the end of the numbered list and then
   bulleted with the bullet button.
3. `remove personal information` and `remove date and time` were set on the document
   before saving, so **Word itself** replaced the revision author with "Author", blanked
   `dc:creator` / `cp:lastModifiedBy` and dropped `word/people.xml`. No byte was edited
   afterwards — the archive is Word's own output.

The suite makes four claims, in the order they can fail:

* the fixture really is Word's output and really carries no personal data
  (`test_the_word_provenance_markers_reject_a_python_docx_file` is the companion that
  proves those markers can fail — without it "the file has Word's marks" is unfalsifiable);
* everything the model **can** hold came back, including paragraph↔table document order on
  Word's XML rather than the writer's;
* every construct the model **cannot** hold is named in the report, and nothing else is —
  over-reporting is a failure mode too, and an exact kind set is what catches it;
* Word's own `w:numPr` outranks a contradicting style name, which is the defect this
  fixture found.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document

from personalclaw.documents.docx_parser import LOSS_KINDS, LossItem, parse_docx
from personalclaw.documents.model import Block, Cell, DocumentModel, PageSetup
from personalclaw.documents.writers.docx_writer import render_docx

FIXTURE = Path(__file__).parent / "fixtures" / "word_authored.docx"

#: Markers that separate a file Word saved from one python-docx wrote. Every one was
#: MEASURED against both before being listed: `w:rsid*` and `customXml/item1.xml` are
#: deliberately absent from this list because python-docx inherits them from its default
#: template (which Word made), so a rail on them would hold for both files and prove
#: nothing. Each entry maps a name to a predicate over the opened archive.
_WORD_MARKERS = {
    # Word stamps a revision-stable id pair on every paragraph it writes.
    "w14:paraId": lambda z, doc: "w14:paraId" in doc and "w14:textId" in doc,
    # The 2012+/2018+ namespace families python-docx's template does not declare.
    "w15/w16 namespaces": lambda z, doc: "w16du" in doc or "w15" in doc,
    "app.xml names Word": lambda z, doc: "<Application>Microsoft Office Word</Application>"
    in z.read("docProps/app.xml").decode(),
    "app.xml names Word 16": lambda z, doc: "<AppVersion>16.0000</AppVersion>"
    in z.read("docProps/app.xml").decode(),
    # Word always writes both note parts; python-docx's template carries neither.
    "footnotes part": lambda z, doc: "word/footnotes.xml" in z.namelist(),
    "endnotes part": lambda z, doc: "word/endnotes.xml" in z.namelist(),
    # python-docx signs its output; Word's privacy scrub blanks the field instead.
    "not signed python-docx": lambda z, doc: "<dc:creator>python-docx</dc:creator>"
    not in z.read("docProps/core.xml").decode(),
}

#: Strings that must not survive in a committed fixture. `Author` is what Word's
#: `remove personal information` leaves behind in `w:ins/@w:author`, and is the evidence
#: the scrub ran rather than a name that slipped through.
_PERSONAL = ("Keyur", "Golani", "golani", "keyur", "@", "Users/", "Macintosh HD")


def _markers(data: bytes) -> dict[str, bool]:
    archive = zipfile.ZipFile(io.BytesIO(data))
    document = archive.read("word/document.xml").decode()
    return {name: bool(check(archive, document)) for name, check in _WORD_MARKERS.items()}


def _fixture_bytes() -> bytes:
    return FIXTURE.read_bytes()


# --------------------------------------------------------------------------------------
# provenance — the claim the rest of the suite rests on
# --------------------------------------------------------------------------------------


def test_the_fixture_is_word_authored_and_carries_no_personal_data():
    data = _fixture_bytes()

    assert _markers(data) == dict.fromkeys(_WORD_MARKERS, True)

    archive = zipfile.ZipFile(io.BytesIO(data))
    leaked: dict[str, list[str]] = {}
    for name in archive.namelist():
        try:
            text = archive.read(name).decode()
        except UnicodeDecodeError:  # pragma: no cover — a binary part carries no names
            continue
        found = [needle for needle in _PERSONAL if needle in text]
        if found:
            leaked[name] = found
    assert leaked == {}, f"personal data in the committed fixture: {leaked}"
    # The scrub ran, rather than the document never having had an author: Word rewrote the
    # revision's author to the literal "Author" and dropped the people part entirely.
    document = archive.read("word/document.xml").decode()
    assert 'w:author="Author"' in document
    assert "word/people.xml" not in archive.namelist()


def test_the_word_provenance_markers_reject_a_python_docx_file():
    """The vacuity floor for the test above.

    "The fixture carries Word's marks" is worth nothing unless the marks can be absent.
    Every marker must read False on this repo's own writer output — if one starts holding
    for python-docx too (a template change, a new namespace), it stops discriminating and
    belongs out of `_WORD_MARKERS`, which this failing here would say.
    """
    ours = render_docx(DocumentModel(blocks=[Block(kind="paragraph", text="not from Word")]))

    assert _markers(ours) == dict.fromkeys(_WORD_MARKERS, False)


# --------------------------------------------------------------------------------------
# what the model CAN hold came back
# --------------------------------------------------------------------------------------


def test_the_word_authored_document_parses_into_the_shipped_model_classes():
    model, report = parse_docx(_fixture_bytes())

    assert type(model) is DocumentModel
    assert model.title == "DFE-3 Word fixture"
    assert type(model.page) is PageSetup
    # Word's template has uniform 1in margins, which `PageSetup.margin_in` CAN hold — so
    # unlike a python-docx document (1.00in top/bottom, 1.25in sides) this fixture carries
    # no unavoidable `page_property` item, and the loss set below can be exact.
    assert (model.page.orientation, model.page.margin_in) == ("portrait", 1.0)
    assert {type(block) for block in model.blocks} == {Block}
    assert {type(item) for item in report.items} == {LossItem}
    table = next(block for block in model.blocks if block.cells)
    assert {type(cell) for row in table.cells for cell in row} == {Cell}


def test_paragraph_and_table_order_is_preserved_on_words_own_xml():
    """The interleaving claim, re-made on a file Word wrote rather than the writer.

    A parser reading `doc.paragraphs` then `doc.tables` yields every paragraph before every
    table; this fixture alternates paragraph→table→paragraph→table→paragraph, so the two
    readings differ and the companion assertion below pins that they do.
    """
    model, _ = parse_docx(_fixture_bytes())

    assert [block.kind for block in model.blocks] == [
        "heading",
        "paragraph",
        "table",
        "paragraph",
        "table",
        "paragraph",
        "bullets",
        "numbered",
        "bullets",
    ]
    assert [block.text for block in model.blocks if block.kind == "paragraph"] == [
        "Plain then bold then italic and a mispeled wrod.",
        "Between the tables.",
        "After the second table.",
    ]
    assert [block.rows for block in model.blocks if block.kind == "table"] == [
        [["r1c1", "r1c2"], ["r2c1", "r2c2"]],
        [["x", "y"], ["z", "w"]],
    ]

    # The fixture discriminates: the naive two-sequence reading really does reorder it.
    document = Document(io.BytesIO(_fixture_bytes()))
    body = [
        child.tag.rsplit("}", 1)[-1]
        for child in document.element.body.iterchildren()
        if child.tag.rsplit("}", 1)[-1] in {"p", "tbl"}
    ]
    naive = ["p"] * body.count("p") + ["tbl"] * body.count("tbl")
    assert body != naive


def test_words_inline_formatting_survives_as_runs():
    model, _ = parse_docx(_fixture_bytes())

    paragraph = model.blocks[1]
    assert [(run.text, run.bold, run.italic) for run in paragraph.runs] == [
        ("Plain then ", False, False),
        ("bold", True, False),
        (" then ", False, False),
        ("italic", False, True),
        (" and a mispeled wrod.", False, False),
    ]


# --------------------------------------------------------------------------------------
# what the model CANNOT hold is named — and nothing else is
# --------------------------------------------------------------------------------------


def test_the_word_authored_fixture_reports_exactly_its_real_losses():
    """Honesty runs both ways.

    The `set ==` is deliberate rather than a containment check: a parser that reported
    every Word construct as lost would satisfy "each unrepresentable construct adds an
    item" and be useless. The three kinds here are the only things in the file the model
    cannot hold; everything else Word wrote (its numbering part, its theme, its rsids, its
    `w14` ids, its sectPr, its table style) either maps onto the model or is a
    representation detail a re-render reproduces, and so must NOT appear.
    """
    _, report = parse_docx(_fixture_bytes())

    assert set(report.kinds()) == {"footnote", "run_property", "tracked_change"}
    assert report.lossless is False
    assert all(item.detail for item in report.items)
    assert all(item.kind in LOSS_KINDS for item in report.items)

    # Each item is located, not just counted — "something was lost somewhere" is not a
    # report a user can act on.
    footnote = report.of_kind("footnote")[0]
    assert (footnote.block_index, footnote.paragraph_ordinal) == (0, 0)
    tracked = report.of_kind("tracked_change")[0]
    assert tracked.block_index == 2  # the first table, which is where Word put it
    assert tracked.where == "block 2"


def test_the_dropped_tracked_insertion_is_a_real_drop_the_report_names():
    """The loss item is checked against the file, not against itself.

    Word's tracked insertion lives inside `w:ins` in the first table's first cell. The
    model holds one document state, so the inserted words are dropped — which is only
    honest if the report says so, and only *provable* by reading the words out of the
    archive and showing the model does not have them.
    """
    document = zipfile.ZipFile(io.BytesIO(_fixture_bytes())).read("word/document.xml").decode()
    assert "<w:ins " in document
    assert "Tracked insertion." in document

    model, report = parse_docx(_fixture_bytes())

    cell = model.blocks[2].cells[0][0]
    assert cell.text == "r1c1"
    assert "Tracked insertion." not in cell.text
    assert report.of_kind("tracked_change")[0].detail.startswith("1 tracked change")


# --------------------------------------------------------------------------------------
# the defect this fixture found
# --------------------------------------------------------------------------------------


def test_words_own_numbering_outranks_a_contradicting_style_name():
    """`w:numPr` beats the style name, because that is what Word renders.

    The last two items were typed at the end of the numbered list, so they inherited its
    "List Number" style, and then bulleted — which adds a bullet `w:numPr` and leaves the
    style alone. Word shows bullets. Reading the style name first read them as two more
    items on the numbered list above, i.e. a paragraph the user sees as a bullet came back
    as a number AND the block boundary vanished. The three assertions are ordered so the
    first two prove the conflict is really in the file before the third makes a claim
    about it.
    """
    document = Document(io.BytesIO(_fixture_bytes()))
    conflicted = [
        paragraph for paragraph in document.paragraphs if paragraph.text.startswith("Word bullet")
    ]
    assert [paragraph.style.name for paragraph in conflicted] == ["List Number"] * 2
    assert all(
        b"<w:numPr>" in paragraph._p.xml.encode() for paragraph in conflicted
    ), "the fixture's conflict is gone: these paragraphs no longer carry a direct w:numPr"

    model, _ = parse_docx(_fixture_bytes())

    assert [(block.kind, block.items) for block in model.blocks[-3:]] == [
        ("bullets", ["Bullet one", "Bullet two"]),
        ("numbered", ["Number one", "Number two"]),
        ("bullets", ["Word bullet A", "Word bullet B"]),
    ]
