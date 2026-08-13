"""Unit tests for the Knowledge Library (store, readers, extractor, retrieval)."""

import importlib
import json
import struct
import sys

import pytest

from personalclaw.knowledge.extractor import EntityExtractor
from personalclaw.knowledge.readers import FileReader
from personalclaw.knowledge.retrieval import HybridRetriever, _attach_locator, _bytes_to_floats
from personalclaw.knowledge.store import KnowledgeStore, SimpleDiGraph, normalize_url


def test_compose_item_text_is_title_and_summary_only():
    """KL-9 clean break: the whole-item vector text is title + summary ONLY. The old
    1000-char body top-up is retired — body-level recall now lives in the chunk index —
    so ``content`` is accepted (stable signature) but never appended."""
    from personalclaw.knowledge.embedder import compose_item_text as c

    # Body is NOT appended even when the summary is thin/absent.
    assert c("Meeting notes", None, "quantum dot solar coating") == "Meeting notes"
    rich = "x" * 90
    assert c("T", rich, "body text") == f"T {rich}"
    # A huge body cannot bloat the item vector text; empties collapse cleanly.
    assert c("t", None, "z" * 5000) == "t"
    assert c("", "", "") == "" and c("Only", None, None) == "Only"


def test_unified_embedder_exposes_model_name(monkeypatch):
    """Regression (bug #15): the knowledge stats + embedding-status endpoints read the
    embedder's model label. The old per-backend embedder had a `.model` attribute;
    UnifiedEmbedder wraps an embed_fn and must expose `model_name` (from the active
    embedding selection) instead. `.model` used to AttributeError → /api/knowledge/
    stats 500 → the FE header showed 'semantic search off' even with embeddings live."""
    import personalclaw.embedding_providers.registry as emb_reg
    from personalclaw.knowledge.embedder import UnifiedEmbedder

    # A bound embedding selection → model_name is the bare model id (prefix stripped).
    monkeypatch.setattr(
        emb_reg, "_active_embedding_spec", lambda: ("sentence-transformers", "all-MiniLM-L6-v2")
    )
    e = UnifiedEmbedder(embed_fn=lambda t: [0.1, 0.2], dim_hint=2)
    assert e.model_name == "all-MiniLM-L6-v2"
    assert e.is_available() is True
    # Nothing bound → empty label, never an AttributeError.
    monkeypatch.setattr(emb_reg, "_active_embedding_spec", lambda: None)
    assert e.model_name == ""
    # The attribute the 500 came from must NOT exist (forces callers onto model_name).
    assert not hasattr(e, "model")


def test_normalize_url_canonicalizes_for_dedup():
    """URL canonicalization underpins bookmark dedup: lowercase host, drop bare trailing
    slash + default port + fragment, strip tracking params, sort the query."""
    n = normalize_url
    assert n("https://example.com/") == "https://example.com"
    assert n("https://Example.COM") == "https://example.com"
    assert n("https://example.com/?utm_source=x&id=5") == "https://example.com?id=5"
    assert n("https://example.com:443/p/") == "https://example.com/p/"  # sub-path slash kept
    assert n("https://example.com/a?b=2&a=1#frag") == "https://example.com/a?a=1&b=2"
    # Non-http inputs and junk pass through untouched.
    assert n("mailto:x@y.com") == "mailto:x@y.com"
    assert n("not a url") == "not a url"
    assert n("") == ""
    # Idempotent.
    assert n(n("https://example.com/?fbclid=z")) == n("https://example.com/?fbclid=z")


def mk(store, title, content, item_type="note", *, summary="", tags=None, embedding=None):
    """Create one logical-document item (the only ingestion shape). Replaces the
    removed legacy add_item chunk-inserter. Embedding is set post-create."""
    iid = store.create_typed_item(
        item_type=item_type,
        title=title,
        content=content,
        summary=summary,
        tags=tags or [],
    )
    if embedding is not None:
        store.db.execute("UPDATE items SET embedding = ? WHERE id = ?", (embedding, iid))
        store.db.commit()
    return iid


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def store(tmp_path):
    s = KnowledgeStore(str(tmp_path / "test.db"))
    yield s
    s.close()


@pytest.fixture()
def store_factory(tmp_path):
    """Return a callable that creates a new store at a given path."""
    stores = []

    def _make(name="test.db"):
        s = KnowledgeStore(str(tmp_path / name))
        stores.append(s)
        return s

    yield _make
    for s in stores:
        s.close()


# ---------------------------------------------------------------------------
# 1. KnowledgeStore
# ---------------------------------------------------------------------------


class TestKnowledgeStore:
    def test_create_and_get_item(self, store):
        item_id = mk(
            store,
            "Auth Design",
            "JWT tokens with 1h expiry",
            "design_doc",
            summary="Auth overview",
            tags=["auth", "jwt"],
        )
        item = store.get_item(item_id)
        assert item is not None
        assert item["title"] == "Auth Design"
        assert item["content"] == "JWT tokens with 1h expiry"
        assert item["item_type"] == "design_doc"
        assert item["summary"] == "Auth overview"
        # P6b: the API serializes tags as a JSON ARRAY (not a string) + exposes `type`.
        assert item["tags"] == ["auth", "jwt"]
        assert item["type"] == "design_doc"

    def test_fts_search(self, store):
        mk(store, "Auth Design", "JWT tokens with refresh flow", "design_doc")
        mk(store, "Database Schema", "DynamoDB table layout", "design_doc")
        results = store.search_items_fts("JWT")
        assert len(results) >= 1
        assert results[0]["title"] == "Auth Design"

    def test_add_entity_and_relation(self, store):
        e1 = store.add_entity("AuthService", "service", description="Handles auth")
        e2 = store.add_entity("DynamoDB", "technology", description="NoSQL DB")
        rid = store.add_entity_relation(e1, e2, "uses", description="Stores tokens")
        assert rid is not None
        assert store.graph.has_edge(e1, e2)
        edge = store.graph.edges[e1, e2]
        assert edge["relation_type"] == "uses"

    def test_add_entity_relation_is_idempotent(self, store):
        """Re-stating the same relation (same src/tgt/type) — common within one document
        and on re-ingest — must not append duplicate edges."""
        e1 = store.add_entity("FeebasService", "service")
        e2 = store.add_entity("DeployService", "service")
        r1 = store.add_entity_relation(e1, e2, "uses", description="deploy")
        r2 = store.add_entity_relation(e1, e2, "uses", description="deploy again")
        assert r1 == r2  # same edge returned, not a new one
        n = store.db.execute(
            "SELECT COUNT(*) FROM entity_relations WHERE source_id=? AND target_id=? AND relation_type='uses'",  # noqa: E501
            (e1, e2),
        ).fetchone()[0]
        assert n == 1
        # A different relation type between the same pair is still distinct.
        store.add_entity_relation(e1, e2, "depends_on")
        assert (
            store.db.execute(
                "SELECT COUNT(*) FROM entity_relations WHERE source_id=? AND target_id=?", (e1, e2)
            ).fetchone()[0]
            == 2
        )

    def test_entity_subgraph(self, store):
        e1 = store.add_entity("ServiceA", "service")
        e2 = store.add_entity("ServiceB", "service")
        e3 = store.add_entity("Database", "technology")
        store.add_entity_relation(e1, e2, "calls")
        store.add_entity_relation(e2, e3, "uses")
        sg = store.get_entity_subgraph(e1, depth=2)
        node_ids = {n["id"] for n in sg["nodes"]}
        assert e1 in node_ids
        assert e2 in node_ids
        assert e3 in node_ids
        assert len(sg["edges"]) == 2
        # Verify D3.js format: nodes have id/name/type, edges have source/target/type
        for n in sg["nodes"]:
            assert "id" in n and "name" in n and "type" in n
        for e in sg["edges"]:
            assert "source" in e and "target" in e and "type" in e

    def test_serialized_item_omits_raw_vector(self, store):
        """The raw embedding vector is an internal pipeline detail no API consumer reads —
        get_item/list expose only a has_embedding flag, never the bytes."""
        import struct

        emb = struct.pack("4f", 0.1, 0.2, 0.3, 0.4)
        iid = mk(store, "Vec", "body", "note", embedding=emb)

        got = store.get_item(iid)
        assert "embedding" not in got, "raw vector must not ride in responses"
        assert got["has_embedding"] is True

        # An item with no vector reports has_embedding False (no key surprises).
        iid2 = mk(store, "NoVec", "body2", "note")
        assert store.get_item(iid2)["has_embedding"] is False

    def test_delete_item(self, store):
        item_id = mk(store, "Temp Doc", "Will be deleted", "personal_notes")
        assert store.get_item(item_id) is not None
        store.delete_item(item_id)
        assert store.get_item(item_id) is None
        # FTS should also be clean
        assert store.search_items_fts("deleted") == []

    def test_find_entity_case_insensitive(self, store):
        store.add_entity("DynamoDB", "technology")
        found = store.find_entity("dynamodb")
        assert found is not None
        assert found["name"] == "DynamoDB"

    def test_merge_entities(self, store):
        e_keep = store.add_entity("AuthService", "service")
        e_merge = store.add_entity("Auth Service", "service")
        e_other = store.add_entity("Database", "technology")
        store.add_entity_relation(e_merge, e_other, "uses")
        item_id = mk(store, "Doc", "content", "design_doc")
        store.add_mention(item_id, e_merge)

        store.merge_entities(e_keep, e_merge)

        # Merged entity should be gone
        assert store.find_entity("Auth Service") is None
        # Relation should point to kept entity
        rels = store.db.execute(
            "SELECT * FROM entity_relations WHERE source_id = ?", (e_keep,)
        ).fetchall()
        assert len(rels) == 1
        assert rels[0]["target_id"] == e_other
        # Mention should reference kept entity
        mentions = store.db.execute(
            "SELECT * FROM mentions WHERE entity_id = ?", (e_keep,)
        ).fetchall()
        assert len(mentions) == 1


# ---------------------------------------------------------------------------
# 3. FileReader
# ---------------------------------------------------------------------------


class TestFileReader:
    def test_read_markdown(self, tmp_path):
        md = tmp_path / "test.md"
        md.write_text("# Hello\nWorld", encoding="utf-8")
        reader = FileReader()
        text, meta = reader.read(str(md))
        assert "# Hello" in text
        assert "World" in text
        assert meta["format"] == "md"
        assert meta["title"] == "test"
        assert meta["line_count"] == 2

    def test_read_html_strips_site_chrome(self, tmp_path):
        """An uploaded .html file ingests the page CONTENT, not its site chrome — the
        reader strips nav/header/footer (same primitive as the bookmark scrape), so a
        saved web page doesn't pollute its document text/embedding with boilerplate."""
        f = tmp_path / "page.html"
        f.write_text(
            "<html><body><header><nav>Skip to content Sign in</nav></header>"
            "<main><h1>Doc Title</h1><p>Real document body.</p></main>"
            "<footer>(c) Site</footer></body></html>",
            encoding="utf-8",
        )
        text, meta = FileReader().read(str(f))
        assert meta["format"] == "html"
        assert "Doc Title" in text and "Real document body." in text
        assert "Skip to content" not in text and "Sign in" not in text and "Site" not in text

    def test_read_csv_renders_markdown_table(self, tmp_path):
        """A .csv is a 'sheet'-type item — read it as a markdown table (like .xlsx) with a
        row_count, not raw comma text. Quoted commas stay one cell; embedded pipes escape."""
        f = tmp_path / "people.csv"
        f.write_text('Name,Role\n"Smith, Alice",Engineer\nBob,"a | b"\n', encoding="utf-8")
        text, meta = FileReader().read(str(f))
        assert meta["format"] == "csv" and meta["content_type"] == "markdown"
        assert meta["row_count"] == 3
        assert "| Name | Role |" in text and "| --- | --- |" in text
        assert "Smith, Alice" in text  # quoted comma kept in one cell
        assert "a \\| b" in text  # embedded pipe escaped

    def test_read_tsv_uses_tab_delimiter(self, tmp_path):
        """A .tsv is the tab-separated sibling of .csv — rendered as a markdown table with
        the TAB delimiter (so a value containing a comma stays one cell)."""
        f = tmp_path / "data.tsv"
        f.write_text("Name\tNote\nAlice\thello, world\nBob\tfoo\n", encoding="utf-8")
        text, meta = FileReader().read(str(f))
        assert meta["format"] == "tsv" and meta["content_type"] == "markdown"
        assert meta["row_count"] == 3
        assert "| Name | Note |" in text
        assert "hello, world" in text  # the comma is data, not a delimiter → one cell

    def test_read_csv_caps_large_table(self, tmp_path):
        """A large CSV renders a capped table (content/embedding don't bloat) but reports
        the true total row_count."""
        f = tmp_path / "big.csv"
        f.write_text("n\n" + "\n".join(str(i) for i in range(600)) + "\n", encoding="utf-8")
        text, meta = FileReader().read(str(f))
        assert meta["row_count"] == 601  # header + 600 data rows
        assert "more rows" in text  # truncation note present
        assert text.count("\n| ") < 600  # not every row rendered

    def test_read_unsupported(self, tmp_path):
        f = tmp_path / "data.xyz"
        f.write_text("binary-ish", encoding="utf-8")
        reader = FileReader()
        # Unsupported extension still falls through to _read_text
        text, meta = reader.read(str(f))
        assert "binary-ish" in text

    def test_supported_formats(self):
        reader = FileReader()
        for ext in (".md", ".txt", ".py", ".html", ".json", ".yaml", ".csv"):
            assert ext in reader.SUPPORTED, f"{ext} missing from SUPPORTED"

    def test_pdf_reader_dependency_present(self):
        """pdfplumber is a declared core dep (#71): the PDF path must NOT
        degrade to the format:'error' missing-dependency sentinel."""
        import personalclaw.knowledge.readers as rd

        assert rd.pdfplumber is not None, "pdfplumber missing — PDF upload silently yields 0 items"

    def test_read_xlsx_as_markdown_tables(self, tmp_path):
        """A spreadsheet extracts as markdown tables (one per sheet), not binary text."""
        import personalclaw.knowledge.readers as rd

        if rd._load_workbook is None:
            import pytest as _pytest

            _pytest.skip("openpyxl not installed")
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Sales"
        ws.append(["product", "qty"])
        ws.append(["widget", 120])
        p = tmp_path / "report.xlsx"
        wb.save(str(p))
        text, meta = FileReader().read(str(p))
        assert meta["format"] == "xlsx" and meta["sheet_count"] == 1
        assert "## Sales" in text and "| product | qty |" in text and "widget" in text
        assert "PK" not in text[:4]  # not raw zip bytes

    def test_xlsx_in_dispatch_and_supported(self):
        reader = FileReader()
        assert ".xlsx" in reader.SUPPORTED and ".xlsx" in reader._DISPATCH

    def test_read_pdf(self, tmp_path):
        # Minimal hand-crafted single-page PDF with extractable text.
        pdf = (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 144]/Contents 4 0 R"
            b"/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
            b"4 0 obj<</Length 54>>stream\n"
            b"BT /F1 18 Tf 20 100 Td (Hello Knowledge PDF) Tj ET\n"
            b"endstream endobj\n"
            b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"trailer<</Size 6/Root 1 0 R>>\n%%EOF"
        )
        f = tmp_path / "doc.pdf"
        f.write_bytes(pdf)
        text, meta = FileReader().read(str(f))
        assert meta["format"] == "pdf", f"PDF read failed: {meta.get('error')}"
        assert "Hello Knowledge PDF" in text

    def test_mislabeled_text_pdf_is_salvaged(self, tmp_path):
        """A .pdf that's really plain text (mis-saved) is recovered as text rather
        than failing — the user's content isn't lost to a wrong extension."""
        f = tmp_path / "notes.pdf"
        f.write_text("# Meeting notes\n\nDiscussed caching strategy and rollout plan.")
        text, meta = FileReader().read(str(f))
        assert meta["format"] == "text" and meta.get("recovered_from") == "pdf"
        assert "Meeting notes" in text

    def test_binary_pdf_still_errors(self, tmp_path):
        """Genuine binary garbage labeled .pdf still errors — not falsely salvaged
        into unreadable noise."""
        f = tmp_path / "junk.pdf"
        f.write_bytes(bytes(range(256)) * 20)
        _, meta = FileReader().read(str(f))
        assert meta["format"] == "error"

    def test_pptx_and_html_readers_present(self):
        """python-pptx + html2text are declared core deps (#71)."""
        import personalclaw.knowledge.readers as rd

        assert rd.Presentation is not None, "python-pptx missing — .pptx upload yields 0 items"
        assert rd._html2text_mod is not None, "html2text missing — HTML extraction degrades"

    def test_missing_reader_dep_surfaces_error(self, tmp_path, monkeypatch):
        """When a reader dependency is genuinely absent, the reader must return a
        format:'error' sentinel (which the ingest handler turns into a visible
        sync_error) rather than a silent empty/0-item result."""
        import personalclaw.knowledge.readers as rd

        monkeypatch.setattr(rd, "pdfplumber", None)
        f = tmp_path / "x.pdf"
        f.write_bytes(b"%PDF-1.4\n%%EOF")
        _text, meta = FileReader().read(str(f))
        assert meta["format"] == "error"
        assert "pdfplumber" in meta.get("error", "")


# ---------------------------------------------------------------------------
# 4. EntityExtractor
# ---------------------------------------------------------------------------


class TestEntityExtractor:
    def test_extract_no_agent(self):
        import asyncio

        ext = EntityExtractor(pool=None)
        result = asyncio.get_event_loop().run_until_complete(ext.extract("some text"))
        assert result == {
            "title": "",
            "entities": [],
            "relations": [],
            "category": "document",
            "summary": "",
        }

    def test_extract_caps_content_sent_to_llm(self):
        """A large document's full text must not flood the extraction prompt (context-window
        blowout → empty graph). extract() sends only the leading _MAX_CHARS to the model."""
        import asyncio

        from personalclaw.knowledge.extractor import _MAX_CHARS

        sent = {}

        class _Pool:
            async def send(self, prompt, timeout=None):
                sent["prompt"] = prompt
                return '{"entities": [], "relations": [], "category": "document", "summary": ""}'

        # Use a sentinel char absent from the prompt template so the count is exact.
        big = "é" * (_MAX_CHARS * 3)
        asyncio.get_event_loop().run_until_complete(EntityExtractor(pool=_Pool()).extract(big))
        assert sent["prompt"].count("é") == _MAX_CHARS  # capped, not the full 3× body

    def test_parse_json_response(self):
        ext = EntityExtractor()
        raw = json.dumps(
            {
                "entities": [{"name": "Svc", "type": "service", "description": "A service"}],
                "relations": [],
                "category": "design_doc",
                "summary": "A service doc.",
            }
        )
        result = ext._parse_response(raw)
        assert len(result["entities"]) == 1
        assert result["category"] == "design_doc"

    def test_parse_code_block_response(self):
        ext = EntityExtractor()
        raw = '```json\n{"entities": [], "relations": [], "category": "runbook", "summary": "ops"}\n```'  # noqa: E501
        result = ext._parse_response(raw)
        assert result["category"] == "runbook"
        assert result["summary"] == "ops"

    def test_parse_normalizes_malformed_entities_and_relations(self):
        """The LLM sometimes returns entities as bare strings or includes junk/incomplete
        entries. Normalize at parse time so a downstream ent.get('name') never raises (which
        the broad except would turn into dropping the WHOLE item's graph). Bare strings →
        {'name': ...}; unnamed/junk dropped; relations need both source + target."""
        ext = EntityExtractor()
        raw = json.dumps(
            {
                "entities": [
                    "MongoDB",
                    "Redis",
                    {"name": "Kafka", "type": "technology"},
                    {"type": "no-name"},
                    42,
                    {"name": "  "},
                ],
                "relations": [
                    {"source": "Kafka", "target": "MongoDB", "type": "feeds"},
                    "bad",
                    {"source": "A"},
                    {"target": "B"},
                ],
                "category": "document",
                "summary": "s",
            }
        )
        result = ext._parse_response(raw)
        assert result["entities"] == [
            {"name": "MongoDB"},
            {"name": "Redis"},
            {"name": "Kafka", "type": "technology"},
        ]
        assert result["relations"] == [{"source": "Kafka", "target": "MongoDB", "type": "feeds"}]
        # Every surviving entity is a dict with a usable name (no downstream AttributeError).
        assert all(isinstance(e, dict) and e.get("name") for e in result["entities"])

    def test_parse_tolerates_non_list_entities(self):
        """A model returning entities/relations as a non-list (e.g. null or a string) must
        normalize to [] — never crash the parse."""
        ext = EntityExtractor()
        result = ext._parse_response('{"entities": "MongoDB", "relations": null, "summary": "s"}')
        assert result["entities"] == [] and result["relations"] == []


# ---------------------------------------------------------------------------
# 5. HybridRetriever
# ---------------------------------------------------------------------------


class TestHybridRetriever:
    def test_keyword_search(self, store):
        mk(store, "Auth Design", "JWT tokens with refresh flow", "design_doc")
        mk(store, "DB Schema", "DynamoDB table layout", "design_doc")
        retriever = HybridRetriever(store)
        results = retriever.search("JWT")
        assert len(results) >= 1
        assert results[0]["title"] == "Auth Design"
        assert "keyword" in results[0]["match_type"]

    def test_search_excludes_archived(self, store):
        """Archived items are hidden from retrieval (keyword + graph + vector), matching
        the default-list semantics — they shouldn't leak into agent search."""
        iid = store.create_typed_item(
            item_type="note",
            title="Flibberprotocol",
            content="The flibberprotocol is a unique searchable term.",
        )
        retriever = HybridRetriever(store)
        assert any(r["id"] == iid for r in retriever.search("flibberprotocol"))
        # Archive it → gone from search.
        store.update_item(iid, is_archived=1)
        store.db.commit()
        assert not any(r["id"] == iid for r in retriever.search("flibberprotocol"))

    def test_search_include_archived_finds_archived(self, store):
        """The Archived UI view searches WITH include_archived=True — so searching there
        finds archived items (the no-query Archived list shows them; search should match).
        Default (agents, chat context) still excludes them."""
        iid = store.create_typed_item(
            item_type="note",
            title="Zarquon",
            content="The zarquon device is a unique searchable term.",
        )
        store.update_item(iid, is_archived=1)
        store.db.commit()
        retriever = HybridRetriever(store)
        # Default: archived item is hidden.
        assert not any(r["id"] == iid for r in retriever.search("zarquon"))
        # include_archived=True (the Archived view's search): now found.
        assert any(r["id"] == iid for r in retriever.search("zarquon", include_archived=True))

    def test_vector_search_skips_dimension_mismatch(self, store):
        """A stored vector from a DIFFERENT embedding model (different dimension) must be
        skipped, not compared via a truncated-prefix cosine. After a model switch, such
        items fall back to keyword/graph retrieval until re-embedded — they never get a
        meaningless similarity score from zip()-truncated dimensions."""

        def _bytes(vec):
            return struct.pack(f"{len(vec)}f", *vec)

        # Two items: one embedded with the CURRENT model (4-dim), one with an OLD model
        # (6-dim) whose text is otherwise identical so only the vector path differs.
        cur = mk(store, "Current model doc", "alpha bravo", embedding=_bytes([1.0, 0.0, 0.0, 0.0]))
        old = mk(
            store,
            "Old model doc",
            "charlie delta",
            embedding=_bytes([1.0, 0.0, 0.0, 0.0, 9.9, 9.9]),
        )
        # Query embedder returns a 4-dim vector aligned with the current-model item.
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        vec = retriever._vector_search("anything", limit=10)
        ids = {iid for iid, _ in (vec or [])}
        assert cur in ids  # current-dimension item is vector-matched
        assert old not in ids  # mismatched-dimension item is skipped (not prefix-scored)

    def test_search_is_or_not_and(self, store):
        # A conversational query must match a doc that contains only SOME terms
        # (OR semantics) — not require every word (the old implicit-AND default).
        mk(store, "Auth Design", "Access tokens expire; refresh tokens rotate.", "design_doc")
        retriever = HybridRetriever(store)
        results = retriever.search("how do we rotate refresh tokens")
        assert any(r["title"] == "Auth Design" for r in results)

    def test_search_prefix_matches_word_variants(self, store):
        # Prefix matching: "token" should hit "tokens" / "tokenize".
        mk(store, "Caching", "We tokenize the request and cache it.", "note")
        retriever = HybridRetriever(store)
        assert any(r["title"] == "Caching" for r in retriever.search("token"))

    def test_title_match_outranks_longer_doc(self, store):
        """An item whose TITLE is the query must rank above a long document that merely
        mentions the terms in passing — title match is a top relevance signal that raw
        BM25 (favoring term-dense long docs) otherwise buries."""
        # A long doc that mentions 'widget config' deep in a wall of other text.
        mk(
            store,
            "Platform Engineering Handbook",
            "intro " * 200 + " the widget config lives in settings " + "more " * 200,
            "document",
        )
        # A short item literally titled for the query.
        mk(store, "Widget Config", "see the dashboard", "note")
        retriever = HybridRetriever(store)
        results = retriever.search("widget config")
        assert results, "expected matches"
        assert (
            results[0]["title"] == "Widget Config"
        ), f"title match should rank first, got {[r['title'] for r in results]}"

    def test_graph_search_matches_multiword_entity(self, store):
        # A 3-word entity name must be found via the graph (its words may not all
        # appear in any single item's text) — the candidate set includes the full query.
        iid = mk(store, "Doc", "content here", "note")
        eid = store.add_entity("MAPLE Payments team", "org")
        store.add_mention(iid, eid)
        results = HybridRetriever(store).search("MAPLE Payments team")
        hit = next((r for r in results if r["id"] == iid), None)
        assert hit is not None and "graph" in hit["match_type"]

    # ── P12 per-item citation locator (source_type/section/line_range/deep_link) ──

    def test_search_result_carries_citation_locator(self, store):
        """Every hit gains the four locator fields; for a doc with a markdown header
        above the match, section names it, line_range spans it, deep_link points in."""
        mk(
            store,
            "Runbook",
            "# Overview\nboilerplate\n## Rollback Procedure\nrun the rollback script to revert\nmore",  # noqa: E501
            "document",
        )
        results = HybridRetriever(store).search("rollback script")
        hit = next((r for r in results if r["title"] == "Runbook"), None)
        assert hit is not None
        assert hit["source_type"] == "document"
        assert hit["section"] == "Rollback Procedure"  # nearest header above the match
        assert hit["line_range"] and len(hit["line_range"]) == 2
        assert hit["deep_link"].startswith(f"/knowledge/items/{hit['id']}")
        assert "loc=L" in hit["deep_link"]

    def test_locator_is_honest_null_for_structureless_item(self):
        # A structureless item (image, no content) never fabricates a section/line —
        # the fields are null and the deep-link is a bare item route.
        loc = _attach_locator({"id": "img1", "item_type": "image", "content": ""}, {"cat"})
        assert loc["section"] is None and loc["line_range"] is None
        assert loc["source_type"] == "image"
        assert loc["deep_link"] == "/knowledge/items/img1"

    def test_locator_no_query_match_yields_no_line_range(self):
        # Content exists but no query term hits any line → no fabricated span.
        loc = _attach_locator(
            {"id": "n1", "item_type": "note", "content": "alpha beta\ngamma delta"}, {"zeta"}
        )
        assert loc["line_range"] is None and loc["section"] is None
        assert loc["deep_link"] == "/knowledge/items/n1"

    def test_rrf_fuse(self):
        list_a = [("item1", 1), ("item2", 2), ("item3", 3)]
        list_b = [("item2", 1), ("item3", 2), ("item4", 3)]
        fused = HybridRetriever._rrf_fuse(list_a, list_b, None, k=60)
        ids = [item_id for item_id, _ in fused]
        # item2 appears in both lists at good ranks, should be top
        assert ids[0] == "item2"
        # All 4 items should be present
        assert set(ids) == {"item1", "item2", "item3", "item4"}


# ---------------------------------------------------------------------------
# 6. SimpleDiGraph
# ---------------------------------------------------------------------------


class TestSimpleDiGraph:
    def test_add_node_and_has_node(self):
        g = SimpleDiGraph()
        g.add_node("a", name="A")
        assert g.has_node("a")
        assert not g.has_node("b")

    def test_add_edge_and_has_edge(self):
        g = SimpleDiGraph()
        g.add_node("a")
        g.add_node("b")
        g.add_edge("a", "b", weight=1.0)
        assert g.has_edge("a", "b")
        assert not g.has_edge("b", "a")

    def test_successors_predecessors(self):
        g = SimpleDiGraph()
        g.add_node("a")
        g.add_node("b")
        g.add_node("c")
        g.add_edge("a", "b")
        g.add_edge("a", "c")
        assert set(g.successors("a")) == {"b", "c"}
        assert set(g.predecessors("b")) == {"a"}
        assert list(g.successors("c")) == []

    def test_degree(self):
        g = SimpleDiGraph()
        g.add_node("a")
        g.add_node("b")
        g.add_node("c")
        g.add_edge("a", "b")
        g.add_edge("c", "a")
        assert g.degree("a") == 2  # 1 outgoing + 1 incoming
        assert g.degree("b") == 1

    def test_nodes_iteration_and_subscript(self):
        g = SimpleDiGraph()
        g.add_node("x", name="X", entity_type="svc")
        g.add_node("y", name="Y", entity_type="db")
        assert set(g.nodes) == {"x", "y"}
        assert g.nodes["x"]["name"] == "X"
        assert "x" in g.nodes
        assert len(g.nodes) == 2

    def test_edges_iteration_and_subscript(self):
        g = SimpleDiGraph()
        g.add_edge("a", "b", relation_type="calls")
        edges = list(g.edges(data=True))
        assert len(edges) == 1
        assert edges[0] == ("a", "b", {"relation_type": "calls"})
        assert g.edges["a", "b"]["relation_type"] == "calls"

    def test_clear(self):
        g = SimpleDiGraph()
        g.add_node("a")
        g.add_edge("a", "b")
        g.clear()
        assert not g.has_node("a")
        assert not g.has_edge("a", "b")
        assert list(g.nodes) == []


# ---------------------------------------------------------------------------
# 7. KnowledgeStore -- additional coverage
# ---------------------------------------------------------------------------


class TestKnowledgeStoreExtended:
    def test_update_item_fts_sync(self, store):
        item_id = mk(store, "Original", "old content about cats", "doc")
        assert len(store.search_items_fts("cats")) == 1
        store.update_item(item_id, title="Updated", content="new content about dogs")
        # After update, new content should be searchable
        assert len(store.search_items_fts("dogs")) == 1
        item = store.get_item(item_id)
        assert item["title"] == "Updated"
        assert item["content"] == "new content about dogs"

    def test_update_item_no_fields(self, store):
        item_id = mk(store, "Doc", "content", "doc")
        store.update_item(item_id)  # no-op, should not crash

    def test_update_item_non_fts_field(self, store):
        item_id = mk(store, "Doc", "content", "doc")
        store.update_item(item_id, status="archived")
        assert store.get_item(item_id)["status"] == "archived"

    def test_get_item_missing(self, store):
        assert store.get_item("nonexistent") is None

    def test_get_neighbors_depth(self, store):
        e1 = store.add_entity("A", "svc")
        e2 = store.add_entity("B", "svc")
        e3 = store.add_entity("C", "svc")
        store.add_entity_relation(e1, e2, "calls")
        store.add_entity_relation(e2, e3, "calls")
        # depth=1 should get B only
        n1 = store.get_neighbors(e1, depth=1)
        assert {n["id"] for n in n1} == {e2}
        # depth=2 should get B and C
        n2 = store.get_neighbors(e1, depth=2)
        assert {n["id"] for n in n2} == {e2, e3}

    def test_get_neighbors_bidirectional(self, store):
        e1 = store.add_entity("A", "svc")
        e2 = store.add_entity("B", "svc")
        store.add_entity_relation(e2, e1, "calls")
        # e1 has no outgoing but has incoming from e2
        neighbors = store.get_neighbors(e1, depth=1)
        assert {n["id"] for n in neighbors} == {e2}

    def test_find_entity_by_alias(self, store):
        store.add_entity("DynamoDB", "technology", aliases=["ddb", "dynamo"])
        found = store.find_entity("ddb")
        assert found is not None
        assert found["name"] == "DynamoDB"

    def test_find_entity_not_found(self, store):
        assert store.find_entity("nonexistent") is None

    def test_backfill_entity_description_only_when_empty(self, store):
        """An entity first extracted without a description can gain one later; an
        existing description is never clobbered."""
        # No description → backfill writes.
        eid = store.add_entity("Cachemere", "technology")
        assert store.backfill_entity_description(eid, "An internal caching layer") is True
        assert store.find_entity("Cachemere")["description"] == "An internal caching layer"
        # Already has one → never overwritten; blank backfill is a no-op.
        assert store.backfill_entity_description(eid, "a worse later description") is False
        assert store.find_entity("Cachemere")["description"] == "An internal caching layer"
        assert store.backfill_entity_description(eid, "   ") is False

    def test_delete_item_cleans_mentions(self, store):
        item_id = mk(store, "Doc", "content", "doc")
        eid = store.add_entity("Svc", "service")
        store.add_mention(item_id, eid, context="test")
        store.delete_item(item_id)
        assert (
            store.db.execute("SELECT * FROM mentions WHERE item_id = ?", (item_id,)).fetchone()
            is None
        )

    def test_get_stats(self, store):
        mk(store, "A", "a", "doc")
        store.add_entity("E", "svc")
        stats = store.get_stats()
        assert stats["items"] == 1
        assert stats["entities"] == 1
        assert stats["relations"] == 0

    def test_all_tags_frequency_ordered_excludes_archived(self, store):
        """all_tags() powers tag autocomplete: distinct tags, most-frequent first,
        archived items' tags excluded."""
        mk(store, "A", "a", "note", tags=["python", "caching"])
        mk(store, "B", "b", "note", tags=["caching", "redis"])
        arch = mk(store, "C", "c", "note", tags=["archived-only"])
        store.update_item(arch, is_archived=1)
        store.db.commit()
        tags = store.all_tags()
        assert tags[0] == "caching"  # appears twice → first
        assert set(tags) == {"caching", "python", "redis"}
        assert "archived-only" not in tags

    def test_corpus_overview_excludes_archived(self, store):
        """corpus_overview() is the agent's gap-detection view — it must reflect the
        live library (active, non-archived), matching all_tags() and retrieval scope,
        so archived items never inflate total/by_type/top_tags."""
        mk(store, "Keep", "k", "note", tags=["live"])
        arch = mk(store, "Gone", "g", "gist", tags=["stale"])
        store.update_item(arch, is_archived=1)
        store.db.commit()
        ov = store.corpus_overview()
        assert ov["total"] == 1
        assert ov["by_type"] == {"note": 1}  # archived gist excluded
        assert [t["tag"] for t in ov["top_tags"]] == ["live"]
        assert "stale" not in [t["tag"] for t in ov["top_tags"]]

    def test_bookmark_url_dedup_ignores_trailing_slash_and_tracking(self, store):
        """A bookmark's URL is canonicalized at storage, and find_active_by_url
        normalizes its lookup — so trailing-slash and tracking-param variants of one
        page dedup against the original instead of creating duplicates."""
        iid = store.create_typed_item(item_type="bookmark", title="Ex", url="https://example.com")
        # Trailing slash, host casing, and a tracking param all resolve to the same item.
        assert store.find_active_by_url("https://example.com/")["id"] == iid
        assert store.find_active_by_url("https://Example.com/?utm_source=news")["id"] == iid
        # A genuinely different path does NOT dedup.
        assert store.find_active_by_url("https://example.com/other") is None
        # The stored URL is the canonical form.
        assert store.get_item(iid)["url"] == "https://example.com"

    def test_corpus_overview_counts_active_library(self, store):
        """corpus_overview reports the live library: total, by-type, top-tags (active,
        non-archived only)."""
        store.create_typed_item(item_type="note", title="d1", content="x", tags=["shared"])
        store.create_typed_item(item_type="gist", title="a1", content="y", tags=["alpha"])
        store.db.commit()
        ov = store.corpus_overview()
        assert ov["total"] == 2
        assert ov["by_type"] == {"note": 1, "gist": 1}
        assert {t["tag"] for t in ov["top_tags"]} == {"shared", "alpha"}

    def test_graph_has_node(self, store):
        eid = store.add_entity("Svc", "service")
        assert store.graph.has_node(eid)
        assert not store.graph.has_node("fake")

    def test_graph_degree(self, store):
        e1 = store.add_entity("A", "svc")
        e2 = store.add_entity("B", "svc")
        store.add_entity_relation(e1, e2, "calls")
        assert store.graph.degree(e1) == 1
        assert store.graph.degree(e2) == 1

    def test_load_graph_on_reopen(self, tmp_path):
        db_path = str(tmp_path / "reload.db")
        s1 = KnowledgeStore(db_path)
        e1 = s1.add_entity("A", "svc")
        e2 = s1.add_entity("B", "svc")
        s1.add_entity_relation(e1, e2, "calls")
        s1.close()
        s2 = KnowledgeStore(db_path)
        assert s2.graph.has_node(e1)
        assert s2.graph.has_edge(e1, e2)
        s2.close()


# ---------------------------------------------------------------------------
# 8. HybridRetriever -- additional coverage
# ---------------------------------------------------------------------------


class TestHybridRetrieverExtended:
    def test_graph_search(self, store):
        e1 = store.add_entity("JWT", "concept")
        item_id = mk(store, "Auth", "JWT token design", "doc")
        store.add_mention(item_id, e1)
        retriever = HybridRetriever(store)
        results = retriever._graph_search("JWT")
        assert len(results) >= 1
        assert results[0][0] == item_id

    def test_graph_search_no_match(self, store):
        retriever = HybridRetriever(store)
        assert retriever._graph_search("nonexistent") == []

    def test_graph_search_with_neighbors(self, store):
        e1 = store.add_entity("Auth", "service")
        e2 = store.add_entity("JWT", "concept")
        store.add_entity_relation(e1, e2, "uses")
        item_id = mk(store, "Token doc", "about tokens", "doc")
        store.add_mention(item_id, e2)
        retriever = HybridRetriever(store)
        results = retriever._graph_search("Auth")
        assert len(results) >= 1

    def test_vector_search_no_embedder(self, store):
        retriever = HybridRetriever(store, embedder=None)
        assert retriever._vector_search("query") is None

    def test_vector_search_with_embedder(self, store):
        emb = struct.pack("4f", 1.0, 0.0, 0.0, 0.0)
        mk(store, "Vec Doc", "vector content", "doc", embedding=emb)
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        results = retriever._vector_search("query")
        assert results is not None
        assert len(results) == 1

    def test_vector_search_drops_below_similarity_floor(self, store):
        """A near-orthogonal item (weak cosine) is dropped, so a precise query isn't
        polluted by semantic noise; a strongly-similar item is kept."""
        # Item A is orthogonal to the query (sim 0), item B is identical (sim 1).
        mk(store, "Orthogonal", "x", "doc", embedding=struct.pack("4f", 0.0, 1.0, 0.0, 0.0))
        b = mk(store, "Aligned", "y", "doc", embedding=struct.pack("4f", 1.0, 0.0, 0.0, 0.0))
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        results = retriever._vector_search("query")
        ids = [r[0] for r in results]
        assert b in ids and len(ids) == 1  # orthogonal one dropped by the floor

    def test_cosine_similarity_identical(self):
        assert HybridRetriever._cosine_similarity([1, 0], [1, 0]) == pytest.approx(1.0)

    def test_cosine_similarity_orthogonal(self):
        assert HybridRetriever._cosine_similarity([1, 0], [0, 1]) == pytest.approx(0.0)

    def test_cosine_similarity_zero_vector(self):
        assert HybridRetriever._cosine_similarity([0, 0], [1, 1]) == 0.0

    def test_search_combined_match_types(self, store):
        e1 = store.add_entity("JWT", "concept")
        emb = struct.pack("4f", 1.0, 0.0, 0.0, 0.0)
        item_id = mk(store, "JWT Auth", "JWT token design", "doc", embedding=emb)
        store.add_mention(item_id, e1)
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        results = retriever.search("JWT")
        assert len(results) >= 1
        # Should have multiple match types
        mt = results[0]["match_type"]
        assert "keyword" in mt

    def test_search_graph_pair_terms(self, store):
        """Graph search tries consecutive word pairs."""
        e1 = store.add_entity("Auth Service", "service")
        item_id = mk(store, "Doc", "about auth service", "doc")
        store.add_mention(item_id, e1)
        retriever = HybridRetriever(store)
        results = retriever._graph_search("Auth Service details")
        assert len(results) >= 1

    def test_bytes_to_floats_valid(self):
        blob = struct.pack("4f", 1.0, 2.0, 3.0, 4.0)
        result = _bytes_to_floats(blob)
        assert result == pytest.approx([1.0, 2.0, 3.0, 4.0])

    def test_bytes_to_floats_empty(self):
        assert _bytes_to_floats(b"") == []
        assert _bytes_to_floats(None) == []

    def test_bytes_to_floats_invalid(self):
        # Too short / not a clean multiple of 4 bytes → no decode.
        assert _bytes_to_floats(b"not json") == []
        assert _bytes_to_floats(b"abc") == []


# ---------------------------------------------------------------------------
# 9. EntityExtractor -- additional coverage
# ---------------------------------------------------------------------------


class TestEntityExtractorExtended:
    def test_extract_empty_text(self):
        import asyncio

        ext = EntityExtractor(pool=None)
        result = asyncio.get_event_loop().run_until_complete(ext.extract(""))
        assert result == {
            "title": "",
            "entities": [],
            "relations": [],
            "category": "document",
            "summary": "",
        }

    def test_extract_with_agent(self):
        import asyncio

        class MockPool:
            async def send(self, prompt, timeout=60.0):
                return json.dumps(
                    {
                        "entities": [{"name": "Svc", "type": "service", "description": "A"}],
                        "relations": [],
                        "category": "design_doc",
                        "summary": "test",
                    }
                )

            async def send_batch(self, prompts, timeout=60.0):
                return [await self.send(p, timeout) for p in prompts]

        ext = EntityExtractor(pool=MockPool())
        result = asyncio.get_event_loop().run_until_complete(ext.extract("some text"))
        assert result["category"] == "design_doc"
        assert len(result["entities"]) == 1

    def test_extract_agent_exception(self):
        import asyncio

        class BadPool:
            async def send(self, prompt, timeout=60.0):
                raise RuntimeError("fail")

            async def send_batch(self, prompts, timeout=60.0):
                raise RuntimeError("fail")

        ext = EntityExtractor(pool=BadPool())
        result = asyncio.get_event_loop().run_until_complete(ext.extract("text"))
        assert result == {
            "title": "",
            "entities": [],
            "relations": [],
            "category": "document",
            "summary": "",
        }

    def test_parse_response_regex_fallback(self):
        ext = EntityExtractor()
        raw = 'Some preamble text {"entities": [], "relations": [], "category": "runbook", "summary": "ok"} trailing'  # noqa: E501
        result = ext._parse_response(raw)
        assert result["category"] == "runbook"

    def test_parse_response_garbage(self):
        ext = EntityExtractor()
        result = ext._parse_response("totally invalid garbage")
        assert result == {
            "title": "",
            "entities": [],
            "relations": [],
            "category": "document",
            "summary": "",
        }

    def test_extract_code_block(self):
        ext = EntityExtractor()
        assert ext._extract_code_block("no block here") is None
        result = ext._extract_code_block('```\n{"a": 1}\n```')
        assert result == '{"a": 1}'

    def test_validate_partial_data(self):
        ext = EntityExtractor()
        result = ext._validate({"category": "runbook"})
        assert result["entities"] == []
        assert result["relations"] == []
        assert result["summary"] == ""
        assert result["category"] == "runbook"


# ---------------------------------------------------------------------------
# 11. FileReader -- additional coverage
# ---------------------------------------------------------------------------


class TestFileReaderExtended:
    def test_read_html_without_html2text(self, tmp_path):
        """Test HTML reading (exercises html2text or regex fallback)."""
        html_file = tmp_path / "test.html"
        html_file.write_text("<html><body><p>Hello</p></body></html>")
        reader = FileReader()
        text, meta = reader.read(str(html_file))
        assert "Hello" in text

    def test_read_latin1_fallback(self, tmp_path):
        f = tmp_path / "latin.txt"
        f.write_bytes(b"caf\xe9")
        reader = FileReader()
        text, meta = reader.read(str(f))
        assert "caf" in text

    def test_read_json_file(self, tmp_path):
        f = tmp_path / "data.json"
        f.write_text('{"key": "value"}')
        reader = FileReader()
        text, meta = reader.read(str(f))
        assert '"key"' in text
        assert meta["format"] == "json"


class TestPysqlite3Fallback:
    """Verify modules fall back to stdlib sqlite3 when pysqlite3 is unavailable."""

    _MODULES = (
        "personalclaw.knowledge.store",
        "personalclaw.knowledge.retrieval",
        "personalclaw.snapshot",
    )

    def _reload_without_pysqlite3(self, module_name: str):
        """Force-reimport a module with pysqlite3 blocked.

        Restores the import system to exactly its prior state afterward. The
        fresh module created here is a NEW object; any other test module that
        already did ``from <module_name> import <name>`` holds a binding into
        the ORIGINAL object. If the reloaded copy is left installed, a later
        ``monkeypatch.setattr("<module_name>.<attr>", …)`` patches the wrong
        object and is silently defeated. So we snapshot every evicted
        ``sys.modules`` entry AND the parent package's attribute (``import a.b as
        c`` binds ``c`` via the parent attribute, not ``sys.modules``) and put
        both back.
        """
        import sqlite3 as stdlib_sqlite3

        parent_name, _, leaf = module_name.rpartition(".")
        parent_mod = sys.modules.get(parent_name) if parent_name else None
        saved_parent_attr = getattr(parent_mod, leaf, None) if parent_mod is not None else None

        evicted: dict[str, object] = {}
        saved_pysqlite3 = sys.modules.pop("pysqlite3", None)
        # The driver choice now lives in `personalclaw.sqlite_compat` (PR-1), which the
        # consumer imports `sqlite3` FROM. It's cached in sys.modules bound to whatever
        # resolved at first import, so evict it too — otherwise reimporting the consumer
        # with pysqlite3 blocked would still get the cached (pysqlite3) binding and the
        # fallback would never be exercised.
        saved_compat = sys.modules.pop("personalclaw.sqlite_compat", None)
        for mod in list(sys.modules):
            if mod == module_name or mod.startswith(module_name + "."):
                evicted[mod] = sys.modules.pop(mod)

        sys.modules["pysqlite3"] = None  # type: ignore[assignment]
        try:
            mod = importlib.import_module(module_name)
            assert mod.sqlite3 is stdlib_sqlite3
        finally:
            del sys.modules["pysqlite3"]
            if saved_pysqlite3 is not None:
                sys.modules["pysqlite3"] = saved_pysqlite3
            # Drop the reloaded copies, then restore the original objects (including
            # the original sqlite_compat, so later tests see the real driver binding).
            for mod in list(sys.modules):
                if mod == module_name or mod.startswith(module_name + "."):
                    sys.modules.pop(mod)
            sys.modules.pop("personalclaw.sqlite_compat", None)
            if saved_compat is not None:
                sys.modules["personalclaw.sqlite_compat"] = saved_compat
            sys.modules.update(evicted)
            if parent_mod is not None and saved_parent_attr is not None:
                setattr(parent_mod, leaf, saved_parent_attr)

    def test_store_falls_back_to_stdlib_sqlite3(self):
        self._reload_without_pysqlite3("personalclaw.knowledge.store")

    def test_retrieval_falls_back_to_stdlib_sqlite3(self):
        self._reload_without_pysqlite3("personalclaw.knowledge.retrieval")

    def test_snapshot_falls_back_to_stdlib_sqlite3(self):
        self._reload_without_pysqlite3("personalclaw.snapshot")


# ---------------------------------------------------------------------------
# 13. FileReader -- .docx content_type metadata
# ---------------------------------------------------------------------------


class TestDocxContentType:
    def test_docx_returns_content_type_markdown(self, tmp_path):
        """Verify _read_docx sets content_type: markdown in metadata."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        # Create a minimal .docx
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Some content here.")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        reader = FileReader()
        text, meta = reader.read(str(path))
        assert meta.get("content_type") == "markdown"
        assert "# Test Heading" in text
        assert "Some content here." in text

    def test_docx_content_type_in_dispatch(self):
        """Verify .docx is in the dispatch table."""
        reader = FileReader()
        assert ".docx" in reader._DISPATCH


# ---------------------------------------------------------------------------
# 14. Chunk-level vector arm (KL-10)
# ---------------------------------------------------------------------------


#: Deterministic 4-dim fixture vectors. No embedding model is involved anywhere in this
#: section: vectors are written straight into the ``items.embedding`` / ``chunks.embedding``
#: BLOBs so the arithmetic under test (cosine, the similarity floor, the roll-up, RRF and
#: the cliff cut) is exactly reproducible on every machine. A test that could not embed
#: would prove nothing about ranking.
def _v(*vals):
    return struct.pack(f"{len(vals)}f", *vals)


def _fixed_corpus(store):
    """The KL-10 pin corpus: four items, NO chunk rows, hand-written item vectors.

    Returns ``{name: item_id}``. Three items sit at descending cosine to the query
    vector [1,0,0,0] (1.00 / 0.80 / 0.60) and the fourth is orthogonal (0.00) so the
    ``_VECTOR_MIN_SIMILARITY`` floor has something to drop. Content deliberately says
    "tokens" (plural) in A and B so a literal query term "token" does NOT match a line
    there — that is the honest-null locator this atom improves on.
    """
    return {
        "A": mk(
            store,
            "Token rotation policy",
            "Refresh tokens rotate on every use.\nStored server side.",
            "design_doc",
            embedding=_v(1.0, 0.0, 0.0, 0.0),
        ),
        "B": mk(
            store,
            "Token storage design",
            "Tokens live in redis.\nEviction is LRU.",
            "design_doc",
            embedding=_v(0.8, 0.6, 0.0, 0.0),
        ),
        "C": mk(
            store,
            "Token vault runbook",
            "The token vault is rotated quarterly.\nAudit yearly.",
            "note",
            embedding=_v(0.6, 0.8, 0.0, 0.0),
        ),
        "D": mk(
            store,
            "Sourdough baking notes",
            "Bake bread at 230C.",
            "note",
            embedding=_v(0.0, 1.0, 0.0, 0.0),
        ),
    }


#: The exact ``search("token")`` ranking on ``_fixed_corpus`` — captured on the tree
#: BEFORE the chunk arm landed and unchanged by it. Pinning the scores (not just the
#: order) is what makes this a fusion regression test: any change to ``_rrf_fuse``, its
#: ``k``, the title boost or the cliff cut moves these numbers.
_PINNED_TOKEN_RANKING = [("B", 0.048916), ("A", 0.048660), ("C", 0.048395)]

#: The result dict's keys. The atom must not add, drop or rename a field.
_PINNED_RESULT_KEYS = {
    "content",
    "deep_link",
    "id",
    "line_range",
    "match_type",
    "provider",
    "score",
    "section",
    "source_type",
    "summary",
    "title",
}


class TestVectorArmUnchangedOnFixedCorpus:
    """Behaviour pins. Every assertion here passed on the tree before the chunk-level
    vector arm existed and must keep passing after it: a library with no chunk rows is
    the pre-change world, so its ranking, thresholds and result shape are frozen."""

    def test_ranking_scores_and_shape_are_pinned(self, store):
        ids = _fixed_corpus(store)
        names = {v: k for k, v in ids.items()}
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        results = retriever.search("token", limit=10)

        assert [names[r["id"]] for r in results] == [n for n, _ in _PINNED_TOKEN_RANKING]
        for got, (_, expected) in zip(results, _PINNED_TOKEN_RANKING):
            assert got["score"] == pytest.approx(expected, abs=1e-6)
        # Item shape: exactly the same fields, and every one of them still an item field
        # (ids are ITEM ids, never chunk ids).
        for r in results:
            assert set(r.keys()) == _PINNED_RESULT_KEYS
            assert r["id"] in names
            assert r["match_type"] == "keyword+vector"

    def test_similarity_floor_still_drops_the_orthogonal_item(self, store):
        """``_VECTOR_MIN_SIMILARITY`` is unchanged: the 0.00-cosine item never reaches
        the vector list, and the 0.60 one does."""
        ids = _fixed_corpus(store)
        names = {v: k for k, v in ids.items()}
        vec = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])._vector_search(
            "token", limit=20
        )
        assert [names[i] for i, _ in vec] == ["A", "B", "C"]
        assert [rank for _, rank in vec] == [1, 2, 3]

    def test_cliff_cut_still_trims_the_weak_tail(self, store):
        """The relevance cliff is unchanged: four indexed items, three returned."""
        ids = _fixed_corpus(store)
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        assert len(retriever.search("token", limit=10)) == 3
        assert ids["D"] not in {r["id"] for r in retriever.search("token", limit=10)}


def _chunk(store, item_id, *, section, line_start, line_end, vec, text="passage"):
    """Write ONE chunk row with a hand-written embedding, via the real store API."""
    from personalclaw.knowledge.chunking import Chunk

    existing = store.get_chunks(item_id)
    store.replace_chunks(
        item_id,
        [
            Chunk(
                text=c["text"],
                section=c["section"],
                line_start=c["line_start"],
                line_end=c["line_end"],
                chunk_index=c["chunk_index"],
                embedding=None,
            )
            for c in existing
        ]
        + [
            Chunk(
                text=text,
                section=section,
                line_start=line_start,
                line_end=line_end,
                chunk_index=len(existing),
                embedding=_v(*vec) if vec else None,
            )
        ],
    )


class TestChunkVectorArm:
    """The vector arm searches chunk vectors and rolls chunk hits up to their item."""

    #: A document whose query term appears TWICE: once in passing under the first heading,
    #: once in the passage that actually answers the question.
    _DOC = (
        "# Overview\n"
        "The vault is mentioned here in passing.\n"
        "\n"
        "# Credential rotation\n"
        "The vault agent cycles secrets every ninety days.\n"
        "Rotation is automatic.\n"
    )

    def test_ids_are_item_ids_and_chunk_hits_roll_up(self, store):
        """A chunk hit must surface as ONE item-shaped result carrying the ITEM's id — a
        chunk id must never reach fusion or the caller, and three chunks of one document
        must not become three results."""
        iid = mk(store, "Platform runbook", self._DOC, "design_doc", embedding=_v(0.6, 0.8, 0, 0))
        for i in range(3):
            _chunk(store, iid, section=f"S{i}", line_start=1, line_end=2, vec=(1.0, 0.0, 0.0, 0.0))
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])

        vec = retriever._vector_search("credentials", limit=20)
        assert vec == [(iid, 1)]  # three chunks → one item at rank 1
        results = retriever.search("credentials", limit=10)
        assert [r["id"] for r in results] == [iid]
        assert set(results[0].keys()) == _PINNED_RESULT_KEYS  # item shape unchanged

    def test_chunk_locator_is_at_least_as_specific_as_the_item_level_one(self, store):
        """The substantive win, asserted as a BEFORE/AFTER comparison on one corpus.

        Case 1 (semantic-only query): the term scan finds nothing, so today's locator is
        the honest null ``section=None, line_range=None``. The winning chunk knows where it
        sits, so it cites a real section and span.

        Case 2 (term appears twice): today's whole-document scan anchors on the FIRST
        mention — the passing one under "Overview". Narrowed to the winning chunk it anchors
        on the passage that matched, at the identical ±1-line window width, and names the
        right heading. Never coarser, and pointed at the right place.
        """
        iid = mk(store, "Platform runbook", self._DOC, "design_doc", embedding=_v(1.0, 0, 0, 0))
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])

        # BEFORE: item vector only, no chunk rows.
        before_semantic = retriever.search("credentials", limit=10)[0]
        assert before_semantic["section"] is None
        assert before_semantic["line_range"] is None
        assert "?loc=" not in before_semantic["deep_link"]
        before_term = retriever.search("vault", limit=10)[0]
        assert before_term["section"] == "Overview"
        assert before_term["line_range"] == [1, 3]

        # AFTER: the "Credential rotation" passage (lines 4-6) carries a stronger vector.
        _chunk(
            store,
            iid,
            section="Credential rotation",
            line_start=4,
            line_end=6,
            vec=(1.0, 0.0, 0.0, 0.0),
            text="The vault agent cycles secrets every ninety days.\nRotation is automatic.",
        )
        after_semantic = retriever.search("credentials", limit=10)[0]
        assert after_semantic["section"] == "Credential rotation"  # was None
        assert after_semantic["line_range"] == [4, 6]  # was None
        assert after_semantic["deep_link"].endswith("?loc=L4-6")
        after_term = retriever.search("vault", limit=10)[0]
        assert after_term["section"] == "Credential rotation"  # was the passing "Overview"
        assert after_term["line_range"] == [4, 6]  # was [1, 3] — same width, right passage
        # Specificity is monotone: the span never widens.
        for before, after in ((before_term, after_term), (before_semantic, after_semantic)):
            if before["line_range"]:
                b = before["line_range"][1] - before["line_range"][0]
                assert (after["line_range"][1] - after["line_range"][0]) <= b

    def test_a_wrong_dimension_chunk_vector_is_skipped_not_truncated(self, store):
        """The dimension guard is load-bearing on the chunk path, not tidy-up.

        A half-re-embedded library holds chunks from the OLD model beside the new query
        vector. Cosine over ``zip()`` truncates to the shorter sequence, so an unguarded
        5-dim chunk would score a meaningless 4-dim prefix — a perfect 1.0 here — and win
        the search on nonsense. The vector must be skipped so the item falls back to
        keyword/graph retrieval until it is re-embedded.

        Note the fixture width: ``_bytes_to_floats`` drops any blob under 16 bytes on its
        own, so a 2-float vector would never reach the dimension comparison and the rail
        would pass with the guard removed.
        """
        stale = mk(store, "Stale-model doc", self._DOC, "design_doc")
        _chunk(
            store,
            stale,
            section="Credential rotation",
            line_start=4,
            line_end=6,
            vec=(1.0, 0.0, 0.0, 0.0, 0.0),
        )
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])

        assert retriever._vector_search("credentials", limit=20) == []

    def test_mixed_chunked_and_unchunked_items_are_both_retrievable(self, store):
        """A partially-chunked library (the mid-backfill state) must degrade, never break:
        the chunked item is found through its chunk, the unchunked one through its
        whole-item vector, in the same search."""
        chunked = mk(store, "Chunked doc", self._DOC, "design_doc", embedding=_v(0, 1.0, 0, 0))
        _chunk(
            store,
            chunked,
            section="Credential rotation",
            line_start=4,
            line_end=6,
            vec=(1.0, 0.0, 0.0, 0.0),
        )
        plain = mk(
            store,
            "Unchunked doc",
            "No chunks exist for this one.",
            "note",
            embedding=_v(0.9, 0.436, 0.0, 0.0),
        )
        # A third item has chunk ROWS but no chunk vectors (embedder unavailable at ingest)
        # — it must still be reachable through its whole-item vector.
        vectorless = mk(
            store, "Vectorless chunks", "Body text here.", "note", embedding=_v(0.8, 0.6, 0.0, 0.0)
        )
        _chunk(store, vectorless, section=None, line_start=1, line_end=1, vec=None)

        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        ranked = retriever._vector_search("anything", limit=20)
        assert [i for i, _ in ranked] == [chunked, plain, vectorless]
        # The chunked item's score came from its chunk (its own vector is orthogonal, 0.0,
        # which the floor would have dropped), so chunks genuinely added recall here.
        assert {r["id"] for r in retriever.search("anything", limit=10)} >= {chunked, plain}

    def test_rollup_is_max_so_many_mediocre_chunks_never_beat_one_strong_one(self, store):
        """Falsifies sum/count aggregation: six 0.5-similarity chunks must not outrank one
        1.0-similarity chunk. Under a summing roll-up the six would score 3.0 and win."""
        strong = mk(store, "Strong single passage", "a", "note")
        _chunk(store, strong, section="S", line_start=1, line_end=1, vec=(1.0, 0.0, 0.0, 0.0))
        many = mk(store, "Many mediocre passages", "b", "note")
        for i in range(6):
            _chunk(
                store, many, section=f"M{i}", line_start=1, line_end=1, vec=(0.5, 0.866, 0.0, 0.0)
            )
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        assert [i for i, _ in retriever._vector_search("q", limit=20)] == [strong, many]

    def test_weak_chunks_neither_lift_an_item_nor_supply_a_locator(self, store):
        """``_VECTOR_MIN_SIMILARITY`` applies per chunk, before the roll-up: an item whose
        every chunk is below the floor gets no vector hit from them, and a below-floor chunk
        can never become the cited passage of an item that matched some other way."""
        iid = mk(store, "Weak chunks only", "line one\nline two\nline three", "note")
        _chunk(store, iid, section="Noise", line_start=3, line_end=3, vec=(0.0, 1.0, 0.0, 0.0))
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        assert retriever._vector_search("q", limit=20) == []
        # Reachable by keyword; its locator must NOT come from the near-orthogonal chunk.
        hit = retriever.search("three", limit=10)[0]
        assert hit["id"] == iid and hit["section"] is None and hit["line_range"] == [2, 3]

    def test_chunk_vectors_from_another_model_are_skipped_not_prefix_scored(self, store):
        """The dimension guard covers chunk vectors too. A half-re-embedded library has
        old-model chunk rows; comparing them via zip()-truncated cosine would score a
        meaningless prefix. The item's own current-model vector still carries it."""
        iid = mk(store, "Half re-embedded", "body", "note", embedding=_v(1.0, 0.0, 0.0, 0.0))
        _chunk(
            store, iid, section="Old", line_start=1, line_end=1, vec=(1.0, 0.0, 0.0, 0.0, 9.9, 9.9)
        )
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        assert retriever._vector_search("q", limit=20) == [(iid, 1)]
        # The 6-dim chunk did not win, so it supplied no locator.
        assert retriever.search("q", limit=10)[0]["section"] is None

    def test_archived_and_inactive_items_chunks_are_excluded(self, store):
        """The chunk arm honours the same visibility rails as the item arm — a chunk cannot
        smuggle an archived item into a default search."""
        iid = mk(store, "Archived doc", "secret body", "note")
        _chunk(store, iid, section="S", line_start=1, line_end=1, vec=(1.0, 0.0, 0.0, 0.0))
        store.update_item(iid, is_archived=1)
        store.db.commit()
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        assert retriever._vector_search("q", limit=20) == []
        assert retriever._vector_search("q", limit=20, include_archived=True) == [(iid, 1)]


# ── KL-11: the chunk ANN index (sqlite-vec) ────────────────────────────────────
#
# The stated recall tolerance. The ANN path is a CANDIDATE GENERATOR: `_consider` still does
# the scoring, so ANN and the exact scan cannot disagree on a similarity value — the only way
# they can differ is candidate truncation. 0.95 is the contract this task promises; the design
# is expected to do better than that (the tests below assert the observed 1.00 as well, so a
# regression to a merely-tolerable 0.96 still fails), and the plan's sharpest named risk is a
# silent recall regression, so the looser number is a floor, never a target.
_ANN_RECALL_TOLERANCE = 0.95


def _unit(cos: float) -> bytes:
    """A 4-dim vector whose cosine against the query [1,0,0,0] is exactly ``cos``."""
    import math

    return _v(cos, math.sqrt(max(0.0, 1.0 - cos * cos)), 0.0, 0.0)


def _collapsing_corpus(store, *, items=40, per_item=12):
    """A corpus where the ANN candidate set COLLAPSES: every chunk clears the similarity
    floor, and one item owns a whole run of the top of the ranking.

    This is the shape that makes ANN diverge from an exact scan, and the reason a 3-row corpus
    proves nothing: with 12 chunks per item, the first 80 candidates come from only ~7 items,
    so a single-shot ANN query cannot fill a 20-item request. Returns the item ids in
    descending best-chunk similarity, i.e. the exact scan's expected ranking.
    """
    from personalclaw.knowledge.chunking import Chunk

    ordered = []
    for i in range(items):
        iid = mk(store, f"doc {i}", f"body of document {i}", "note")
        top = 0.99 - i * 0.015  # 0.99 down to 0.405 — every one above the 0.25 floor
        store.replace_chunks(
            iid,
            [
                Chunk(
                    text=f"passage {j} of doc {i}",
                    section=f"S{j}",
                    line_start=j * 3 + 1,
                    line_end=j * 3 + 3,
                    chunk_index=j,
                    embedding=_unit(top - j * 0.0005),
                )
                for j in range(per_item)
            ],
        )
        ordered.append(iid)
    return ordered


def _recall(ann: list, exact: list, k: int) -> float:
    """Fraction of the exact scan's top-``k`` items that the ANN path also returned in its
    top-``k``."""
    want = [i for i, _ in exact[:k]]
    got = {i for i, _ in ann[:k]}
    return (sum(1 for i in want if i in got) / len(want)) if want else 1.0


@pytest.fixture()
def clean_vec_probe():
    """The sqlite-vec capability probe is cached process-wide and its INFO log fires once, so
    every test that touches either resets both sides."""
    from personalclaw.knowledge import vector_index

    vector_index.reset_probe()
    yield vector_index
    vector_index.reset_probe()


class TestChunkAnnIndex:
    """H1.4 — a vec0 index over chunk vectors, fail-soft to the exact scan."""

    def test_sqlite_vec_loads_in_this_environment(self, clean_vec_probe):
        """sqlite-vec is a CORE dependency, so it must resolve here. Asserted on its own so a
        build that cannot load extensions produces ONE failure naming the reason, instead of a
        skip that would let every ANN assertion below silently pass."""
        cap = clean_vec_probe.probe()
        assert cap.available, f"sqlite-vec did not load: {cap.reason}"
        assert cap.version

    def test_the_index_is_built_and_covers_every_embedded_chunk(self, store, clean_vec_probe):
        """The index lives in the SAME database file and is written through by the store."""
        _collapsing_corpus(store, items=3, per_item=4)
        cov = store.vec_index.coverage()
        assert cov["extension_available"] and cov["loaded"]
        assert cov["dimensions"] == {"4": {"indexed": 12, "live": 12}}
        names = [
            r[0]
            for r in store.db.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name='chunk_vec_4'"
            )
        ]
        assert names == ["chunk_vec_4"], "the ANN index must live inside knowledge.db"

    def test_ann_matches_the_exact_scan_within_the_stated_tolerance(self, store, clean_vec_probe):
        """THE recall gate. ANN vs exact on a corpus built to collapse the candidate set, and
        the escalation is asserted to have actually fired — an ANN run that never escalated
        would agree with exact trivially and prove nothing."""
        ordered = _collapsing_corpus(store)
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])

        calls = []
        real = store.vec_index.candidate_chunk_ids

        def counting(blob, dim, k):
            calls.append(k)
            return real(blob, dim, k)

        store.vec_index.candidate_chunk_ids = counting
        ann = retriever._vector_search("q", limit=20)
        store.vec_index.candidate_chunk_ids = real

        assert len(calls) > 1, f"the candidate set never escalated (k tried: {calls})"
        # The exact scan, from the same store with the index refused.
        store.vec_index.candidate_chunk_ids = lambda blob, dim, k: None
        exact = retriever._vector_search("q", limit=20)
        store.vec_index.candidate_chunk_ids = real

        assert [i for i, _ in exact] == ordered[:20], "the exact scan itself must rank by cosine"
        for k in (1, 5, 10, 20):
            assert _recall(ann, exact, k) >= _ANN_RECALL_TOLERANCE
        # Stronger than the tolerance, and asserted so a regression to "merely tolerable" fails.
        assert ann == exact

    def test_without_escalation_the_same_corpus_loses_recall(self, store, monkeypatch):
        """The escalation is load-bearing, not decoration: capped at one ANN query the SAME
        corpus falls below the tolerance. This is the truncation failure mode the tolerance
        test exists to catch, made visible."""
        from personalclaw.knowledge import retrieval as retr

        _collapsing_corpus(store)
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        exact_index = store.vec_index.candidate_chunk_ids
        store.vec_index.candidate_chunk_ids = lambda blob, dim, k: None
        exact = retriever._vector_search("q", limit=20)
        store.vec_index.candidate_chunk_ids = exact_index

        monkeypatch.setattr(retr, "_ANN_MAX_ATTEMPTS", 1)
        monkeypatch.setattr(retr, "_ANN_OVERFETCH", 1)
        truncated = retriever._vector_search("q", limit=20)
        assert _recall(truncated, exact, 20) < _ANN_RECALL_TOLERANCE

    def test_force_disabled_extension_still_returns_correct_results(
        self, store_factory, monkeypatch, clean_vec_probe
    ):
        """Fail SOFT, never closed: with extension loading refused the way a SQLite built
        without it presents, search keeps working through the exact scan and returns the same
        ranking — just at the old cost."""
        store = store_factory("with_index.db")
        ordered = _collapsing_corpus(store, items=12, per_item=4)
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        with_ann = retriever._vector_search("q", limit=10)
        assert store.vec_index.enabled

        def no_extension(conn):
            raise AttributeError("forced: sqlite3 built without enable_load_extension")

        monkeypatch.setattr(clean_vec_probe, "_load_extension", no_extension)
        clean_vec_probe.reset_probe()
        degraded = store_factory("degraded.db")
        assert degraded.vec_index.enabled is False
        _collapsing_corpus(degraded, items=12, per_item=4)  # writes must not raise either
        exact = HybridRetriever(degraded, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])._vector_search(
            "q", limit=10
        )
        assert [r for _, r in exact] == [r for _, r in with_ann]
        assert len(exact) == 10
        # Same corpus shape, so the same TITLES come back in the same order (ids differ per DB).
        assert [degraded.get_item(i)["title"] for i, _ in exact] == [
            store.get_item(i)["title"] for i, _ in with_ann
        ]
        assert [store.get_item(i)["title"] for i, _ in with_ann][0] == store.get_item(ordered[0])[
            "title"
        ]

    def test_the_probe_is_cached_and_the_degradation_logs_once(
        self, store_factory, monkeypatch, caplog, clean_vec_probe
    ):
        """The probe runs ONCE however many searches happen, and the degradation is announced
        exactly one time — a per-search INFO line would be its own defect."""
        attempts = []

        def no_extension(conn):
            attempts.append(1)
            raise AttributeError("forced: no enable_load_extension")

        monkeypatch.setattr(clean_vec_probe, "_load_extension", no_extension)
        clean_vec_probe.reset_probe()
        with caplog.at_level("INFO", logger="personalclaw.knowledge.vector_index"):
            store = store_factory("cached_probe.db")
            _collapsing_corpus(store, items=4, per_item=3)
            retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
            for _ in range(5):
                retriever._vector_search("q", limit=10)
        assert len(attempts) == 1, f"the probe ran {len(attempts)} times, not once"
        lines = [
            r
            for r in caplog.records
            if r.name == "personalclaw.knowledge.vector_index" and r.levelname == "INFO"
        ]
        assert len(lines) == 1, [r.getMessage() for r in lines]
        assert "exact scan" in lines[0].getMessage()

    def test_the_index_stays_in_step_through_rechunk_delete_and_clear(self, store):
        """Staleness in the other costume: a re-chunk mints new chunk ids, so an index that
        only removed the ids it just wrote would keep every previous generation as an orphan
        candidate."""
        from personalclaw.knowledge.chunking import Chunk

        ids = _collapsing_corpus(store, items=3, per_item=4)
        dims = lambda: store.vec_index.coverage()["dimensions"]["4"]  # noqa: E731
        assert dims() == {"indexed": 12, "live": 12}
        store.replace_chunks(
            ids[0],
            [
                Chunk(
                    text="only",
                    section="S",
                    line_start=1,
                    line_end=1,
                    chunk_index=0,
                    embedding=_unit(0.9),
                )
            ],
        )
        assert dims() == {"indexed": 9, "live": 9}
        store.clear_chunks(ids[1])
        assert dims() == {"indexed": 5, "live": 5}
        store.delete_item(ids[2])
        assert dims() == {"indexed": 1, "live": 1}

    def test_a_database_written_without_the_index_is_reconciled_on_the_next_search(
        self, store, clean_vec_probe, caplog
    ):
        """The pre-KL-11 / wrote-while-degraded case: chunk rows with no index entries. The
        first search must rebuild rather than quietly return fewer results."""
        _collapsing_corpus(store, items=6, per_item=4)
        store.db.execute("DELETE FROM chunk_vec_4 WHERE chunk_id IN (SELECT id FROM chunks)")
        store.vec_index._synced.clear()
        assert store.vec_index.coverage()["dimensions"]["4"]["indexed"] == 0
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        with caplog.at_level("INFO", logger="personalclaw.knowledge.vector_index"):
            found = retriever._vector_search("q", limit=6)
        assert len(found) == 6
        assert store.vec_index.coverage()["dimensions"]["4"] == {"indexed": 24, "live": 24}
        assert any("rebuilding" in r.getMessage() for r in caplog.records)

    def test_a_stale_extra_candidate_is_dropped_rather_than_returned(self, store):
        """An index entry whose chunk row is gone must not become a result — the reader joins
        candidates back to the live table, so a stale EXTRA costs a slot, never correctness."""
        ids = _collapsing_corpus(store, items=3, per_item=2)
        # Delete the chunk rows behind the index's back (no drop_item), then pin the count so
        # reconciliation does not repair it before the read path is exercised.
        store.db.execute("DELETE FROM chunks WHERE item_id = ?", (ids[0],))
        store.db.commit()
        store.vec_index._synced.add(4)
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        found = retriever._vector_search("q", limit=10)
        assert ids[0] not in {i for i, _ in found}
        assert {i for i, _ in found} == set(ids[1:])

    def test_mixed_dimension_chunks_get_their_own_index_and_never_cross(self, store):
        """A half-re-embedded library keeps one self-consistent index per dimension, matching
        the reader's dimension guard: a 6-dim chunk is unscoreable against a 4-dim query
        either way, so it must not be indexed alongside one."""
        from personalclaw.knowledge.chunking import Chunk

        ids = _collapsing_corpus(store, items=2, per_item=2)
        store.replace_chunks(
            ids[0],
            [
                Chunk(
                    text="old model",
                    section="S",
                    line_start=1,
                    line_end=1,
                    chunk_index=0,
                    embedding=_v(1.0, 0.0, 0.0, 0.0, 9.9, 9.9),
                )
            ],
        )
        cov = store.vec_index.coverage()["dimensions"]
        assert cov == {"4": {"indexed": 2, "live": 2}, "6": {"indexed": 1, "live": 1}}
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0, 0.0])
        assert {i for i, _ in retriever._vector_search("q", limit=10)} == {ids[1]}

    def test_faiss_stays_an_extra_and_sqlite_vec_is_core(self):
        """The dependency ruling, asserted: sqlite-vec joins core, faiss does NOT move."""
        import pathlib
        import re

        text = (pathlib.Path(__file__).resolve().parents[1] / "pyproject.toml").read_text()

        def requirements(block: str) -> list[str]:
            """The REQUIREMENT strings of a dependency block, with comment lines dropped —
            both blocks explain in prose why faiss is not core, and a substring scan that
            counted that prose would fail on the very comment that documents the rule."""
            body = re.search(rf"^{block} = \[(.*?)^\]", text, re.S | re.M).group(1)
            return [
                line.strip() for line in body.splitlines() if line.strip().startswith(('"', "'"))
            ]

        core = requirements("dependencies")
        assert any(r.startswith('"sqlite-vec>=0.1,<1"') for r in core)
        assert not any("faiss" in r for r in core), "faiss must stay in the [embeddings] extra"
        assert any("faiss" in r for r in requirements("embeddings"))


class TestDoctorVectorIndexLine:
    """The Doctor capability line — the honest-degradation surface for a stripped SQLite."""

    def _run(self, home):
        import asyncio

        from personalclaw.resilience.doctor import (
            DoctorContext,
            _probe_knowledge_vector_index,
        )

        return asyncio.run(_probe_knowledge_vector_index(DoctorContext(home=home)))

    def test_reports_the_active_index(self, tmp_path, clean_vec_probe):
        db_dir = tmp_path / "workspace" / "knowledge"
        db_dir.mkdir(parents=True)
        s = KnowledgeStore(str(db_dir / "knowledge.db"))
        _collapsing_corpus(s, items=3, per_item=4)
        s.close()
        res = self._run(tmp_path)
        assert res.ok and "12 chunk vector(s) indexed" in res.detail
        assert res.evidence["extension_available"] is True
        assert res.evidence["dimensions"] == {"4": {"indexed": 12, "live": 12}}
        assert "degraded" not in res.evidence

    def test_reports_the_degraded_line_when_the_extension_cannot_load(
        self, tmp_path, monkeypatch, clean_vec_probe
    ):
        def no_extension(conn):
            raise AttributeError("forced: no enable_load_extension")

        monkeypatch.setattr(clean_vec_probe, "_load_extension", no_extension)
        clean_vec_probe.reset_probe()
        res = self._run(tmp_path)
        # Degraded is NOT failed: a correct-but-slower search must not read as an outage.
        assert res.ok is True
        assert res.evidence["degraded"] is True
        assert res.evidence["extension_available"] is False
        assert "exact scan" in res.detail
        assert "sqlite-vec" in res.evidence["remedy"]
